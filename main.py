"""
Doqurix - Intelligent Document Analysis
Using llama-cpp-python for efficient CPU inference with quantization
Best-in-class RAG implementation
"""

import os
import sys
from pathlib import Path
from llama_cpp import Llama
from sentence_transformers import SentenceTransformer, CrossEncoder
import chromadb
from chromadb.config import Settings
import PyPDF2
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
from rank_bm25 import BM25Okapi
import numpy as np
import re
from huggingface_hub import hf_hub_download
import urllib.request
import time
import hashlib
import json
from datetime import datetime, timedelta
try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Word document support disabled.")


# ============================================================================
# AGENTIC WORKFLOWS SYSTEM - LOCAL & CPU-ONLY
# ============================================================================

class AgentWorkflows:
    """
    Lightweight AI agents for automating multi-document tasks.
    Privacy-first, local-only processing without external dependencies.
    """
    
    def __init__(self, llm, embedder, reranker):
        self.llm = llm
        self.embedder = embedder
        self.reranker = reranker
        self.document_cache = {}  # Cache for document metadata
        self.language_cache = {}  # Cache for detected languages
        
    def auto_triage_documents(self, documents, query):
        """
        Automatically triage documents based on relevance and content type.
        Returns prioritized list with confidence scores.
        """
        if not documents:
            return []
            
        # Quick semantic relevance scoring
        query_embedding = self.embedder.encode(query)
        scored_docs = []
        
        for doc in documents:
            doc_text = doc.get('content', '')[:1000]  # First 1000 chars for efficiency
            doc_embedding = self.embedder.encode(doc_text)
            
            # Calculate semantic similarity
            similarity = np.dot(query_embedding, doc_embedding) / (
                np.linalg.norm(query_embedding) * np.linalg.norm(doc_embedding)
            )
            
            # Document type detection (technical, legal, financial, etc.)
            doc_type = self._detect_document_type(doc_text)
            
            # Content density scoring
            density_score = self._calculate_content_density(doc_text)
            
            # Combined triage score
            triage_score = similarity * 0.6 + density_score * 0.4
            
            scored_docs.append({
                'document': doc,
                'relevance_score': float(similarity),
                'content_density': float(density_score),
                'document_type': doc_type,
                'triage_score': float(triage_score),
                'priority': 'high' if triage_score > 0.7 else 'medium' if triage_score > 0.4 else 'low'
            })
        
        # Sort by triage score
        return sorted(scored_docs, key=lambda x: x['triage_score'], reverse=True)
    
    def citation_aware_refinement(self, answer, source_contexts):
        """
        Refine answers with accurate citations and cross-references.
        Local-only citation validation and enhancement.
        """
        if not source_contexts:
            return answer
            
        # Extract existing citations from answer
        citation_pattern = r'\[(\d+)\]'
        citations = re.findall(citation_pattern, answer)
        
        # Build citation map
        citation_map = {}
        for i, ctx in enumerate(source_contexts, 1):
            source_name = ctx.get('metadata', {}).get('source', f'Source {i}')
            page = ctx.get('metadata', {}).get('page', 'Unknown')
            citation_map[str(i)] = {
                'source': source_name,
                'page': page,
                'content_snippet': ctx.get('doc', '')[:200] + '...'
            }
        
        # Enhance answer with citation details
        enhanced_answer = answer
        
        # Add citation validation prompt
        validation_prompt = f"""<|im_start|>system
You are a citation validator. Review the answer and ensure all claims are properly supported by the provided sources.
Add [VERIFY] tags to unsupported claims.<|im_end|>
<|im_start|>user
Answer: {answer}

Sources available: {len(source_contexts)}
Validate citations and mark unsupported claims.<|im_end|>
<|im_start|>assistant
"""
        
        try:
            validation_result = self.llm(
                validation_prompt,
                max_tokens=300,
                temperature=0.3,
                stop=["<|im_end|>", "<|im_start|>"]
            )
            
            validated_answer = validation_result['choices'][0]['text'].strip()
            
            # Use validated version if it's more accurate
            if '[VERIFY]' not in validated_answer and len(validated_answer) > len(answer) * 0.8:
                enhanced_answer = validated_answer
                
        except Exception:
            pass  # Fall back to original answer
        
        return {
            'refined_answer': enhanced_answer,
            'citation_map': citation_map,
            'validation_status': 'validated' if '[VERIFY]' not in enhanced_answer else 'needs_review'
        }
    
    def cross_language_retrieval(self, query, documents):
        """
        Intelligent cross-language document retrieval.
        Detects query language and searches across multilingual content.
        """
        # Detect query language
        query_lang = self._detect_language(query)
        
        # Group documents by language
        lang_groups = {}
        for doc in documents:
            content = doc.get('content', '')
            doc_lang = self._detect_language(content[:500])
            
            if doc_lang not in lang_groups:
                lang_groups[doc_lang] = []
            lang_groups[doc_lang].append(doc)
        
        # Multi-language search strategy
        results = []
        
        # 1. Direct language match (highest priority)
        if query_lang in lang_groups:
            direct_matches = self._semantic_search(query, lang_groups[query_lang])
            for match in direct_matches:
                match['language_match'] = 'direct'
                match['priority_score'] = match.get('score', 0) * 1.2
            results.extend(direct_matches)
        
        # 2. English fallback (medium priority)
        if 'english' in lang_groups and query_lang != 'english':
            english_matches = self._semantic_search(query, lang_groups['english'])
            for match in english_matches:
                match['language_match'] = 'english_fallback'
                match['priority_score'] = match.get('score', 0) * 1.0
            results.extend(english_matches[:3])  # Limit fallback results
        
        # 3. Cross-language semantic search (lower priority)
        other_langs = [lang for lang in lang_groups.keys() 
                      if lang not in [query_lang, 'english']]
        
        for lang in other_langs:
            if len(results) < 10:  # Limit total results
                cross_matches = self._semantic_search(query, lang_groups[lang])
                for match in cross_matches[:2]:  # Limited cross-language matches
                    match['language_match'] = f'cross_{lang}'
                    match['priority_score'] = match.get('score', 0) * 0.8
                results.extend(cross_matches[:2])
        
        # Sort by priority score
        return sorted(results, key=lambda x: x.get('priority_score', 0), reverse=True)
    
    def generate_followup_questions(self, query, answer, source_contexts):
        """
        Generate intelligent follow-up questions based on the conversation context.
        Identifies information gaps and suggests deeper exploration paths.
        """
        # Analyze the current answer for potential follow-ups
        context_summary = self._summarize_contexts(source_contexts)
        
        followup_prompt = f"""<|im_start|>system
You are an intelligent research assistant. Based on the user's question and the answer provided, generate 3-4 insightful follow-up questions that would help the user explore the topic deeper. Focus on:
1. Clarifying ambiguous points
2. Exploring related concepts
3. Identifying practical applications
4. Uncovering additional insights from the documents

Be concise and specific.<|im_end|>
<|im_start|>user
Original Question: {query}

Answer: {answer}

Available Context: {context_summary}

Generate follow-up questions:<|im_end|>
<|im_start|>assistant
"""
        
        try:
            result = self.llm(
                followup_prompt,
                max_tokens=200,
                temperature=0.7,
                stop=["<|im_end|>", "<|im_start|>"]
            )
            
            followup_text = result['choices'][0]['text'].strip()
            
            # Parse questions from the response
            questions = []
            lines = followup_text.split('\n')
            
            for line in lines:
                line = line.strip()
                if line and ('?' in line):
                    # Clean up question formatting
                    question = re.sub(r'^\d+[\.\)]\s*', '', line)  # Remove numbering
                    question = question.strip()
                    if question:
                        questions.append(question)
            
            return questions[:4]  # Limit to 4 questions
            
        except Exception:
            # Fallback: rule-based question generation
            return self._generate_fallback_questions(query, answer)
    
    def auto_document_insights(self, documents):
        """
        Automatically extract key insights and themes from document collection.
        Provides intelligent, Claude/ChatGPT-style overview using LLM.
        """
        if not documents:
            return {'themes': [], 'key_entities': [], 'summary': '', 'intelligent_summary': ''}
        
        # Sample documents for analysis (use first 5 for comprehensive analysis)
        sample_docs = documents[:5] if len(documents) > 5 else documents
        
        # Build representative text from samples
        sample_texts = []
        for doc in sample_docs:
            content = doc.get('content', '')
            # Take meaningful chunks (beginning and key sections)
            if len(content) > 2000:
                # Take beginning, middle, and end for better representation
                chunk1 = content[:700]
                chunk2 = content[len(content)//2:len(content)//2 + 600]
                chunk3 = content[-700:]
                combined = f"{chunk1}... {chunk2}... {chunk3}"
                sample_texts.append(combined)
            else:
                sample_texts.append(content)
        
        combined_sample = "\n\n---\n\n".join(sample_texts)
        
        # Use LLM for intelligent analysis
        analysis_prompt = f"""<|im_start|>system
You are an expert document analyst. Analyze the provided document collection and create a comprehensive, professional summary that identifies:
1. Main topics and themes
2. Document type and purpose (legal, technical, financial, etc.)
3. Key stakeholders, parties, or entities mentioned
4. Critical dates, deadlines, or timelines
5. Important obligations, requirements, or actions
6. Risk factors or notable clauses
7. Overall significance and context

Be specific, actionable, and professional. Format the response clearly.<|im_end|>
<|im_start|>user
Analyze this document collection ({len(documents)} document chunks total, analyzing representative samples):

{combined_sample[:8000]}

Provide a comprehensive analysis in a clear, professional format.<|im_end|>
<|im_start|>assistant
"""
        
        try:
            # Generate intelligent analysis using LLM
            result = self.llm(
                analysis_prompt,
                max_tokens=800,
                temperature=0.3,  # Lower temperature for factual analysis
                stop=["<|im_end|>", "<|im_start|>"],
                top_p=0.9
            )
            
            intelligent_summary = result['choices'][0]['text'].strip()
            
            # Also extract structured data for backward compatibility
            themes = self._extract_themes_llm(combined_sample[:3000])
            entities = self._extract_entities_enhanced(combined_sample[:3000])
            
            return {
                'intelligent_summary': intelligent_summary,
                'themes': themes,
                'key_entities': entities,
                'document_count': len(documents),
                'analysis_timestamp': datetime.now().isoformat(),
                'analyzed_samples': len(sample_docs)
            }
            
        except Exception as e:
            # Fallback to basic analysis if LLM fails
            themes = self._extract_themes(sample_docs)
            entities = self._extract_entities(sample_docs)
            summary = self._generate_collection_summary(sample_docs, themes)
            
            return {
                'intelligent_summary': f"Basic analysis mode:\n\n{summary}\n\nNote: Advanced AI analysis unavailable.",
                'themes': themes,
                'key_entities': entities,
                'document_count': len(documents),
                'analysis_timestamp': datetime.now().isoformat()
            }
    
    def _extract_themes_llm(self, text):
        """Extract themes using LLM for better understanding"""
        try:
            theme_prompt = f"""<|im_start|>system
Extract 5-7 main themes or topics from this text. Be specific and use full phrases, not single words.<|im_end|>
<|im_start|>user
Text: {text[:2000]}

List the main themes:<|im_end|>
<|im_start|>assistant
"""
            
            result = self.llm(
                theme_prompt,
                max_tokens=150,
                temperature=0.3,
                stop=["<|im_end|>"]
            )
            
            theme_text = result['choices'][0]['text'].strip()
            
            # Parse themes from response
            themes = []
            lines = theme_text.split('\n')
            for line in lines:
                line = line.strip()
                if line and len(line) > 5:
                    # Remove numbering and bullets
                    clean_line = re.sub(r'^[\d\.\-\*\•]+\s*', '', line)
                    if clean_line and len(clean_line.split()) >= 2:  # At least 2 words
                        themes.append({
                            'term': clean_line[:80],  # Limit length
                            'relevance': 1.0
                        })
            
            return themes[:7]
            
        except:
            return self._extract_themes([{'content': text}])
    
    def _extract_entities_enhanced(self, text):
        """Enhanced entity extraction with better patterns"""
        entities = []
        
        # Enhanced patterns for better extraction
        patterns = {
            'parties': r'\b(?:Party|Parties|Company|Corporation|LLC|Ltd|Inc|GmbH|AG)\b[^.]{0,100}',
            'dates': r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b|\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
            'monetary': r'\$\s*\d{1,3}(?:,\d{3})*(?:\.\d{2})?(?:\s*(?:million|billion|thousand|USD|EUR|GBP))?',
            'percentages': r'\b\d{1,3}(?:\.\d{1,2})?\s*%',
            'obligations': r'\b(?:shall|must|required to|obligated to|agree to)\b[^.]{0,100}',
            'deadlines': r'\b(?:within|by|before|after|no later than)\s+\d+\s+(?:days|weeks|months|years)\b'
        }
        
        for entity_type, pattern in patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                # Clean and deduplicate
                cleaned = list(set([m.strip()[:100] for m in matches]))
                if cleaned:
                    entities.append({
                        'type': entity_type.replace('_', ' ').title(),
                        'items': cleaned[:5]  # Limit to 5 examples
                    })
        
        return entities
    
    # Helper methods
    def _detect_document_type(self, text):
        """Detect document type based on content patterns"""
        text_lower = text.lower()
        
        # Technical document indicators
        if any(word in text_lower for word in ['algorithm', 'implementation', 'system', 'technical', 'specification']):
            return 'technical'
        
        # Legal document indicators
        if any(word in text_lower for word in ['shall', 'hereby', 'whereas', 'agreement', 'contract', 'legal']):
            return 'legal'
        
        # Financial document indicators
        if any(word in text_lower for word in ['financial', 'revenue', 'budget', 'cost', 'investment', 'profit']):
            return 'financial'
        
        # Research document indicators
        if any(word in text_lower for word in ['research', 'study', 'analysis', 'findings', 'methodology']):
            return 'research'
        
        return 'general'
    
    def _calculate_content_density(self, text):
        """Calculate content density score (information richness)"""
        if not text:
            return 0.0
        
        # Metrics: sentence length variety, vocabulary richness, technical terms
        sentences = re.split(r'[.!?]+', text)
        if not sentences:
            return 0.0
        
        # Sentence length variance (more varied = higher density)
        lengths = [len(s.split()) for s in sentences if s.strip()]
        if not lengths:
            return 0.0
        
        avg_length = sum(lengths) / len(lengths)
        length_variance = sum((l - avg_length) ** 2 for l in lengths) / len(lengths)
        
        # Unique word ratio
        words = text.lower().split()
        unique_ratio = len(set(words)) / len(words) if words else 0
        
        # Technical term density
        technical_terms = len(re.findall(r'\b[A-Z]{2,}\b|\b\w+[_-]\w+\b', text))
        technical_density = min(technical_terms / len(words) if words else 0, 1.0)
        
        # Combined density score
        density = (
            min(length_variance / 100, 1.0) * 0.3 +
            unique_ratio * 0.4 +
            technical_density * 0.3
        )
        
        return min(density, 1.0)
    
    def _detect_language(self, text):
        """Simple language detection based on character patterns"""
        if not text:
            return 'unknown'
        
        # Cache for performance
        text_hash = hash(text[:200])
        if text_hash in self.language_cache:
            return self.language_cache[text_hash]
        
        sample = text[:500].lower()
        
        # Character-based detection
        arabic_chars = sum(1 for c in sample if '\u0600' <= c <= '\u06FF')
        cyrillic_chars = sum(1 for c in sample if '\u0400' <= c <= '\u04FF')
        cjk_chars = sum(1 for c in sample if '\u4e00' <= c <= '\u9fff')
        german_chars = sum(1 for c in sample if c in 'äöüß')
        
        total_chars = len(sample)
        if total_chars == 0:
            return 'unknown'
        
        # Language determination
        if arabic_chars / total_chars > 0.1:
            lang = 'arabic'
        elif cyrillic_chars / total_chars > 0.1:
            lang = 'russian'
        elif cjk_chars / total_chars > 0.1:
            lang = 'chinese'
        elif german_chars > 0:
            lang = 'german'
        else:
            lang = 'english'
        
        self.language_cache[text_hash] = lang
        return lang
    
    def _semantic_search(self, query, documents):
        """Lightweight semantic search within document subset"""
        if not documents:
            return []
        
        query_embedding = self.embedder.encode(query)
        results = []
        
        for doc in documents:
            content = doc.get('content', '')
            if len(content) > 1000:
                content = content[:1000]  # Truncate for efficiency
            
            doc_embedding = self.embedder.encode(content)
            similarity = np.dot(query_embedding, doc_embedding) / (
                np.linalg.norm(query_embedding) * np.linalg.norm(doc_embedding)
            )
            
            results.append({
                'document': doc,
                'score': float(similarity),
                'content': content
            })
        
        return sorted(results, key=lambda x: x['score'], reverse=True)
    
    def _summarize_contexts(self, contexts):
        """Create brief summary of available contexts"""
        if not contexts:
            return "No context available"
        
        sources = []
        total_length = 0
        
        for ctx in contexts:
            source = ctx.get('metadata', {}).get('source', 'Unknown')
            content_preview = ctx.get('doc', '')[:100]
            sources.append(f"{source}: {content_preview}...")
            total_length += len(ctx.get('doc', ''))
        
        return f"Available sources ({len(contexts)}): {'; '.join(sources[:3])}"
    
    def _generate_fallback_questions(self, query, answer):
        """Generate fallback questions using rule-based approach"""
        questions = []
        
        # Question patterns based on query type
        if 'what' in query.lower():
            questions.append(f"How does this relate to other aspects of the topic?")
            questions.append(f"What are the practical implications of this?")
        
        if 'how' in query.lower():
            questions.append(f"What are the potential challenges with this approach?")
            questions.append(f"Are there alternative methods to consider?")
        
        if 'why' in query.lower():
            questions.append(f"What evidence supports this explanation?")
            questions.append(f"How might this impact other related areas?")
        
        # Generic follow-ups
        questions.extend([
            "Can you provide more specific examples?",
            "What additional details are available in the documents?",
            "How does this compare to industry standards or best practices?"
        ])
        
        return questions[:4]
    
    def _extract_themes(self, documents):
        """Extract key themes using lightweight text analysis"""
        if not documents:
            return []
        
        # Combine all document texts
        all_text = " ".join([doc.get('content', '')[:500] for doc in documents])
        
        # Simple keyword extraction based on frequency and uniqueness
        words = re.findall(r'\b[a-zA-Z]{4,}\b', all_text.lower())
        word_freq = {}
        
        for word in words:
            word_freq[word] = word_freq.get(word, 0) + 1
        
        # Filter out common words and get themes
        common_words = {'this', 'that', 'with', 'have', 'they', 'will', 'from', 'been', 'each', 'more', 'some'}
        
        themes = []
        for word, freq in sorted(word_freq.items(), key=lambda x: x[1], reverse=True):
            if word not in common_words and freq > 2 and len(themes) < 10:
                themes.append({
                    'term': word.title(),
                    'frequency': freq,
                    'relevance': min(freq / len(documents), 1.0)
                })
        
        return themes
    
    def _extract_entities(self, documents):
        """Extract key entities using pattern matching"""
        entities = []
        all_text = " ".join([doc.get('content', '')[:1000] for doc in documents])
        
        # Pattern-based entity extraction
        patterns = {
            'organizations': r'\b[A-Z][a-z]+ (?:Inc|Corp|LLC|Ltd|Company|Organization)\b',
            'dates': r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b(?:January|February|March|April|May|June|July|August|September|October|November|December) \d{1,2}, \d{4}\b',
            'numbers': r'\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\s*(?:%|percent|dollars?|euros?)\b',
            'codes': r'\b[A-Z]{2,5}[-_]?\d{2,6}\b'
        }
        
        for entity_type, pattern in patterns.items():
            matches = re.findall(pattern, all_text)
            if matches:
                entities.append({
                    'type': entity_type,
                    'items': list(set(matches[:10]))  # Limit and dedupe
                })
        
        return entities
    
    def _generate_collection_summary(self, documents, themes):
        """Generate a brief summary of the document collection"""
        doc_count = len(documents)
        
        if not themes:
            return f"Collection of {doc_count} documents covering various topics."
        
        top_themes = [theme['term'] for theme in themes[:3]]
        theme_text = ", ".join(top_themes)
        
        return f"Collection of {doc_count} documents primarily covering: {theme_text}. Key themes identified through content analysis."


# ============================================================================
# LICENSE MANAGEMENT SYSTEM
# ============================================================================

class LicenseManager:
    """Manages trial period and license key validation"""
    
    # Secret key for license generation (keep this secure!)
    _SECRET_KEY = "DQX2025-UNIPER-SECURE-KEY"
    _TRIAL_DAYS = 30
    
    # Valid license keys (pre-generated)
    # Format: DQRX-XXXX-XXXX-XXXX
    VALID_KEYS = [
        "DQRX-7K9M-P2XN-4HTL",
        "DQRX-3FBW-8QYC-J6VR",
        "DQRX-9LDH-5MKS-2ATP",
        "DQRX-6NWT-R4GZ-8CVE",
        "DQRX-1YPJ-X7BF-Q3UM",
        "DQRX-4KCR-2HVL-9DWN",
        "DQRX-8MSG-6TYP-Z5JB",
        "DQRX-2QXA-K9FE-7LHC",
        "DQRX-5VBN-3DWT-R8YK",
        "DQRX-7ZJL-P4MQ-1SGF",
        "DQRX-9CTH-6XVR-4KBW",
        "DQRX-3PME-Y8NJ-2FAL",
        "DQRX-6WKS-1QDC-5TXH",
        "DQRX-4LBV-7GRZ-9MYN",
        "DQRX-8FXP-2KCJ-3WQT",
        "DQRX-1DHY-5SLM-6VBG",
        "DQRX-9NRA-4ZWK-8PFC",
        "DQRX-2TCM-X6HQ-7JLS",
        "DQRX-5GVE-9BYN-1KDW",
        "DQRX-7QJZ-3FPT-4XMH",
    ]
    
    def __init__(self):
        # License data stored in user's AppData
        appdata = Path(os.environ.get('LOCALAPPDATA', os.path.expanduser('~')))
        self.license_dir = appdata / "Doqurix"
        self.license_dir.mkdir(exist_ok=True)
        self.license_file = self.license_dir / ".license"
        self.trial_file = self.license_dir / ".trial"
    
    def _get_machine_id(self):
        """Generate a unique machine identifier"""
        import platform
        machine_info = f"{platform.node()}-{platform.machine()}-{os.environ.get('USERNAME', 'user')}"
        return hashlib.sha256(machine_info.encode()).hexdigest()[:16]
    
    def _hash_data(self, data):
        """Create a hash for data integrity"""
        return hashlib.sha256(f"{data}{self._SECRET_KEY}".encode()).hexdigest()
    
    def start_trial(self):
        """Start a new trial period"""
        machine_id = self._get_machine_id()
        start_date = datetime.now().isoformat()
        
        trial_data = {
            "machine_id": machine_id,
            "start_date": start_date,
            "hash": self._hash_data(f"{machine_id}{start_date}")
        }
        
        with open(self.trial_file, 'w') as f:
            json.dump(trial_data, f)
        
        return self._TRIAL_DAYS
    
    def get_trial_status(self):
        """
        Check trial status
        Returns: (is_valid, days_remaining, message)
        """
        if not self.trial_file.exists():
            # First run - start trial
            days = self.start_trial()
            return True, days, f"Trial started! {days} days remaining."
        
        try:
            with open(self.trial_file, 'r') as f:
                trial_data = json.load(f)
            
            # Verify data integrity
            machine_id = self._get_machine_id()
            expected_hash = self._hash_data(f"{trial_data['machine_id']}{trial_data['start_date']}")
            
            if trial_data['hash'] != expected_hash or trial_data['machine_id'] != machine_id:
                return False, 0, "Trial data corrupted or transferred. Please enter a license key."
            
            # Calculate remaining days
            start_date = datetime.fromisoformat(trial_data['start_date'])
            elapsed = datetime.now() - start_date
            remaining = self._TRIAL_DAYS - elapsed.days
            
            if remaining > 0:
                return True, remaining, f"Trial: {remaining} days remaining"
            else:
                return False, 0, "Trial period expired. Please enter a license key to continue."
            
        except (json.JSONDecodeError, KeyError, ValueError):
            return False, 0, "Trial data corrupted. Please enter a license key."
    
    def validate_license_key(self, key):
        """
        Validate a license key
        Returns: (is_valid, message)
        """
        # Normalize key format
        key = key.strip().upper().replace(" ", "-")
        
        # Check format
        if not re.match(r'^DQRX-[A-Z0-9]{4}-[A-Z0-9]{4}-[A-Z0-9]{4}$', key):
            return False, "Invalid key format. Expected: DQRX-XXXX-XXXX-XXXX"
        
        # Check against valid keys
        if key in self.VALID_KEYS:
            self._save_license(key)
            return True, "License activated successfully! Thank you for your purchase."
        
        return False, "Invalid license key. Please check and try again."
    
    def _save_license(self, key):
        """Save activated license"""
        machine_id = self._get_machine_id()
        activation_date = datetime.now().isoformat()
        
        license_data = {
            "key": key,
            "machine_id": machine_id,
            "activation_date": activation_date,
            "hash": self._hash_data(f"{key}{machine_id}{activation_date}")
        }
        
        with open(self.license_file, 'w') as f:
            json.dump(license_data, f)
    
    def check_license(self):
        """
        Check if software is licensed
        Returns: (is_licensed, message, days_remaining)
        - days_remaining is -1 if not licensed, otherwise days until expiration
        """
        if not self.license_file.exists():
            return False, "Not licensed", -1
        
        try:
            with open(self.license_file, 'r') as f:
                license_data = json.load(f)
            
            machine_id = self._get_machine_id()
            expected_hash = self._hash_data(
                f"{license_data['key']}{license_data['machine_id']}{license_data['activation_date']}"
            )
            
            if license_data['hash'] != expected_hash:
                return False, "License data corrupted", -1
            
            if license_data['machine_id'] != machine_id:
                return False, "License not valid for this machine", -1
            
            if license_data['key'] not in self.VALID_KEYS:
                return False, "Invalid license key", -1
            
            # Check if license has expired (1 year = 365 days)
            activation_date = datetime.fromisoformat(license_data['activation_date'])
            days_since_activation = (datetime.now() - activation_date).days
            license_days_remaining = 365 - days_since_activation
            
            if license_days_remaining <= 0:
                return False, "License expired (1 year limit reached)", 0
            
            return True, "Licensed", license_days_remaining
            
        except (json.JSONDecodeError, KeyError):
            return False, "License data corrupted", -1
    
    def get_status(self):
        """
        Get overall license status
        Returns: (can_run, is_trial, days_remaining, message)
        - For licensed users: days_remaining = days until license expires
        - For trial users: days_remaining = days left in trial
        """
        # First check if licensed
        is_licensed, license_msg, license_days = self.check_license()
        if is_licensed:
            if license_days <= 30:
                return True, False, license_days, f"✓ Licensed ({license_days} days left)"
            else:
                return True, False, license_days, "✓ Licensed"
        
        # If license expired, show that message
        if license_msg == "License expired (1 year limit reached)":
            return False, False, 0, license_msg
        
        # Check trial
        trial_valid, days_remaining, trial_msg = self.get_trial_status()
        if trial_valid:
            return True, True, days_remaining, trial_msg
        
        # Neither licensed nor valid trial
        return False, False, 0, trial_msg


class LicenseDialog:
    """Dialog for entering license key"""
    
    def __init__(self, parent, license_manager, days_remaining=0, is_expired=False):
        self.result = False
        self.license_manager = license_manager
        
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Doqurix - License Activation")
        self.dialog.geometry("500x400")
        self.dialog.resizable(False, False)
        self.dialog.transient(parent)
        self.dialog.grab_set()
        
        # Center the dialog
        self.dialog.update_idletasks()
        x = (self.dialog.winfo_screenwidth() // 2) - 250
        y = (self.dialog.winfo_screenheight() // 2) - 200
        self.dialog.geometry(f"+{x}+{y}")
        
        self.dialog.configure(bg='#f5f5f5')
        
        # Main frame
        main_frame = tk.Frame(self.dialog, bg='#f5f5f5', padx=30, pady=20)
        main_frame.pack(fill='both', expand=True)
        
        # Logo/Title
        title_label = tk.Label(main_frame, text="📚 Doqurix",
                               font=('Segoe UI', 24, 'bold'),
                               bg='#f5f5f5', fg='#2c3e50')
        title_label.pack(pady=(0, 5))
        
        subtitle = tk.Label(main_frame, text="License Activation",
                           font=('Segoe UI', 12),
                           bg='#f5f5f5', fg='#7f8c8d')
        subtitle.pack(pady=(0, 20))
        
        # Status message
        if is_expired:
            status_text = "⚠️ Your trial period has expired."
            status_color = '#e74c3c'
        else:
            status_text = f"📅 Trial: {days_remaining} days remaining"
            status_color = '#27ae60'
        
        status_label = tk.Label(main_frame, text=status_text,
                               font=('Segoe UI', 11),
                               bg='#f5f5f5', fg=status_color)
        status_label.pack(pady=(0, 20))
        
        # License key entry
        key_frame = tk.Frame(main_frame, bg='#f5f5f5')
        key_frame.pack(fill='x', pady=10)
        
        tk.Label(key_frame, text="Enter License Key:",
                font=('Segoe UI', 10, 'bold'),
                bg='#f5f5f5').pack(anchor='w')
        
        self.key_entry = tk.Entry(key_frame, font=('Consolas', 14),
                                  width=30, justify='center')
        self.key_entry.pack(fill='x', pady=(5, 0), ipady=8)
        self.key_entry.insert(0, "DQRX-XXXX-XXXX-XXXX")
        self.key_entry.bind('<FocusIn>', self._clear_placeholder)
        self.key_entry.bind('<Return>', lambda e: self._activate())
        
        # Message label
        self.message_label = tk.Label(main_frame, text="",
                                      font=('Segoe UI', 9),
                                      bg='#f5f5f5', fg='#e74c3c')
        self.message_label.pack(pady=10)
        
        # Buttons
        button_frame = tk.Frame(main_frame, bg='#f5f5f5')
        button_frame.pack(pady=20)
        
        activate_btn = tk.Button(button_frame, text="🔑 Activate License",
                                 command=self._activate,
                                 font=('Segoe UI', 10, 'bold'),
                                 bg='#27ae60', fg='white',
                                 activebackground='#229954',
                                 relief='flat', padx=20, pady=10,
                                 cursor='hand2')
        activate_btn.pack(side='left', padx=5)
        
        if not is_expired:
            continue_btn = tk.Button(button_frame, text="Continue Trial",
                                     command=self._continue_trial,
                                     font=('Segoe UI', 10),
                                     bg='#3498db', fg='white',
                                     activebackground='#2980b9',
                                     relief='flat', padx=20, pady=10,
                                     cursor='hand2')
            continue_btn.pack(side='left', padx=5)
        
        exit_btn = tk.Button(button_frame, text="Exit",
                            command=self._exit,
                            font=('Segoe UI', 10),
                            bg='#95a5a6', fg='white',
                            activebackground='#7f8c8d',
                            relief='flat', padx=20, pady=10,
                            cursor='hand2')
        exit_btn.pack(side='left', padx=5)
        
        # Purchase info
        info_label = tk.Label(main_frame, 
                             text="Need a license? Contact: sales@doqurix.com",
                             font=('Segoe UI', 9, 'italic'),
                             bg='#f5f5f5', fg='#7f8c8d')
        info_label.pack(side='bottom', pady=10)
        
        # Handle window close
        self.dialog.protocol("WM_DELETE_WINDOW", self._exit)
    
    def _clear_placeholder(self, event):
        if self.key_entry.get() == "DQRX-XXXX-XXXX-XXXX":
            self.key_entry.delete(0, 'end')
    
    def _activate(self):
        key = self.key_entry.get()
        is_valid, message = self.license_manager.validate_license_key(key)
        
        if is_valid:
            self.message_label.config(text=message, fg='#27ae60')
            self.result = True
            self.dialog.after(1500, self.dialog.destroy)
        else:
            self.message_label.config(text=message, fg='#e74c3c')
    
    def _continue_trial(self):
        self.result = True
        self.dialog.destroy()
    
    def _exit(self):
        self.result = False
        self.dialog.destroy()
    
    def show(self):
        self.dialog.wait_window()
        return self.result


class DownloadProgressWindow:
    """A progress window that shows download status to users"""
    
    def __init__(self, parent, title="Downloading..."):
        self.window = tk.Toplevel(parent)
        self.window.title(title)
        self.window.geometry("500x200")
        self.window.resizable(False, False)
        self.window.transient(parent)
        self.window.grab_set()
        
        # Center the window
        self.window.update_idletasks()
        x = (self.window.winfo_screenwidth() // 2) - (250)
        y = (self.window.winfo_screenheight() // 2) - (100)
        self.window.geometry(f"+{x}+{y}")
        
        # Configure style
        self.window.configure(bg='#f5f5f5')
        
        # Main frame
        main_frame = tk.Frame(self.window, bg='#f5f5f5', padx=30, pady=20)
        main_frame.pack(fill='both', expand=True)
        
        # Title label
        self.title_label = tk.Label(main_frame, text="⏳ First Time Setup", 
                                    font=('Segoe UI', 14, 'bold'),
                                    bg='#f5f5f5', fg='#2c3e50')
        self.title_label.pack(pady=(0, 10))
        
        # Status label
        self.status_label = tk.Label(main_frame, text="Preparing download...", 
                                     font=('Segoe UI', 10),
                                     bg='#f5f5f5', fg='#555')
        self.status_label.pack(pady=(0, 15))
        
        # Progress bar
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(main_frame, variable=self.progress_var,
                                            maximum=100, length=400, mode='determinate')
        self.progress_bar.pack(pady=(0, 10))
        
        # Percentage and size label
        self.detail_label = tk.Label(main_frame, text="0% - Calculating...", 
                                     font=('Segoe UI', 9),
                                     bg='#f5f5f5', fg='#777')
        self.detail_label.pack()
        
        # ETA label
        self.eta_label = tk.Label(main_frame, text="", 
                                  font=('Segoe UI', 9),
                                  bg='#f5f5f5', fg='#777')
        self.eta_label.pack()
        
        self.start_time = None
        self.window.protocol("WM_DELETE_WINDOW", lambda: None)  # Prevent closing
        
    def update_progress(self, downloaded, total, speed=None):
        """Update the progress bar and labels"""
        if total > 0:
            percent = (downloaded / total) * 100
            self.progress_var.set(percent)
            
            # Format sizes
            downloaded_mb = downloaded / (1024 * 1024)
            total_mb = total / (1024 * 1024)
            
            self.detail_label.config(text=f"{percent:.1f}% - {downloaded_mb:.1f} MB / {total_mb:.1f} MB")
            
            # Calculate ETA
            if self.start_time and downloaded > 0:
                elapsed = time.time() - self.start_time
                if elapsed > 0:
                    speed_calc = downloaded / elapsed
                    remaining = total - downloaded
                    eta_seconds = remaining / speed_calc if speed_calc > 0 else 0
                    
                    if eta_seconds < 60:
                        eta_str = f"{int(eta_seconds)} seconds"
                    elif eta_seconds < 3600:
                        eta_str = f"{int(eta_seconds // 60)} min {int(eta_seconds % 60)} sec"
                    else:
                        eta_str = f"{int(eta_seconds // 3600)} hr {int((eta_seconds % 3600) // 60)} min"
                    
                    speed_mbps = (speed_calc * 8) / (1024 * 1024)
                    self.eta_label.config(text=f"Speed: {speed_mbps:.1f} Mbps • ETA: {eta_str}")
        else:
            self.progress_bar.config(mode='indeterminate')
            self.progress_bar.start(10)
            
        self.window.update()
    
    def set_status(self, status):
        """Update status text"""
        self.status_label.config(text=status)
        self.window.update()
    
    def set_title(self, title):
        """Update title text"""
        self.title_label.config(text=title)
        self.window.update()
    
    def start_timer(self):
        """Start the download timer for ETA calculation"""
        self.start_time = time.time()
    
    def set_indeterminate(self):
        """Set progress bar to indeterminate mode"""
        self.progress_bar.config(mode='indeterminate')
        self.progress_bar.start(10)
        self.detail_label.config(text="Please wait...")
        self.eta_label.config(text="")
        self.window.update()
    
    def close(self):
        """Close the progress window"""
        self.window.grab_release()
        self.window.destroy()

class DocumentQAApp:
    def __init__(self, trial_info=None):
        self.app_dir = Path(__file__).parent
        self.trial_info = trial_info  # (is_trial, days_remaining)
        
        # Use AppData for user-writable directories (works in Program Files)
        appdata = Path(os.environ.get('LOCALAPPDATA', os.path.expanduser('~')))
        self.user_data_dir = appdata / "Doqurix"
        self.user_data_dir.mkdir(exist_ok=True)
        
        # All user data stored in AppData (Program Files is read-only)
        self.models_dir = self.user_data_dir / "models"
        self.data_dir = self.user_data_dir / "data"
        self.vector_store_dir = self.data_dir / "vector_store"
        
        # Create directories
        self.models_dir.mkdir(exist_ok=True)
        self.data_dir.mkdir(exist_ok=True)
        self.vector_store_dir.mkdir(exist_ok=True)
        
        # BM25 index
        self.bm25 = None
        self.bm25_corpus = []
        self.current_contexts = []
        
        # Progress window reference
        self.progress_window = None
        
        # Setup GUI first
        self.setup_gui()
        
        # Initialize models in background
        self.models_loaded = False
        threading.Thread(target=self.load_models, daemon=True).start()
    
    def show_progress_window(self, title="Setting Up..."):
        """Show progress window on main thread"""
        if self.progress_window is None:
            self.progress_window = DownloadProgressWindow(self.root, title)
        return self.progress_window
    
    def close_progress_window(self):
        """Close progress window safely"""
        if self.progress_window:
            self.progress_window.close()
            self.progress_window = None
    
    def load_models(self):
        """Load all models in background"""
        try:
            self.update_status("🔄 Loading AI engine...")
            
            # Check if any models need downloading
            needs_download = self.check_models_need_download()
            
            if needs_download:
                # Show progress window on main thread
                self.root.after(0, lambda: self.show_progress_window("First Time Setup"))
                self.root.after(100, self._continue_loading)
            else:
                self._load_all_models()
                
        except Exception as e:
            self.root.after(0, self.close_progress_window)
            self.update_status(f"❌ Error: {str(e)}")
            self.root.after(0, lambda: messagebox.showerror("Error", 
                f"Failed to load models:\n{str(e)}\n\nMake sure you have a good internet connection"))
    
    def check_models_need_download(self):
        """Check if any models need to be downloaded"""
        model_path = self.models_dir / "qwen2.5-1.5b-instruct-q4_k_m.gguf"
        
        # Check for the main GGUF model
        if not model_path.exists():
            return True
        
        # Check for sentence transformers cache
        cache_dir = Path(os.environ.get('HF_HOME', Path.home() / '.cache' / 'huggingface'))
        st_cache = cache_dir / 'hub'
        
        # If cache doesn't exist or is empty, models will need to download
        if not st_cache.exists():
            return True
            
        return False
    
    def _continue_loading(self):
        """Continue loading in background after progress window is shown"""
        threading.Thread(target=self._load_all_models, daemon=True).start()
    
    def _load_all_models(self):
        """Actually load all the models"""
        try:
            self.init_llm()
            self.init_embeddings()
            self.init_reranker()
            self.init_vector_db()
            self.init_agent_workflows()
            
            self.root.after(0, self.close_progress_window)
            self.models_loaded = True
            self.update_status("✓ Ready! Upload documents and start asking questions.")
            self.root.after(0, lambda: self.ask_btn.config(state='normal'))
            self.root.after(0, lambda: self.upload_btn.config(state='normal'))
            self.root.after(0, lambda: self.summary_btn.config(state='normal'))
            self.root.after(0, lambda: self.insights_btn.config(state='normal'))
            # Enable delete button only if there are documents
            if self.doc_listbox.size() > 0:
                self.root.after(0, lambda: self.delete_btn.config(state='normal'))
        except Exception as e:
            self.root.after(0, self.close_progress_window)
            self.update_status(f"❌ Error: {str(e)}")
            self.root.after(0, lambda: messagebox.showerror("Error", 
                f"Failed to load models:\n{str(e)}\n\nMake sure you have a good internet connection"))
    
    def download_with_progress(self, url, dest_path, description="Downloading..."):
        """Download a file with progress bar updates"""
        import urllib.request
        
        # Update progress window
        if self.progress_window:
            self.root.after(0, lambda: self.progress_window.set_title("⬇️ Downloading AI Model"))
            self.root.after(0, lambda: self.progress_window.set_status(description))
            self.root.after(0, lambda: self.progress_window.start_timer())
        
        # Get file size first
        try:
            with urllib.request.urlopen(url, timeout=30) as response:
                total_size = int(response.headers.get('Content-Length', 0))
                
                # Download in chunks with progress
                downloaded = 0
                chunk_size = 1024 * 1024  # 1MB chunks
                
                with open(dest_path, 'wb') as f:
                    while True:
                        chunk = response.read(chunk_size)
                        if not chunk:
                            break
                        f.write(chunk)
                        downloaded += len(chunk)
                        
                        # Update progress bar
                        if self.progress_window and total_size > 0:
                            self.root.after(0, lambda d=downloaded, t=total_size: 
                                           self.progress_window.update_progress(d, t))
        except Exception as e:
            # Clean up partial download
            if dest_path.exists():
                dest_path.unlink()
            raise e
        
        return dest_path
    
    def init_llm(self):
        """Initialize LLM with quantization via llama.cpp"""
        self.update_status("📥 Checking AI model...")
        
        # Download quantized GGUF model from HuggingFace
        model_path = self.models_dir / "qwen2.5-1.5b-instruct-q4_k_m.gguf"
        
        if not model_path.exists():
            self.update_status("⬇️ Downloading AI model...")
            
            # Direct download URL from HuggingFace
            model_url = "https://huggingface.co/Qwen/Qwen2.5-1.5B-Instruct-GGUF/resolve/main/qwen2.5-1.5b-instruct-q4_k_m.gguf"
            
            try:
                self.download_with_progress(
                    url=model_url,
                    dest_path=model_path,
                    description="Downloading language model (~1 GB)..."
                )
            except Exception as e:
                self.update_status(f"Error downloading: {str(e)}")
                raise
        
        self.update_status("🧠 Loading AI engine into memory...")
        if self.progress_window:
            self.root.after(0, lambda: self.progress_window.set_title("🧠 Loading AI Engine"))
            self.root.after(0, lambda: self.progress_window.set_status("Loading model into memory..."))
            self.root.after(0, lambda: self.progress_window.set_indeterminate())
        
        # Load with llama.cpp (optimized for CPU with max performance)
        self.llm = Llama(
            model_path=str(model_path),
            n_ctx=3072,  # Increased context window
            n_threads=os.cpu_count() or 4,  # Use all CPU cores
            n_batch=512,  # Larger batch for faster processing
            n_gpu_layers=0,  # CPU only
            verbose=False,
            use_mlock=True  # Lock model in RAM for faster access
        )
        
        self.update_status("✓ AI engine loaded successfully")
    
    def init_embeddings(self):
        """Initialize faster embedding model"""
        self.update_status("📥 Loading search engine...")
        
        if self.progress_window:
            self.root.after(0, lambda: self.progress_window.set_title("📥 Loading Search Engine"))
            self.root.after(0, lambda: self.progress_window.set_status("Downloading embedding model..."))
            self.root.after(0, lambda: self.progress_window.set_indeterminate())
        
        # Using faster, smaller embedding model
        self.embedder = SentenceTransformer('all-MiniLM-L6-v2')
        self.update_status("✓ Embeddings loaded")
    
    def init_reranker(self):
        """Initialize faster reranker"""
        self.update_status("📥 Loading ranking engine...")
        
        if self.progress_window:
            self.root.after(0, lambda: self.progress_window.set_title("📥 Loading Ranking Engine"))
            self.root.after(0, lambda: self.progress_window.set_status("Downloading reranker model..."))
            self.root.after(0, lambda: self.progress_window.set_indeterminate())
        
        # Using smaller, faster reranker
        self.reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
        self.update_status("✓ Reranker loaded")
    
    def init_vector_db(self):
        """Initialize ChromaDB"""
        self.update_status("📥 Initializing vector database...")
        
        if self.progress_window:
            self.root.after(0, lambda: self.progress_window.set_title("📥 Initializing Database"))
            self.root.after(0, lambda: self.progress_window.set_status("Setting up vector database..."))
            self.root.after(0, lambda: self.progress_window.set_indeterminate())
        
        self.chroma_client = chromadb.Client(Settings(
            persist_directory=str(self.vector_store_dir),
            anonymized_telemetry=False
        ))
        
        try:
            self.collection = self.chroma_client.get_collection("documents")
            self.rebuild_bm25_index()
            doc_count = len(self.collection.get()['documents'])
            if doc_count > 0:
                self.update_status(f"✓ Database ready ({doc_count} chunks)")
        except:
            self.collection = self.chroma_client.create_collection("documents")
            self.update_status("✓ Database ready")
    
    def rebuild_bm25_index(self):
        """Rebuild BM25 index"""
        all_docs = self.collection.get()
        if all_docs['documents']:
            self.bm25_corpus = all_docs['documents']
            tokenized_corpus = [doc.lower().split() for doc in self.bm25_corpus]
            self.bm25 = BM25Okapi(tokenized_corpus)
    
    def init_agent_workflows(self):
        """Initialize intelligent agent workflows"""
        self.update_status("🤖 Initializing AI agents...")
        
        if self.progress_window:
            self.root.after(0, lambda: self.progress_window.set_title("🤖 Setting up AI Agents"))
            self.root.after(0, lambda: self.progress_window.set_status("Configuring intelligent workflows..."))
        
        # Initialize agent system with loaded models
        self.agents = AgentWorkflows(
            llm=self.llm,
            embedder=self.embedder,
            reranker=self.reranker
        )
        
        # Initialize specialized agent collections
        self.init_tax_agent_collection()
        self.init_buerokratai_collection()
        
        self.update_status("✓ AI agents ready for automation")
    
    def init_tax_agent_collection(self):
        """Initialize German Tax Agent collection with pre-loaded knowledge"""
        try:
            # Check if tax collection already exists
            try:
                self.tax_collection = self.chroma_client.get_collection("tax_agent_germany")
                doc_count = len(self.tax_collection.get()['documents'])
                if doc_count > 0:
                    self.update_status(f"✓ Tax Agent ready ({doc_count} knowledge chunks)")
                    return
            except:
                # Create new collection
                self.tax_collection = self.chroma_client.create_collection("tax_agent_germany")
            
            # Load tax knowledge documents
            tax_knowledge_dir = self.app_dir / "tax_knowledge"
            if not tax_knowledge_dir.exists():
                self.update_status("⚠ Tax knowledge not found, agent will have limited capability")
                return
            
            self.update_status("📚 Loading German Tax Agent knowledge base...")
            
            # Process each tax knowledge file
            all_chunks = []
            all_ids = []
            all_metadatas = []
            chunk_id = 0
            
            for tax_file in tax_knowledge_dir.glob("*.txt"):
                with open(tax_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Split into chunks (approximately 500 words each for better context)
                words = content.split()
                chunk_size = 500
                overlap = 50
                
                for i in range(0, len(words), chunk_size - overlap):
                    chunk_words = words[i:i + chunk_size]
                    chunk_text = ' '.join(chunk_words)
                    
                    if len(chunk_text.strip()) > 100:  # Skip very small chunks
                        all_chunks.append(chunk_text)
                        all_ids.append(f"tax_{chunk_id}")
                        all_metadatas.append({
                            'source': tax_file.name,
                            'type': 'tax_knowledge',
                            'agent': 'tax_germany'
                        })
                        chunk_id += 1
            
            if all_chunks:
                # Embed and store in batches
                batch_size = 100
                for i in range(0, len(all_chunks), batch_size):
                    batch_chunks = all_chunks[i:i+batch_size]
                    batch_ids = all_ids[i:i+batch_size]
                    batch_metadatas = all_metadatas[i:i+batch_size]
                    
                    # Generate embeddings
                    embeddings = self.embedder.encode(batch_chunks).tolist()
                    
                    # Add to collection
                    self.tax_collection.add(
                        documents=batch_chunks,
                        embeddings=embeddings,
                        metadatas=batch_metadatas,
                        ids=batch_ids
                    )
                
                self.update_status(f"✓ Tax Agent loaded with {len(all_chunks)} knowledge chunks")
            else:
                self.update_status("⚠ No tax knowledge loaded")
                
        except Exception as e:
            self.update_status(f"⚠ Tax Agent initialization failed: {str(e)}")
            print(f"Tax agent error: {e}")
    
    def init_buerokratai_collection(self):
        """Initialize BürokratAI Agent collection with pre-loaded immigration knowledge"""
        try:
            # Check if BürokratAI collection already exists
            try:
                self.buerokratai_collection = self.chroma_client.get_collection("buerokratai_agent")
                doc_count = len(self.buerokratai_collection.get()['documents'])
                if doc_count > 0:
                    self.update_status(f"✓ BürokratAI Agent ready ({doc_count} knowledge chunks)")
                    return
            except:
                # Create new collection
                self.buerokratai_collection = self.chroma_client.create_collection("buerokratai_agent")
            
            # Load BürokratAI knowledge documents
            buerokratai_knowledge_dir = self.app_dir / "buerokratai_knowledge"
            if not buerokratai_knowledge_dir.exists():
                self.update_status("⚠ BürokratAI knowledge not found, agent will have limited capability")
                return
            
            self.update_status("📚 Loading BürokratAI immigration knowledge base...")
            
            # Process each knowledge file
            all_chunks = []
            all_ids = []
            all_metadatas = []
            chunk_id = 0
            
            for knowledge_file in buerokratai_knowledge_dir.glob("*.txt"):
                with open(knowledge_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Split into chunks (approximately 500 words each for better context)
                words = content.split()
                chunk_size = 500
                overlap = 50
                
                for i in range(0, len(words), chunk_size - overlap):
                    chunk_words = words[i:i + chunk_size]
                    chunk_text = ' '.join(chunk_words)
                    
                    if len(chunk_text.strip()) > 100:  # Skip very small chunks
                        all_chunks.append(chunk_text)
                        all_ids.append(f"buerokratai_{chunk_id}")
                        all_metadatas.append({
                            'source': knowledge_file.name,
                            'type': 'immigration_knowledge',
                            'agent': 'buerokratai_germany'
                        })
                        chunk_id += 1
            
            if all_chunks:
                # Embed and store in batches
                batch_size = 100
                for i in range(0, len(all_chunks), batch_size):
                    batch_chunks = all_chunks[i:i+batch_size]
                    batch_ids = all_ids[i:i+batch_size]
                    batch_metadatas = all_metadatas[i:i+batch_size]
                    
                    # Generate embeddings
                    embeddings = self.embedder.encode(batch_chunks).tolist()
                    
                    # Add to collection
                    self.buerokratai_collection.add(
                        documents=batch_chunks,
                        embeddings=embeddings,
                        metadatas=batch_metadatas,
                        ids=batch_ids
                    )
                
                self.update_status(f"✓ BürokratAI Agent loaded with {len(all_chunks)} knowledge chunks")
            else:
                self.update_status("⚠ No BürokratAI knowledge loaded")
                
        except Exception as e:
            self.update_status(f"⚠ BürokratAI Agent initialization failed: {str(e)}")
            print(f"BürokratAI agent error: {e}")
    
    def on_agent_change(self, event=None):
        """Handle agent selection change"""
        # Map display names to internal values
        display_to_internal = {
            "None": "none",
            "Insights": "insights",
            "Tax Germany": "tax_germany",
            "E-Commerce Germany": "ecommerce_germany",
            "BürokratAI": "buerokratai_germany"
        }
        
        agent_display = self.agent_mode.get()
        agent = display_to_internal.get(agent_display, "none")
        
        # Update info label based on selection
        agent_descriptions = {
            "none": "None - Use your uploaded documents",
            "insights": "Insights Agent - Automatic document analysis and insights",
            "tax_germany": "German Tax Agent - Expert knowledge on German taxation (No PDF upload needed)",
            "ecommerce_germany": "E-Commerce Agent - Product search & price comparison for German market (No PDF upload needed)",
            "buerokratai_germany": "BürokratAI - Immigration & bureaucracy assistant for Germany (No PDF upload needed)"
        }
        
        description = agent_descriptions.get(agent, "Unknown agent")
        self.agent_info.config(text=description)
        
        # Update UI based on agent
        if agent == "buerokratai_germany":
            # BürokratAI doesn't need document upload
            self.update_status("✓ BürokratAI Agent selected - Ask about German bureaucracy!")
            self.answer_text.delete('1.0', 'end')
            self.answer_text.insert('1.0', 
                "🏛️ BürokratAI - Immigration Assistant Activated\n\n"
                "I help immigrants navigate German bureaucracy. I have knowledge about:\n\n"
                "📋 Registration & Documents:\n"
                "• Anmeldung (Address Registration)\n"
                "• Tax ID (Steuer-ID) Application\n"
                "• Health Insurance Requirements\n\n"
                "🛂 Visas & Residence Permits:\n"
                "• EU Blue Card\n"
                "• Student Visa\n"
                "• Job Seeker Visa\n"
                "• Family Reunification\n"
                "• Opportunity Card (Chancenkarte)\n\n"
                "🏠 Living in Germany:\n"
                "• Renting an Apartment\n"
                "• Opening a Bank Account\n"
                "• Driver's License Conversion\n"
                "• Integration Course\n\n"
                "💼 Work & Career:\n"
                "• Work Permits\n"
                "• Recognition of Qualifications\n"
                "• Permanent Residence Path\n\n"
                "Ask me anything! For example:\n"
                "- 'What documents do I need for Anmeldung?'\n"
                "- 'How do I get a Tax ID?'\n"
                "- 'What are the requirements for EU Blue Card?'\n"
                "- 'How can I convert my driver's license?'\n\n"
                "⚠️ Note: Information is based on official German sources.\n"
                "Always verify with authorities for current regulations."
            )
        elif agent == "tax_germany":
            # Tax agent doesn't need document upload
            self.update_status("✓ German Tax Agent selected - Ask tax questions directly!")
            self.answer_text.delete('1.0', 'end')
            self.answer_text.insert('1.0', 
                "🇩🇪 German Tax Agent Activated\n\n"
                "I'm your expert on German taxation. I have comprehensive knowledge about:\n\n"
                "• Income Tax (Einkommensteuer)\n"
                "• Value Added Tax (Umsatzsteuer/VAT)\n"
                "• Corporate Tax (Körperschaftsteuer)\n"
                "• Trade Tax (Gewerbesteuer)\n"
                "• Church Tax (Kirchensteuer)\n"
                "• Capital Gains Tax (Abgeltungsteuer)\n"
                "• Real Estate Transfer Tax (Grunderwerbsteuer)\n"
                "• Inheritance & Gift Tax\n"
                "• Social Security Contributions\n"
                "• Tax Classes, Deductions, Filing Requirements\n\n"
                "Ask me anything about German taxes! For example:\n"
                "- 'What are the income tax brackets for 2025?'\n"
                "- 'How does the Kleinunternehmerregelung work?'\n"
                "- 'What can I deduct as a home office?'\n"
                "- 'Explain trade tax calculation'\n"
                "- 'What is the solidarity surcharge?'\n\n"
                "No document upload needed - I have pre-loaded expert knowledge!"
            )
        elif agent == "ecommerce_germany":
            # E-commerce agent for product search
            self.update_status("✓ E-Commerce Agent selected - Search for products!")
            self.answer_text.delete('1.0', 'end')
            self.answer_text.insert('1.0',
                "🛒 E-Commerce Agent Activated (German Market)\n\n"
                "I can help you find and compare products! Just tell me what you're looking for.\n\n"
                "What I can do:\n"
                "• Search for products across German online shops\n"
                "• Compare prices from multiple retailers\n"
                "• Provide direct links to purchase\n"
                "• Summarize product features and specifications\n"
                "• Help you make informed buying decisions\n\n"
                "Example searches:\n"
                "- 'Find me a gaming laptop under 1500 euros'\n"
                "- 'Compare prices for iPhone 15 Pro'\n"
                "- 'Best wireless headphones for running'\n"
                "- 'Affordable coffee machines with milk frother'\n"
                "- 'Top rated 4K monitors for home office'\n\n"
                "No document upload needed - Just ask what you want to buy!"
            )
        elif agent == "insights":
            self.update_status("✓ Insights Agent selected - Upload documents for automatic analysis")
        else:
            self.update_status("✓ Standard mode - Upload your documents to begin")
            if self.answer_text.get('1.0', 'end').strip() == "":
                self.answer_text.delete('1.0', 'end')
                self.answer_text.insert('1.0', "Upload documents and ask questions to get started.")
    
    def get_document_collection_insights(self):
        """Get automated insights about the document collection"""
        if not self.models_loaded:
            return "Models not loaded yet"
        
        try:
            # Get all documents from collection
            all_docs = self.collection.get()
            if not all_docs['documents']:
                return "No documents available for analysis"
            
            # Prepare documents for analysis
            documents = []
            for i, (doc, metadata) in enumerate(zip(all_docs['documents'], all_docs['metadatas'])):
                documents.append({
                    'content': doc,
                    'metadata': metadata,
                    'id': all_docs['ids'][i]
                })
            
            # Get automated insights using LLM
            insights = self.agents.auto_document_insights(documents)
            
            # Format insights in a clean, professional layout
            insight_text = f"INTELLIGENT DOCUMENT ANALYSIS\n\n"
            insight_text += f"Document Collection: {insights['document_count']} chunks analyzed\n"
            insight_text += f"Analysis Depth: {insights.get('analyzed_samples', 'N/A')} representative samples\n"
            insight_text += f"Generated: {insights['analysis_timestamp'][:19].replace('T', ' ')}\n\n\n"
            
            # Clean the LLM output of markdown formatting
            if insights.get('intelligent_summary'):
                clean_summary = insights['intelligent_summary']
                
                # Remove markdown bold/italic
                clean_summary = re.sub(r'\*\*(.+?)\*\*', r'\1', clean_summary)  # **bold**
                clean_summary = re.sub(r'\*(.+?)\*', r'\1', clean_summary)      # *italic*
                clean_summary = re.sub(r'__(.+?)__', r'\1', clean_summary)      # __bold__
                clean_summary = re.sub(r'_(.+?)_', r'\1', clean_summary)        # _italic_
                
                # Remove markdown headers
                clean_summary = re.sub(r'^#{1,6}\s+', '', clean_summary, flags=re.MULTILINE)
                
                # Clean bullet points and numbering
                clean_summary = re.sub(r'^\s*[-\*\•]\s+', '  • ', clean_summary, flags=re.MULTILINE)
                clean_summary = re.sub(r'^\s*\d+[\.\)]\s+', '  ', clean_summary, flags=re.MULTILINE)
                
                # Remove extra blank lines
                clean_summary = re.sub(r'\n{3,}', '\n\n', clean_summary)
                
                insight_text += clean_summary.strip()
                insight_text += "\n\n\n"
            
            # Add structured insights if available
            if insights.get('themes'):
                insight_text += "IDENTIFIED THEMES\n\n"
                for i, theme in enumerate(insights['themes'][:5], 1):
                    insight_text += f"{i}. {theme['term']}\n"
                insight_text += "\n\n"
            
            if insights.get('key_entities'):
                insight_text += "EXTRACTED ENTITIES\n\n"
                for entity_group in insights['key_entities']:
                    insight_text += f"{entity_group['type']}:\n"
                    for item in entity_group['items'][:3]:
                        insight_text += f"  • {item}\n"
                    insight_text += "\n"
                insight_text += "\n"
            
            insight_text += "TIP: Use these insights to formulate targeted questions or explore specific topics in greater depth."
            
            return insight_text
            
        except Exception as e:
            return f"Error generating insights:\n\n{str(e)}\n\nPlease try again or check your documents."
    
    def show_document_insights(self):
        """Show automated document collection insights"""
        if not self.models_loaded:
            messagebox.showwarning("Please Wait", "Models still loading.")
            return
        
        if self.doc_listbox.size() == 0:
            messagebox.showwarning("No Documents", "Upload documents first to generate insights.")
            return
        
        # Disable buttons during processing
        self.insights_btn.config(state='disabled')
        self.ask_btn.config(state='disabled')
        self.summary_btn.config(state='disabled')
        
        # Show insights in the answer area with animation
        self.answer_text.delete('1.0', 'end')
        self.answer_text.insert('1.0', "🤖 Generating document collection insights...\n\n⏳ Processing")
        self.update_status("🤖 AI agents analyzing document collection...")
        
        # Start processing animation
        self.processing = True
        self.insights_start_time = time.time()
        self._animate_insights_processing()
        
        def generate_insights():
            try:
                insights = self.get_document_collection_insights()
                self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
                self.root.after(0, lambda: self.answer_text.insert('1.0', insights))
                
                elapsed = time.time() - self.insights_start_time
                self.root.after(0, lambda: self.update_status(f"✓ Document insights generated in {elapsed:.2f}s"))
            except Exception as e:
                error_msg = f"❌ Error generating insights: {str(e)}"
                self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
                self.root.after(0, lambda: self.answer_text.insert('1.0', error_msg))
                self.root.after(0, lambda: self.update_status(f"❌ Error: {str(e)}"))
            finally:
                self.processing = False
                self.root.after(0, lambda: self.insights_btn.config(state='normal'))
                self.root.after(0, lambda: self.ask_btn.config(state='normal'))
                self.root.after(0, lambda: self.summary_btn.config(state='normal'))
        
        # Run in background thread
        threading.Thread(target=generate_insights, daemon=True).start()
    
    def _animate_insights_processing(self):
        """Animate insights processing with elapsed time"""
        if not hasattr(self, 'processing') or not self.processing:
            return
        
        try:
            elapsed = time.time() - self.insights_start_time
            dots = "." * ((int(elapsed * 2) % 4))  # Animated dots
            
            # Update the processing message with elapsed time
            current_text = self.answer_text.get('1.0', 'end-1c')
            if "🤖 Generating" in current_text:
                self.answer_text.delete('1.0', 'end')
                self.answer_text.insert('1.0', 
                    f"🤖 Generating document collection insights...\n\n"
                    f"⏳ Analyzing documents{dots}\n"
                    f"⏱️  Elapsed: {elapsed:.1f}s")
            
            # Schedule next update
            self.root.after(500, self._animate_insights_processing)
            
        except Exception:
            pass
    
    def setup_gui(self):
        """Create modern, commercial-grade GUI"""
        self.root = tk.Tk()
        self.root.title("Doqurix - Intelligent Document Analysis")
        self.root.geometry("1500x900")
        self.root.minsize(1200, 700)
        
        # Modern color palette
        self.colors = {
            'bg_primary': '#0f172a',      # Dark navy background
            'bg_secondary': '#1e293b',    # Slightly lighter navy
            'bg_card': '#ffffff',          # White cards
            'bg_input': '#f8fafc',         # Light gray input bg
            'accent_primary': '#6366f1',   # Indigo accent
            'accent_secondary': '#8b5cf6', # Purple accent
            'accent_success': '#10b981',   # Emerald green
            'accent_warning': '#f59e0b',   # Amber
            'accent_danger': '#ef4444',    # Red
            'text_primary': '#0f172a',     # Dark text
            'text_secondary': '#64748b',   # Gray text
            'text_light': '#ffffff',       # White text
            'border': '#e2e8f0',           # Light border
            'highlight': '#fef08a',        # Yellow highlight
            'gradient_start': '#6366f1',   # Gradient start
            'gradient_end': '#8b5cf6',     # Gradient end
        }
        
        self.root.configure(bg=self.colors['bg_primary'])
        
        # Initialize mode variable
        self.rag_mode = tk.StringVar(value="basic")
        
        # Initialize agent selection variable
        self.agent_mode = tk.StringVar(value="None")
        
        # Configure ttk styles for modern look
        self.style = ttk.Style()
        self.style.theme_use('clam')
        
        # Custom button style
        self.style.configure('Modern.TButton',
                            font=('Segoe UI', 10, 'bold'),
                            padding=(20, 12))
        
        # Create menu bar with modern styling
        menubar = tk.Menu(self.root, bg=self.colors['bg_secondary'], fg=self.colors['text_light'],
                         activebackground=self.colors['accent_primary'], activeforeground='white',
                         font=('Segoe UI', 10))
        self.root.config(menu=menubar)
        
        # File menu
        file_menu = tk.Menu(menubar, tearoff=0, bg=self.colors['bg_card'], fg=self.colors['text_primary'],
                           activebackground=self.colors['accent_primary'], activeforeground='white',
                           font=('Segoe UI', 10))
        menubar.add_cascade(label="  File  ", menu=file_menu)
        file_menu.add_command(label="📂 Open Documents", command=self.upload_document)
        file_menu.add_separator()
        file_menu.add_command(label="🌐 Launch Web Version", command=self.launch_web_version)
        file_menu.add_separator()
        file_menu.add_command(label="❌ Exit", command=self.root.quit)
        
        # Help menu
        help_menu = tk.Menu(menubar, tearoff=0, bg=self.colors['bg_card'], fg=self.colors['text_primary'],
                           activebackground=self.colors['accent_primary'], activeforeground='white',
                           font=('Segoe UI', 10))
        menubar.add_cascade(label="  Help  ", menu=help_menu)
        help_menu.add_command(label="🔑 Activate License", command=self.show_license_dialog)
        help_menu.add_separator()
        help_menu.add_command(label="ℹ️ About Doqurix", command=self.show_about)
        
        # ============== HEADER ==============
        header = tk.Frame(self.root, bg=self.colors['bg_primary'], height=100)
        header.pack(fill='x', side='top')
        header.pack_propagate(False)
        
        # Header content container
        header_content = tk.Frame(header, bg=self.colors['bg_primary'])
        header_content.pack(fill='both', expand=True, padx=40, pady=15)
        
        # Left side - Logo and title
        title_frame = tk.Frame(header_content, bg=self.colors['bg_primary'])
        title_frame.pack(side='left', fill='y')
        
        # App icon and name
        logo_frame = tk.Frame(title_frame, bg=self.colors['bg_primary'])
        logo_frame.pack(anchor='w')
        
        title = tk.Label(logo_frame, text="◆ DOQURIX", 
                        font=('Segoe UI', 28, 'bold'), 
                        bg=self.colors['bg_primary'], fg=self.colors['text_light'])
        title.pack(side='left')
        
        version_label = tk.Label(logo_frame, text="  v1.0", 
                                font=('Segoe UI', 10), 
                                bg=self.colors['bg_primary'], fg=self.colors['text_secondary'])
        version_label.pack(side='left', pady=(12, 0))
        
        subtitle = tk.Label(title_frame, text="AI-Powered Document Intelligence Platform", 
                           font=('Segoe UI', 11), 
                           bg=self.colors['bg_primary'], fg=self.colors['text_secondary'])
        subtitle.pack(anchor='w', pady=(2, 0))
        
        # Right side - Web button and License indicator
        right_header = tk.Frame(header_content, bg=self.colors['bg_primary'])
        right_header.pack(side='right', fill='y')
        
        # Web Version button
        self.web_btn = tk.Button(right_header, text="🌐 Launch Web",
                                command=self.launch_web_version,
                                font=('Segoe UI', 11, 'bold'),
                                bg='#10a37f', fg='white',
                                activebackground='#0d8a6b',
                                relief='flat', padx=20, pady=10,
                                cursor='hand2', borderwidth=0)
        self.web_btn.pack(side='right', padx=(15, 0))
        
        # Add hover effects for web button
        def on_web_enter(e):
            self.web_btn.config(bg='#0d8a6b', relief='raised')
        
        def on_web_leave(e):
            self.web_btn.config(bg='#10a37f', relief='flat')
        
        self.web_btn.bind("<Enter>", on_web_enter)
        self.web_btn.bind("<Leave>", on_web_leave)
        
        # Trial/License indicator
        if self.trial_info and self.trial_info[0]:  # is_trial
            days = self.trial_info[1]
            if days <= 7:
                trial_color = self.colors['accent_danger']
                trial_bg = '#fef2f2'
            elif days <= 14:
                trial_color = self.colors['accent_warning']
                trial_bg = '#fffbeb'
            else:
                trial_color = self.colors['accent_success']
                trial_bg = '#ecfdf5'
            
            trial_frame = tk.Frame(right_header, bg=trial_bg, padx=12, pady=6)
            trial_frame.pack(side='right')
            
            self.trial_label = tk.Label(trial_frame, text=f"⏱ {days} days left",
                                  font=('Segoe UI', 9, 'bold'),
                                  bg=trial_bg, fg=trial_color,
                                  cursor='hand2')
            self.trial_label.pack()
            self.trial_label.bind('<Button-1>', lambda e: self.show_license_dialog())
            trial_frame.bind('<Button-1>', lambda e: self.show_license_dialog())
        elif self.trial_info:
            license_frame = tk.Frame(right_header, bg='#ecfdf5', padx=12, pady=6)
            license_frame.pack(side='right')
            self.trial_label = tk.Label(license_frame, text="✓ Licensed",
                                     font=('Segoe UI', 9, 'bold'),
                                     bg='#ecfdf5', fg=self.colors['accent_success'])
            self.trial_label.pack()
        
        # ============== MAIN CONTENT ==============
        main_container = tk.Frame(self.root, bg=self.colors['bg_primary'])
        main_container.pack(fill='both', expand=True, padx=30, pady=(0, 20))
        
        # ============== LEFT PANEL - DOCUMENTS ==============
        left_panel = tk.Frame(main_container, bg=self.colors['bg_card'], width=320)
        left_panel.pack(side='left', fill='y', padx=(0, 15))
        left_panel.pack_propagate(False)
        
        # Panel header with gradient effect simulation
        left_header = tk.Frame(left_panel, bg=self.colors['accent_primary'], height=60)
        left_header.pack(fill='x')
        left_header.pack_propagate(False)
        
        header_inner = tk.Frame(left_header, bg=self.colors['accent_primary'])
        header_inner.pack(expand=True)
        
        tk.Label(header_inner, text="📁", 
                font=('Segoe UI', 18), 
                bg=self.colors['accent_primary'], fg='white').pack(side='left', padx=(0, 8))
        tk.Label(header_inner, text="Documents", 
                font=('Segoe UI', 14, 'bold'), 
                bg=self.colors['accent_primary'], fg='white').pack(side='left')
        
        # Upload section with modern styling
        upload_section = tk.Frame(left_panel, bg=self.colors['bg_card'])
        upload_section.pack(fill='x', padx=20, pady=20)
        
        # Drag & drop style upload area
        upload_area = tk.Frame(upload_section, bg='#f1f5f9', highlightbackground=self.colors['border'],
                              highlightthickness=2)
        upload_area.pack(fill='x', pady=(0, 15))
        
        upload_inner = tk.Frame(upload_area, bg='#f1f5f9')
        upload_inner.pack(pady=20)
        
        tk.Label(upload_inner, text="📄", 
                font=('Segoe UI', 24), 
                bg='#f1f5f9', fg=self.colors['accent_primary']).pack()
        
        self.upload_btn = tk.Button(upload_inner, text="+ Upload PDF",
                                    command=self.upload_document,
                                    font=('Segoe UI', 11, 'bold'),
                                    bg=self.colors['accent_primary'], fg='white',
                                    activebackground='#4f46e5',
                                    relief='flat', padx=25, pady=10,
                                    cursor='hand2', state='disabled',
                                    borderwidth=0)
        self.upload_btn.pack(pady=(10, 5))
        
        tk.Label(upload_inner, text="or drag files here", 
                font=('Segoe UI', 9), 
                bg='#f1f5f9', fg=self.colors['text_secondary']).pack()
        
        # Document count with icon
        count_frame = tk.Frame(upload_section, bg=self.colors['bg_card'])
        count_frame.pack(fill='x')
        
        self.doc_count_label = tk.Label(count_frame, text="0 documents loaded", 
                                       font=('Segoe UI', 10), 
                                       bg=self.colors['bg_card'], fg=self.colors['text_secondary'])
        self.doc_count_label.pack(side='left')
        
        self.delete_btn = tk.Button(count_frame, text="🗑",
                                    command=self.delete_document,
                                    font=('Segoe UI', 12),
                                    bg=self.colors['bg_card'], fg=self.colors['accent_danger'],
                                    activebackground='#fee2e2',
                                    relief='flat', padx=8, pady=2,
                                    cursor='hand2', state='disabled',
                                    borderwidth=0)
        self.delete_btn.pack(side='right')
        
        # Document list with modern styling
        list_frame = tk.Frame(left_panel, bg=self.colors['bg_card'])
        list_frame.pack(fill='both', expand=True, padx=20, pady=(0, 20))
        
        # List label
        tk.Label(list_frame, text="Uploaded Files", 
                font=('Segoe UI', 10, 'bold'), 
                bg=self.colors['bg_card'], fg=self.colors['text_primary']).pack(anchor='w', pady=(0, 8))
        
        # Listbox container with border
        listbox_container = tk.Frame(list_frame, bg=self.colors['border'], padx=1, pady=1)
        listbox_container.pack(fill='both', expand=True)
        
        list_inner = tk.Frame(listbox_container, bg=self.colors['bg_input'])
        list_inner.pack(fill='both', expand=True)
        
        scrollbar = tk.Scrollbar(list_inner, width=8)
        scrollbar.pack(side='right', fill='y')
        
        self.doc_listbox = tk.Listbox(list_inner, yscrollcommand=scrollbar.set,
                                      font=('Segoe UI', 10),
                                      relief='flat', bd=0,
                                      bg=self.colors['bg_input'],
                                      fg=self.colors['text_primary'],
                                      selectbackground=self.colors['accent_primary'],
                                      selectforeground='white',
                                      highlightthickness=0,
                                      activestyle='none')
        self.doc_listbox.pack(side='left', fill='both', expand=True, padx=8, pady=8)
        scrollbar.config(command=self.doc_listbox.yview)
        
        # ============== MIDDLE PANEL - Q&A ==============
        middle_panel = tk.Frame(main_container, bg=self.colors['bg_card'])
        middle_panel.pack(side='left', fill='both', expand=True, padx=(0, 15))
        
        # Q&A Header
        middle_header = tk.Frame(middle_panel, bg=self.colors['accent_success'], height=60)
        middle_header.pack(fill='x')
        middle_header.pack_propagate(False)
        
        header_inner = tk.Frame(middle_header, bg=self.colors['accent_success'])
        header_inner.pack(expand=True)
        
        tk.Label(header_inner, text="💬", 
                font=('Segoe UI', 18), 
                bg=self.colors['accent_success'], fg='white').pack(side='left', padx=(0, 8))
        tk.Label(header_inner, text="Ask Questions", 
                font=('Segoe UI', 14, 'bold'), 
                bg=self.colors['accent_success'], fg='white').pack(side='left')
        
        # Question input section
        question_section = tk.Frame(middle_panel, bg=self.colors['bg_card'])
        question_section.pack(fill='x', padx=20, pady=18)
        
        # Question label with icon
        q_label_frame = tk.Frame(question_section, bg=self.colors['bg_card'])
        q_label_frame.pack(fill='x', pady=(0, 10))
        
        tk.Label(q_label_frame, text="Your Question", 
                font=('Segoe UI', 12, 'bold'), 
                bg=self.colors['bg_card'], fg=self.colors['text_primary']).pack(side='left')
        
        tk.Label(q_label_frame, text="Ask anything about your documents", 
                font=('Segoe UI', 9), 
                bg=self.colors['bg_card'], fg=self.colors['text_secondary']).pack(side='right')
        
        # Question entry with modern border
        entry_container = tk.Frame(question_section, bg=self.colors['border'], padx=2, pady=2)
        entry_container.pack(fill='x')
        
        self.question_entry = scrolledtext.ScrolledText(entry_container, 
                                                        height=2,
                                                        font=('Segoe UI', 11),
                                                        wrap='word',
                                                        relief='flat', bd=0,
                                                        bg=self.colors['bg_input'],
                                                        fg=self.colors['text_primary'],
                                                        insertbackground=self.colors['accent_primary'],
                                                        padx=15, pady=12)
        self.question_entry.pack(fill='x')
        
        # Mode selector with modern toggle style
        mode_frame = tk.Frame(question_section, bg=self.colors['bg_card'])
        mode_frame.pack(fill='x', pady=(12, 0))
        
        tk.Label(mode_frame, text="Search Mode", 
                font=('Segoe UI', 10, 'bold'), 
                bg=self.colors['bg_card'], fg=self.colors['text_primary']).pack(side='left', padx=(0, 10))
        
        # Radio buttons with modern styling
        radio_frame = tk.Frame(mode_frame, bg=self.colors['bg_card'])
        radio_frame.pack(side='left')
        
        basic_radio = tk.Radiobutton(radio_frame, text="⚡ Quick", 
                                     variable=self.rag_mode, value="basic",
                                     font=('Segoe UI', 9),
                                     bg=self.colors['bg_card'], fg=self.colors['text_primary'],
                                     activebackground=self.colors['bg_card'],
                                     selectcolor=self.colors['bg_input'],
                                     cursor='hand2')
        basic_radio.pack(side='left', padx=(0, 10))
        
        advanced_radio = tk.Radiobutton(radio_frame, text="🎯 Detailed", 
                                       variable=self.rag_mode, value="advanced",
                                       font=('Segoe UI', 9),
                                       bg=self.colors['bg_card'], fg=self.colors['text_primary'],
                                       activebackground=self.colors['bg_card'],
                                       selectcolor=self.colors['bg_input'],
                                       cursor='hand2')
        advanced_radio.pack(side='left')
        
        # Agent selection section
        agent_frame = tk.Frame(question_section, bg=self.colors['bg_card'])
        agent_frame.pack(fill='x', pady=(12, 0))
        
        tk.Label(agent_frame, text="AI Agent", 
                font=('Segoe UI', 10, 'bold'), 
                bg=self.colors['bg_card'], fg=self.colors['text_primary']).pack(side='left', padx=(0, 10))
        
        # Agent selector dropdown
        agent_dropdown_container = tk.Frame(agent_frame, bg=self.colors['border'], padx=1, pady=1)
        agent_dropdown_container.pack(side='left', fill='x', expand=True)
        
        # Configure combobox style
        self.style.configure('Agent.TCombobox',
                            fieldbackground=self.colors['bg_input'],
                            background=self.colors['bg_input'],
                            foreground=self.colors['text_primary'],
                            arrowcolor=self.colors['accent_primary'],
                            borderwidth=0,
                            relief='flat')
        
        self.agent_selector = ttk.Combobox(agent_dropdown_container,
                                          textvariable=self.agent_mode,
                                          values=["None", "Insights", "Tax Germany", "E-Commerce Germany", "BürokratAI"],
                                          state='readonly',
                                          font=('Segoe UI', 10),
                                          style='Agent.TCombobox',
                                          width=20)
        self.agent_selector.pack(fill='x', padx=5, pady=5)
        
        # Agent info label
        self.agent_info = tk.Label(agent_frame, text="None - Use your uploaded documents", 
                                   font=('Segoe UI', 9), 
                                   bg=self.colors['bg_card'], fg=self.colors['text_secondary'],
                                   wraplength=300, justify='left')
        self.agent_info.pack(side='left', padx=(10, 0))
        
        # Bind agent selection change
        self.agent_selector.bind('<<ComboboxSelected>>', self.on_agent_change)
        
        # Action buttons - on the same row, right side
        self.insights_btn = tk.Button(mode_frame, text="🤖 Insights",
                                     command=self.show_document_insights,
                                     font=('Segoe UI', 9, 'bold'),
                                     bg=self.colors['accent_secondary'], fg='white',
                                     activebackground='#7c3aed',
                                     relief='flat', padx=12, pady=6,
                                     cursor='hand2', state='disabled',
                                     borderwidth=0)
        self.insights_btn.pack(side='right', padx=(5, 0))
        
        self.summary_btn = tk.Button(mode_frame, text="📝 Summarize",
                                     command=lambda: self.ask_question_thread('summarize'),
                                     font=('Segoe UI', 9, 'bold'),
                                     bg=self.colors['accent_warning'], fg='white',
                                     activebackground='#d97706',
                                     relief='flat', padx=12, pady=6,
                                     cursor='hand2', state='disabled',
                                     borderwidth=0)
        self.summary_btn.pack(side='right', padx=(5, 0))
        
        self.ask_btn = tk.Button(mode_frame, text="🔍 Get Answer",
                                command=lambda: self.ask_question_thread('answer'),
                                font=('Segoe UI', 9, 'bold'),
                                bg=self.colors['accent_success'], fg='white',
                                activebackground='#059669',
                                relief='flat', padx=12, pady=6,
                                cursor='hand2', state='disabled',
                                borderwidth=0)
        self.ask_btn.pack(side='right')

        
        # Answer section
        answer_section = tk.Frame(middle_panel, bg=self.colors['bg_card'])
        answer_section.pack(fill='both', expand=True, padx=25, pady=(10, 25))
        
        # Answer label with icon
        a_label_frame = tk.Frame(answer_section, bg=self.colors['bg_card'])
        a_label_frame.pack(fill='x', pady=(0, 10))
        
        tk.Label(a_label_frame, text="💡 Answer", 
                font=('Segoe UI', 12, 'bold'), 
                bg=self.colors['bg_card'], fg=self.colors['text_primary']).pack(side='left')
        
        # Answer text with modern border
        answer_container = tk.Frame(answer_section, bg=self.colors['border'], padx=2, pady=2)
        answer_container.pack(fill='both', expand=True)
        
        self.answer_text = scrolledtext.ScrolledText(answer_container,
                                                     font=('Segoe UI', 11),
                                                     wrap='word',
                                                     relief='flat', bd=0,
                                                     bg='#fafbfc',
                                                     fg=self.colors['text_primary'],
                                                     padx=15, pady=15)
        self.answer_text.pack(fill='both', expand=True)
        
        # ============== RIGHT PANEL - CONTEXT ==============
        right_panel = tk.Frame(main_container, bg=self.colors['bg_card'], width=380)
        right_panel.pack(side='right', fill='y')
        right_panel.pack_propagate(False)
        
        # Context Header
        right_header = tk.Frame(right_panel, bg=self.colors['accent_secondary'], height=60)
        right_header.pack(fill='x')
        right_header.pack_propagate(False)
        
        header_inner = tk.Frame(right_header, bg=self.colors['accent_secondary'])
        header_inner.pack(expand=True)
        
        tk.Label(header_inner, text="🔍", 
                font=('Segoe UI', 18), 
                bg=self.colors['accent_secondary'], fg='white').pack(side='left', padx=(0, 8))
        tk.Label(header_inner, text="Source Context", 
                font=('Segoe UI', 14, 'bold'), 
                bg=self.colors['accent_secondary'], fg='white').pack(side='left')
        
        # Context content
        context_section = tk.Frame(right_panel, bg=self.colors['bg_card'])
        context_section.pack(fill='both', expand=True, padx=20, pady=20)
        
        # Info text
        info_label = tk.Label(context_section, text="Relevant passages with keywords", 
                font=('Segoe UI', 9), 
                bg=self.colors['bg_card'], fg=self.colors['text_secondary'],
                wraplength=250, justify='left')
        info_label.pack(anchor='w', pady=(0, 10))
        
        # Context text with warm background
        context_container = tk.Frame(context_section, bg=self.colors['border'], padx=2, pady=2)
        context_container.pack(fill='both', expand=True)
        
        self.context_text = scrolledtext.ScrolledText(context_container,
                                                      font=('Segoe UI', 10),
                                                      wrap='word',
                                                      relief='flat', bd=0,
                                                      bg='#fffbeb',
                                                      fg=self.colors['text_primary'],
                                                      padx=15, pady=15)
        self.context_text.pack(fill='both', expand=True)
        
        # Configure highlight tags
        self.context_text.tag_config('highlight', background=self.colors['highlight'], 
                                     foreground=self.colors['text_primary'], font=('Segoe UI', 10, 'bold'))
        self.context_text.tag_config('source', font=('Segoe UI', 10, 'bold'), 
                                     foreground=self.colors['accent_primary'])
        
        # ============== STATUS BAR ==============
        status_bar = tk.Frame(self.root, bg=self.colors['bg_secondary'], height=40)
        status_bar.pack(fill='x', side='bottom')
        status_bar.pack_propagate(False)
        
        status_inner = tk.Frame(status_bar, bg=self.colors['bg_secondary'])
        status_inner.pack(fill='both', expand=True, padx=30)
        
        # Status indicator dot
        self.status_dot = tk.Label(status_inner, text="●", 
                                   font=('Segoe UI', 10),
                                   bg=self.colors['bg_secondary'], fg=self.colors['accent_warning'])
        self.status_dot.pack(side='left', pady=10)
        
        self.status_label = tk.Label(status_inner, text=" Initializing...", 
                                     font=('Segoe UI', 10),
                                     bg=self.colors['bg_secondary'], fg=self.colors['text_secondary'],
                                     anchor='w')
        self.status_label.pack(side='left', fill='x', pady=10)
        
        # Copyright
        tk.Label(status_inner, text="© 2025 Doqurix", 
                font=('Segoe UI', 9),
                bg=self.colors['bg_secondary'], fg='#475569').pack(side='right', pady=10)
    
    def launch_web_version(self):
        """Launch the Bottle web version"""
        import threading
        import webbrowser
        import time
        import urllib.request
        
        bottle_app_path = self.app_dir / "bottle_app.py"
        
        # Check in _internal for PyInstaller builds
        if not bottle_app_path.exists():
            internal_path = self.app_dir / "_internal" / "bottle_app.py"
            if internal_path.exists():
                bottle_app_path = internal_path
            else:
                messagebox.showerror("Error", "Web app file not found.\n\nPlease ensure bottle_app.py exists in the application directory.")
                return
        
        if not self.models_loaded:
            messagebox.showwarning("Please Wait", "Models are still loading. Please wait and try again.")
            return
        
        # Check if server already running
        if hasattr(self, 'bottle_server_running') and self.bottle_server_running:
            webbrowser.open('http://localhost:8502')
            self.update_status("✓ Web version already running at http://localhost:8502")
            return
        
        try:
            self.update_status("🌐 Starting web server...")
            
            # Start Bottle server in background thread
            def run_bottle_server():
                try:
                    import sys
                    sys.path.insert(0, str(bottle_app_path.parent))
                    import bottle_app
                    bottle_app.run_server(self, port=8502)
                except Exception as e:
                    print(f"Bottle server error: {e}")
            
            server_thread = threading.Thread(target=run_bottle_server, daemon=True)
            server_thread.start()
            
            # Wait for server to be ready (with health check)
            self.update_status("🌐 Waiting for server...")
            server_ready = False
            
            for i in range(50):  # 5 seconds max
                try:
                    response = urllib.request.urlopen('http://localhost:8502', timeout=0.5)
                    server_ready = True
                    break
                except:
                    time.sleep(0.1)
            
            if server_ready:
                self.bottle_server_running = True
                webbrowser.open('http://localhost:8502')
                self.update_status("✓ Web version launched at http://localhost:8502")
            else:
                messagebox.showerror("Error", "Web server failed to start.\n\nPlease try again.")
                self.update_status("❌ Web server failed to start")
                
        except Exception as e:
            messagebox.showerror("Error", f"Failed to launch web version:\n{str(e)}")
            self.update_status(f"❌ Error: {str(e)}")
    
  
    def show_license_dialog(self):
        """Show license activation dialog"""
        license_manager = LicenseManager()
        
        # Check current status
        is_licensed, _, _ = license_manager.check_license()
        
        if is_licensed:
            messagebox.showinfo("Already Licensed", 
                "Doqurix is already activated!\n\nThank you for your purchase.")
            return
        
        # Get trial status
        trial_valid, days_remaining, _ = license_manager.get_trial_status()
        
        # Show dialog
        dialog = LicenseDialog(self.root, license_manager, 
                              days_remaining=days_remaining, 
                              is_expired=not trial_valid)
        result = dialog.show()
        
        if result:
            # Check if license was activated
            is_licensed, _, _ = license_manager.check_license()
            if is_licensed:
                # Update the trial label to show licensed
                if hasattr(self, 'trial_label'):
                    self.trial_label.config(text="✓ Licensed", fg='#27ae60', cursor='arrow')
                    self.trial_label.unbind('<Button-1>')
                messagebox.showinfo("Success", 
                    "License activated successfully!\n\nThank you for purchasing Doqurix.")
    
    def show_about(self):
        """Show about dialog"""
        license_manager = LicenseManager()
        is_licensed, _, _ = license_manager.check_license()
        
        if is_licensed:
            status = "Licensed Version"
        else:
            trial_valid, days, _ = license_manager.get_trial_status()
            if trial_valid:
                status = f"Trial Version ({days} days remaining)"
            else:
                status = "Trial Expired"
        
        about_text = f"""Doqurix v1.0.0

Intelligent Document Analysis
Smart Search • Context Highlighting

Status: {status}

Powered by:
• Advanced AI Language Model
• RAG Pipeline Technology

© 2025 Doqurix. All rights reserved.

Contact: sales@doqurix.com"""
        
        messagebox.showinfo("About Doqurix", about_text)
    
    def update_status(self, message):
        """Update status with visual indicator"""
        if hasattr(self, 'root'):
            def _update():
                self.status_label.config(text=f" {message}")
                # Update status dot color based on message
                if "✓" in message or "Ready" in message:
                    self.status_dot.config(fg='#10b981')  # Green
                elif "❌" in message or "Error" in message:
                    self.status_dot.config(fg='#ef4444')  # Red
                elif "🔄" in message or "Loading" in message or "Processing" in message:
                    self.status_dot.config(fg='#f59e0b')  # Yellow
                else:
                    self.status_dot.config(fg='#6366f1')  # Blue
            self.root.after(0, _update)
    
    def update_doc_count(self):
        """Update document count"""
        count = self.doc_listbox.size()
        self.doc_count_label.config(text=f"{count} document{'s' if count != 1 else ''} loaded")
    
    def upload_document(self):
        """Handle upload - supports multiple file selection"""
        if not self.models_loaded:
            messagebox.showwarning("Please Wait", "Models still loading.")
            return
        
        filetypes = [("Supported Documents", "*.pdf *.docx"), 
                     ("PDF files", "*.pdf"),
                     ("Word Documents", "*.docx"),
                     ("All files", "*.*")]
        
        file_paths = filedialog.askopenfilenames(
            title="Select Documents (Multiple Selection Enabled)",
            filetypes=filetypes
        )
        
        if file_paths:
            file_count = len(file_paths)
            if file_count > 1:
                self.update_status(f"📚 Processing {file_count} documents...")
            
            for file_path in file_paths:
                threading.Thread(target=self.process_document, args=(file_path, file_count), daemon=True).start()
    
    def process_document(self, file_path, total_files=1):
        """Process document"""
        try:
            filename = os.path.basename(file_path)
            self.update_status(f"📄 Processing {filename}...")
            chunks = self.add_document(file_path)
            
            self.root.after(0, lambda: self.doc_listbox.insert('end', f"📄 {filename} ({chunks} chunks)"))
            self.root.after(0, self.update_doc_count)
            self.root.after(0, lambda: self.delete_btn.config(state='normal'))
            self.update_status(f"✓ Added {filename}")
            
            # Show confirmation message only for single file or suppress for batch
            if total_files == 1:
                self.root.after(0, lambda: messagebox.showinfo("Document Added", 
                    f"'{filename}' has been successfully uploaded.\n\nExtracted {chunks} text chunks for analysis."))
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("Error", f"Failed to process document:\n{str(e)}"))
    
    def delete_document(self):
        """Delete selected document from the list and vector store"""
        selection = self.doc_listbox.curselection()
        
        if not selection:
            messagebox.showwarning("No Selection", "Please select a document to remove.")
            return
        
        # Get the selected item text
        selected_idx = selection[0]
        selected_text = self.doc_listbox.get(selected_idx)
        
        # Extract filename from the listbox text (format: "📄 filename.pdf (X chunks)")
        # Remove the emoji and extract just the filename
        filename = selected_text.replace("📄 ", "").split(" (")[0]
        
        # Confirm deletion
        confirm = messagebox.askyesno("Confirm Deletion", 
            f"Are you sure you want to remove '{filename}'?\n\nThis will delete the document from the database.")
        
        if not confirm:
            return
        
        try:
            # Remove from ChromaDB - get all IDs that contain this filename
            all_data = self.collection.get()
            ids_to_delete = []
            
            for i, metadata in enumerate(all_data['metadatas']):
                if metadata.get('source') == filename:
                    ids_to_delete.append(all_data['ids'][i])
            
            if ids_to_delete:
                self.collection.delete(ids=ids_to_delete)
            
            # Remove from listbox
            self.doc_listbox.delete(selected_idx)
            self.update_doc_count()
            
            # Rebuild BM25 index
            self.rebuild_bm25_index()
            
            # Disable delete button if no more documents
            if self.doc_listbox.size() == 0:
                self.delete_btn.config(state='disabled')
            
            self.update_status(f"✓ Removed {filename}")
            
            # Show confirmation
            messagebox.showinfo("Document Removed", 
                f"'{filename}' has been successfully removed.\n\n{len(ids_to_delete)} chunks deleted from the database.")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to remove document:\n{str(e)}")

    def add_document(self, file_path):
        """Add document - supports PDF and DOCX"""
        # Detect file type and extract text accordingly
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            text = self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        chunks = self.smart_chunk_text(text)
        
        for i, chunk in enumerate(chunks):
            embedding = self.embedder.encode(chunk).tolist()
            self.collection.add(
                embeddings=[embedding],
                documents=[chunk],
                metadatas=[{
                    "source": os.path.basename(file_path),
                    "chunk_id": i,
                    "page": i // 3
                }],
                ids=[f"{file_path}_{i}"]
            )
        
        self.rebuild_bm25_index()
        return len(chunks)
    
    def extract_text_from_pdf(self, file_path):
        """Extract from PDF"""
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def extract_text_from_docx(self, file_path):
        """Advanced extraction from DOCX with comprehensive content parsing"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library is required for Word document support. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract headers from all sections
            for section in doc.sections:
                if section.header:
                    header_text = self._extract_header_footer_text(section.header)
                    if header_text.strip():
                        full_text.append(f"[HEADER] {header_text}")
            
            # Process document body with structured extraction
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # Extract paragraph text with formatting context
                    paragraph = Paragraph(element, doc)
                    para_text = paragraph.text.strip()
                    
                    if para_text:
                        # Detect heading styles for better structure
                        if paragraph.style.name.startswith('Heading'):
                            level = paragraph.style.name.replace('Heading ', '')
                            full_text.append(f"\n[HEADING {level}] {para_text}\n")
                        elif paragraph.style.name == 'Title':
                            full_text.append(f"\n[TITLE] {para_text}\n")
                        elif paragraph.style.name == 'List Paragraph':
                            full_text.append(f"• {para_text}")
                        else:
                            full_text.append(para_text)
                
                elif isinstance(element, CT_Tbl):
                    # Advanced table extraction with structure preservation
                    table = Table(element, doc)
                    table_text = self._extract_table_content(table)
                    if table_text:
                        full_text.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
            
            # Extract footers from all sections
            for section in doc.sections:
                if section.footer:
                    footer_text = self._extract_header_footer_text(section.footer)
                    if footer_text.strip():
                        full_text.append(f"[FOOTER] {footer_text}")
            
            # Join with proper spacing and clean up
            text = '\n'.join(full_text)
            text = re.sub(r'\n{3,}', '\n\n', text)  # Remove excessive newlines
            
            return text
            
        except Exception as e:
            raise Exception(f"Error extracting text from Word document: {str(e)}")
    
    def _extract_header_footer_text(self, header_footer):
        """Extract text from header or footer"""
        text_parts = []
        for paragraph in header_footer.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        return ' '.join(text_parts)
    
    def _extract_table_content(self, table):
        """Extract and structure table content intelligently"""
        table_data = []
        
        # Process each row
        for i, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                # Extract text from each cell, handling merged cells
                cell_text = ' '.join(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip())
                row_data.append(cell_text)
            
            if any(row_data):  # Only add non-empty rows
                # Format first row as header if it looks like headers
                if i == 0 and all(cell.strip() for cell in row_data):
                    table_data.append(' | '.join(row_data))
                    table_data.append('-' * 50)  # Separator line
                else:
                    table_data.append(' | '.join(row_data))
        
        return '\n'.join(table_data)
    
    def smart_chunk_text(self, text, chunk_size=600, overlap=200):
        """Smart chunking"""
        text = re.sub(r'\s+', ' ', text).strip()
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            words = sentence.split()
            sentence_length = len(words)
            
            if current_length + sentence_length > chunk_size and current_chunk:
                chunk_text = ' '.join(current_chunk)
                if len(chunk_text.strip()) > 100:
                    chunks.append(chunk_text)
                
                overlap_sentences = []
                overlap_length = 0
                for s in reversed(current_chunk):
                    if overlap_length + len(s.split()) <= overlap:
                        overlap_sentences.insert(0, s)
                        overlap_length += len(s.split())
                    else:
                        break
                
                current_chunk = overlap_sentences
                current_length = overlap_length
            
            current_chunk.append(sentence)
            current_length += sentence_length
        
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            if len(chunk_text.strip()) > 100:
                chunks.append(chunk_text)
        
        return chunks
    
    def advanced_hybrid_search(self, question, n_results=15, alpha=0.7):
        """
        Advanced hybrid search with semantic prioritization
        
        Args:
            question: The search query
            n_results: Number of results to return
            alpha: Weight for vector search (0-1). Higher = more semantic.
                   Default 0.7 = 70% semantic, 30% keyword matching
        """
        # Check if collection has documents
        collection_data = self.collection.get()
        if not collection_data['documents'] or len(collection_data['documents']) == 0:
            return []
        
        question_embedding = self.embedder.encode(question).tolist()
        vector_results = self.collection.query(
            query_embeddings=[question_embedding],
            n_results=n_results
        )
        
        combined_docs = {}
        k = 60
        
        # Add vector results with HIGHER weight (alpha = 0.7 by default)
        for rank, (doc, metadata, distance) in enumerate(zip(
            vector_results['documents'][0],
            vector_results['metadatas'][0],
            vector_results['distances'][0]
        )):
            # Filter by minimum similarity
            # For cosine distance: 0 = identical, 2 = opposite
            similarity = 1 - (distance / 2)  # Convert to similarity score [0-1]
            
            # Skip documents that are too dissimilar (< 30% similar)
            if similarity < 0.3:
                continue
            
            doc_key = doc[:150]
            if doc_key not in combined_docs:
                combined_docs[doc_key] = {
                    'doc': doc,
                    'metadata': metadata,
                    'score': 0,
                    'vector_similarity': similarity  # Store for debugging
                }
            # Weight by alpha (default 70% for semantic understanding)
            combined_docs[doc_key]['score'] += alpha * (1 / (k + rank))
        
        # BM25 search (keywords) with LOWER weight (1 - alpha = 0.3 by default)
        if self.bm25:
            tokenized_query = question.lower().split()
            bm25_scores = self.bm25.get_scores(tokenized_query)
            top_bm25_indices = np.argsort(bm25_scores)[-n_results:][::-1]
            
            for rank, idx in enumerate(top_bm25_indices):
                if idx < len(self.bm25_corpus):
                    doc = self.bm25_corpus[idx]
                    doc_key = doc[:150]
                    
                    if doc_key not in combined_docs:
                        all_docs = self.collection.get()
                        if idx < len(all_docs['metadatas']):
                            metadata = all_docs['metadatas'][idx]
                            combined_docs[doc_key] = {
                                'doc': doc,
                                'metadata': metadata,
                                'score': 0
                            }
                    
                    if doc_key in combined_docs:
                        # Weight by (1 - alpha) (default 30% for keyword matching)
                        combined_docs[doc_key]['score'] += (1 - alpha) * (1 / (k + rank))
        
        sorted_docs = sorted(combined_docs.values(), key=lambda x: x['score'], reverse=True)
        return sorted_docs[:n_results]
    
    def rerank_documents(self, question, doc_results, top_k=5):
        """Rerank documents"""
        if not doc_results:
            return []
        
        documents = [d['doc'] for d in doc_results]
        pairs = [[question, doc] for doc in documents]
        scores = self.reranker.predict(pairs)
        
        for i, doc_result in enumerate(doc_results):
            doc_result['rerank_score'] = float(scores[i])
            doc_result['final_score'] = (
                0.3 * doc_result['score'] +
                0.7 * doc_result['rerank_score']
            )
        
        reranked = sorted(doc_results, key=lambda x: x['final_score'], reverse=True)
        return reranked[:top_k]
    
    def detect_language(self, text):
        """Detect primary language from text sample"""
        # Simple heuristic based on character ranges
        sample = text[:500]  # Check first 500 chars
        
        # Count Arabic characters
        arabic_chars = sum(1 for c in sample if '\u0600' <= c <= '\u06FF' or '\u0750' <= c <= '\u077F')
        # Count German-specific characters
        german_chars = sum(1 for c in sample if c in 'äöüßÄÖÜ')
        # Count Cyrillic (Russian, etc.)
        cyrillic_chars = sum(1 for c in sample if '\u0400' <= c <= '\u04FF')
        # Count Chinese/Japanese
        cjk_chars = sum(1 for c in sample if '\u4E00' <= c <= '\u9FFF' or '\u3040' <= c <= '\u30FF')
        
        total_chars = len(sample)
        if total_chars == 0:
            return 'english'
        
        # If > 20% Arabic, it's Arabic
        if arabic_chars / total_chars > 0.2:
            return 'arabic'
        # If has German chars and Latin script, likely German
        if german_chars > 0:
            return 'german'
        # If > 20% Cyrillic, it's Russian/Slavic
        if cyrillic_chars / total_chars > 0.2:
            return 'russian'
        # If > 20% CJK, it's Chinese/Japanese
        if cjk_chars / total_chars > 0.2:
            return 'chinese'
        
        # Default to English
        return 'english'

    def generate_answer(self, question, contexts, mode='answer', custom_system_prompt=None):
        """Generate answer using the AI model with language preservation"""
        # Get current RAG mode
        current_mode = self.rag_mode.get()
        
        # Use different number of contexts based on mode
        if current_mode == "basic":
            num_contexts = 2  # Basic: faster, 2 contexts
        else:
            num_contexts = 3  # Advanced: more detailed, 3 contexts
        
        # Handle both dict and string context formats
        if contexts and isinstance(contexts[0], dict):
            context_text = "\n\n".join([c['doc'] for c in contexts[:num_contexts]])
        else:
            # If contexts is already a string (for tax agent), use it directly
            context_text = contexts if isinstance(contexts, str) else "\n\n".join(contexts[:num_contexts])
        
        # Detect language from context
        detected_lang = self.detect_language(context_text)
        
        # Create language-specific instruction
        lang_instructions = {
            'arabic': 'IMPORTANT: The documents are in Arabic. You MUST answer in Arabic only. Do not translate to English.',
            'german': 'IMPORTANT: The documents are in German. You MUST answer in German only. Do not translate to English.',
            'russian': 'IMPORTANT: The documents are in Russian. You MUST answer in Russian only. Do not translate to English.',
            'chinese': 'IMPORTANT: The documents are in Chinese. You MUST answer in Chinese only. Do not translate to English.',
            'english': 'Answer clearly and concisely in English.'
        }
        
        lang_instruction = lang_instructions.get(detected_lang, lang_instructions['english'])
        
        # Use custom system prompt if provided (for specialized agents like tax agent)
        if custom_system_prompt:
            system_prompt = custom_system_prompt
        elif mode == 'summarize':
            system_prompt = f"""You are a professional assistant providing clear, well-structured summaries.
{lang_instruction}"""
        else:
            system_prompt = f"""You are a professional assistant providing clear, concise answers.
{lang_instruction}
Write in complete sentences and paragraphs, not bullet points or lists."""
        
        if mode == 'summarize':
            prompt = f"""<|im_start|>system
{system_prompt}<|im_end|>
<|im_start|>user
Provide a clear and professional summary of the following content. Use proper paragraphs and avoid bullet points or numbered lists. Write in a natural, flowing narrative style.

Content:
{context_text}

Summary:<|im_end|>
<|im_start|>assistant
"""
        else:
            prompt = f"""<|im_start|>system
{system_prompt}<|im_end|>
<|im_start|>user
Based on the following context, answer the question in a clear, professional manner. Write your answer in flowing paragraphs without using bullet points, numbered lists, or special formatting.

Context:
{context_text}

Question: {question}

Answer:<|im_end|>
<|im_start|>assistant
"""
        
        # Generate with llama.cpp - Optimized for speed
        output = self.llm(
            prompt,
            max_tokens=300,  # Reduced for faster generation
            temperature=0.7,
            top_p=0.9,
            repeat_penalty=1.1,
            stop=["<|im_end|>", "<|im_start|>"],
            top_k=40  # Limit sampling for faster generation
        )
        
        answer = output['choices'][0]['text'].strip()
        return answer
    
    def display_contexts(self, contexts, question):
        """Display contexts with highlighting"""
        self.context_text.delete('1.0', 'end')
        
        for i, ctx in enumerate(contexts, 1):
            source_text = f"\n{'─'*50}\n"
            source_text += f"📄 Source {i}: {ctx['metadata']['source']}\n"
            source_text += f"📍 Page ~{ctx['metadata']['page']}\n"
            source_text += f"{'─'*50}\n\n"
            
            self.context_text.insert('end', source_text, 'source')
            
            doc_text = ctx['doc']
            keywords = [w.lower() for w in question.split() if len(w) > 3]
            
            last_pos = 0
            for match in re.finditer(r'\b\w+\b', doc_text):
                word = match.group()
                start, end = match.span()
                
                self.context_text.insert('end', doc_text[last_pos:start])
                
                if word.lower() in keywords:
                    self.context_text.insert('end', word, 'highlight')
                else:
                    self.context_text.insert('end', word)
                
                last_pos = end
            
            self.context_text.insert('end', doc_text[last_pos:] + '\n\n')
    
    def ask_question_thread(self, mode='answer'):
        """Handle question"""
        question = self.question_entry.get('1.0', 'end-1c').strip()
        
        if mode == 'summarize':
            if self.doc_listbox.size() == 0:
                messagebox.showwarning("No Documents", "Upload documents first.")
                return
            question = "Provide a comprehensive summary"
        elif not question:
            messagebox.showwarning("Input Required", "Enter a question.")
            return
        
        if not self.models_loaded:
            messagebox.showwarning("Please Wait", "Models still loading.")
            return
        
        self.ask_btn.config(state='disabled')
        self.summary_btn.config(state='disabled')
        self.answer_text.delete('1.0', 'end')
        self.context_text.delete('1.0', 'end')
        
        # Show animated processing indicator
        self.processing = True
        self.answer_text.insert('1.0', "🔍 Analyzing your question...\n\n⏳ Processing")
        self.update_status("🤔 Processing your question...")
        
        # Start animation
        self._animate_processing()
        
        threading.Thread(target=self.process_question, args=(question, mode), daemon=True).start()
    
    def _animate_processing(self):
        """Animate processing dots with elapsed time"""
        if not hasattr(self, 'processing') or not self.processing:
            return
        
        try:
            import time
            
            # Calculate elapsed time if available
            elapsed_str = ""
            if hasattr(self, 'processing_start_time'):
                elapsed = time.time() - self.processing_start_time
                elapsed_str = f" ({elapsed:.1f}s)"
            
            # Get current text and update dots
            current_text = self.answer_text.get('1.0', 'end')
            if "Processing" in current_text:
                # Count current dots
                dot_count = current_text.count('•')
                if dot_count >= 5:
                    # Reset dots with time
                    self.answer_text.delete('1.0', 'end')
                    self.answer_text.insert('1.0', f"🔍 Analyzing your question...\n\n⏳ Processing{elapsed_str} •")
                else:
                    # Update with dots and time - delete old content and rewrite
                    self.answer_text.delete('1.0', 'end')
                    dots = '•' * (dot_count + 1)
                    self.answer_text.insert('1.0', f"🔍 Analyzing your question...\n\n⏳ Processing{elapsed_str} {dots}")
            
            # Schedule next animation frame
            self.root.after(400, self._animate_processing)
        except:
            pass
    
    def process_question(self, question, mode):
        """Process question with intelligent agent workflows"""
        import time
        start_time = time.time()  # Track start time
        self.processing_start_time = start_time  # Store for animation
        
        try:
            # Get current agent mode - map display to internal
            display_to_internal = {
                "None": "none",
                "Insights": "insights",
                "Tax Germany": "tax_germany",
                "E-Commerce Germany": "ecommerce_germany",
                "BürokratAI": "buerokratai_germany"
            }
            agent_display = self.agent_mode.get()
            agent = display_to_internal.get(agent_display, "none")
            
            # Handle different agents
            if agent == "tax_germany":
                # Use tax agent collection instead of user documents
                self.process_tax_agent_question(question, mode, start_time)
                return
            elif agent == "ecommerce_germany":
                # Use e-commerce agent for product search
                self.process_ecommerce_question(question, mode, start_time)
                return
            elif agent == "buerokratai_germany":
                # Use BürokratAI agent for immigration questions
                self.process_buerokratai_question(question, mode, start_time)
                return
            elif agent == "insights":
                # Use insights agent with user documents
                # Enhanced document analysis
                pass  # Continue with normal flow but with insights
            
            # Standard processing for "none" agent or insights agent with documents
            # Get current RAG mode
            current_mode = self.rag_mode.get()
            
            # Agent-enhanced document retrieval
            if hasattr(self, 'agents'):
                # Use cross-language retrieval for better document discovery
                all_docs = self.collection.get()
                if all_docs['documents']:
                    documents = []
                    for i, (doc, metadata) in enumerate(zip(all_docs['documents'], all_docs['metadatas'])):
                        documents.append({
                            'content': doc,
                            'metadata': metadata,
                            'id': all_docs['ids'][i]
                        })
                    
                    # Auto-triage documents for better relevance
                    triaged_docs = self.agents.auto_triage_documents(documents[:50], question)  # Limit for performance
                    
                    # Use only high-priority documents for search if available
                    high_priority = [d for d in triaged_docs if d['priority'] == 'high']
                    if high_priority and len(high_priority) >= 3:
                        # Focus search on high-priority docs
                        priority_ids = [d['document']['id'] for d in high_priority[:20]]
                        doc_results = self.advanced_hybrid_search_filtered(question, priority_ids, n_results=15 if current_mode == "advanced" else 8)
                    else:
                        # Fallback to standard search
                        doc_results = self.advanced_hybrid_search(question, n_results=15 if current_mode == "advanced" else 8)
                else:
                    doc_results = []
            else:
                # Fallback to standard search
                if current_mode == "basic":
                    doc_results = self.advanced_hybrid_search(question, n_results=8)
                    top_k = 3
                else:
                    doc_results = self.advanced_hybrid_search(question, n_results=15)
                    top_k = 5
            
            if not doc_results:
                self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
                self.root.after(0, lambda: self.answer_text.insert('1.0', 
                    "❌ No documents found. Upload PDFs first."))
                return
            
            # Set top_k based on mode
            top_k = 5 if current_mode == "advanced" else 3
            mode_label = "🎯 Advanced" if current_mode == "advanced" else "⚡ Basic"
            
            # Rerank documents
            reranked_contexts = self.rerank_documents(question, doc_results, top_k=top_k)
            self.root.after(0, lambda: self.display_contexts(reranked_contexts, question))
            
            # Generate initial answer
            answer = self.generate_answer(question, reranked_contexts, mode)
            
            # Agent-enhanced answer refinement
            if hasattr(self, 'agents') and current_mode == "advanced":
                try:
                    # Citation-aware answer refinement
                    refined_result = self.agents.citation_aware_refinement(answer, reranked_contexts)
                    if refined_result['validation_status'] == 'validated':
                        answer = refined_result['refined_answer']
                except Exception:
                    pass  # Use original answer if refinement fails
            
            # Calculate elapsed time
            elapsed_time = time.time() - start_time
            
            # Format enterprise-grade response with timing
            result = f"{answer}\n\n"
            result += f"⏱️ Processed in {elapsed_time:.2f} seconds\n\n"
            result += "REFERENCES\n\n"
            
            for i, ctx in enumerate(reranked_contexts, 1):
                source = ctx['metadata']['source']
                page = ctx['metadata']['page']
                result += f"  [{i}] {source}\n"
                result += f"      Page: {page}\n\n"
            
            # Generate follow-up questions for advanced mode
            if hasattr(self, 'agents') and current_mode == "advanced":
                try:
                    followup_questions = self.agents.generate_followup_questions(question, answer, reranked_contexts)
                    if followup_questions:
                        result += "💡 SUGGESTED FOLLOW-UP QUESTIONS\n\n"
                        for i, fq in enumerate(followup_questions, 1):
                            result += f"  {i}. {fq}\n"
                        result += "\n"
                except Exception:
                    pass  # Skip follow-ups if generation fails
            
            self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
            self.root.after(0, lambda: self.answer_text.insert('1.0', result))
            self.update_status(f"✓ Answer generated ({mode_label} mode) in {elapsed_time:.2f}s")
            
        except Exception as e:
            error_msg = f"❌ Error: {str(e)}"
            self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
            self.root.after(0, lambda: self.answer_text.insert('1.0', error_msg))
            self.update_status(f"❌ Error: {str(e)}")
        finally:
            self.processing = False  # Stop animation
            self.root.after(0, lambda: self.ask_btn.config(state='normal'))
            self.root.after(0, lambda: self.summary_btn.config(state='normal'))
    
    def process_tax_agent_question(self, question: str, mode: str, start_time: float):
        """
        Process questions using the tax agent - searches pre-loaded German tax knowledge.
        Does not require user-uploaded documents.
        """
        try:
            # Check if tax collection is initialized
            if not hasattr(self, 'tax_collection') or self.tax_collection is None:
                error_msg = "❌ Tax agent collection is not initialized. Please restart the application."
                self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
                self.root.after(0, lambda: self.answer_text.insert('1.0', error_msg))
                self.update_status(error_msg)
                return
            
            self.update_status(f"🔍 Searching German tax knowledge base...")
            
            # Generate query embedding
            query_embedding = self.embedder.encode([question], convert_to_tensor=False)[0]
            
            # Search tax collection (semantic search only - no BM25 for tax knowledge)
            search_results = self.tax_collection.query(
                query_embeddings=[query_embedding.tolist()],
                n_results=15,  # Get more initial results for better reranking
                include=['documents', 'metadatas', 'distances']
            )
            
            # Extract results
            documents = search_results['documents'][0] if search_results['documents'] else []
            metadatas = search_results['metadatas'][0] if search_results['metadatas'] else []
            distances = search_results['distances'][0] if search_results['distances'] else []
            
            if not documents:
                no_results_msg = ("I don't have specific information about that in my German tax knowledge base. "
                                "Please try rephrasing your question or ask about:\n\n"
                                "• Income Tax (Einkommensteuer) - brackets, classes, deductions\n"
                                "• VAT (Umsatzsteuer) - rates, registration, exemptions\n"
                                "• Corporate Tax (Körperschaftsteuer) - rates, regulations\n"
                                "• Trade Tax (Gewerbesteuer) - municipal rates, calculations\n"
                                "• Church Tax (Kirchensteuer) - rates, opt-out\n"
                                "• Capital Gains Tax (Abgeltungsteuer)\n"
                                "• Real Estate Transfer Tax (Grunderwerbsteuer)\n"
                                "• Inheritance & Gift Tax (Erbschaft- und Schenkungsteuer)\n"
                                "• Social Security Contributions")
                self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
                self.root.after(0, lambda: self.answer_text.insert('1.0', no_results_msg))
                self.update_status("⚠ No relevant tax information found")
                return
            
            self.update_status(f"📊 Reranking {len(documents)} tax knowledge chunks...")
            
            # Prepare contexts for reranking
            contexts_for_reranking = []
            for doc, meta, dist in zip(documents, metadatas, distances):
                contexts_for_reranking.append({
                    'text': doc,
                    'metadata': meta,
                    'distance': dist
                })
            
            # Rerank using CrossEncoder
            pairs = [[question, ctx['text']] for ctx in contexts_for_reranking]
            rerank_scores = self.reranker.predict(pairs)
            
            # Sort by rerank score (higher is better)
            for i, ctx in enumerate(contexts_for_reranking):
                ctx['rerank_score'] = float(rerank_scores[i])
            
            reranked_contexts = sorted(contexts_for_reranking, 
                                      key=lambda x: x['rerank_score'], 
                                      reverse=True)
            
            # Take top 3 after reranking (reduced for smaller context window)
            top_contexts = reranked_contexts[:3]
            
            # Build context string for LLM - truncate each context to 400 chars to fit in model's context window
            context_str = ""
            for i, ctx in enumerate(top_contexts, 1):
                source = ctx['metadata'].get('source', 'Unknown')
                # Truncate context text to fit within token limits
                truncated_text = ctx['text'][:400] + "..." if len(ctx['text']) > 400 else ctx['text']
                context_str += f"[Source {i} - {source}]\n{truncated_text}\n\n"
            
            self.update_status(f"💭 Generating expert tax answer...")
            
            # Generate answer using LLM with tax-specific system prompt (shortened for context window)
            tax_system_prompt = """You are a German tax expert. Provide accurate information about German taxes including rates, thresholds, and procedures. Use the provided context to give precise answers."""

            answer = self.generate_answer(question, context_str, custom_system_prompt=tax_system_prompt)
            
            # Calculate elapsed time
            elapsed_time = time.time() - start_time
            
            # Format result for display
            result = f"🇩🇪 GERMAN TAX AGENT ANSWER\n{'='*60}\n\n"
            result += f"Question: {question}\n\n"
            result += f"{'─'*60}\n\n"
            result += f"{answer}\n\n"
            result += f"{'─'*60}\n\n"
            result += f"📚 KNOWLEDGE BASE SOURCES USED\n\n"
            
            for i, ctx in enumerate(top_contexts, 1):
                source = ctx['metadata'].get('source', 'Unknown')
                score = ctx['rerank_score']
                snippet = ctx['text'][:200].replace('\n', ' ') + "..."
                result += f"  [{i}] {source} (Relevance: {score:.3f})\n"
                result += f"      {snippet}\n\n"
            
            result += f"✓ Processed in {elapsed_time:.2f}s using German tax knowledge base\n"
            
            # Update UI
            self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
            self.root.after(0, lambda: self.answer_text.insert('1.0', result))
            
            # Also show context in context pane
            context_display = "🇩🇪 TAX KNOWLEDGE BASE CONTEXT\n" + "="*60 + "\n\n"
            for i, ctx in enumerate(top_contexts, 1):
                source = ctx['metadata'].get('source', 'Unknown')
                score = ctx['rerank_score']
                context_display += f"SOURCE {i}: {source}\n"
                context_display += f"Relevance Score: {score:.3f}\n"
                context_display += f"{'-'*60}\n"
                context_display += f"{ctx['text']}\n\n"
                context_display += f"{'='*60}\n\n"
            
            self.root.after(0, lambda: self.context_text.delete('1.0', 'end'))
            self.root.after(0, lambda: self.context_text.insert('1.0', context_display))
            
            self.update_status(f"✓ Tax agent answer generated in {elapsed_time:.2f}s")
            
        except Exception as e:
            error_msg = f"❌ Error processing tax question: {str(e)}\n\nPlease try again or contact support."
            self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
            self.root.after(0, lambda: self.answer_text.insert('1.0', error_msg))
            self.update_status(f"❌ Error: {str(e)}")
        finally:
            self.processing = False  # Stop animation
            self.root.after(0, lambda: self.ask_btn.config(state='normal'))
            self.root.after(0, lambda: self.summary_btn.config(state='normal'))
    
    def process_buerokratai_question(self, question: str, mode: str, start_time: float):
        """
        Process questions using the BürokratAI agent - searches pre-loaded German immigration knowledge.
        Does not require user-uploaded documents.
        """
        try:
            from buerokratai_agent import BUEROKRATAI_SYSTEM_PROMPT, classify_topic, get_relevant_links
            
            # Check if BürokratAI collection is initialized
            if not hasattr(self, 'buerokratai_collection') or self.buerokratai_collection is None:
                error_msg = "❌ BürokratAI agent collection is not initialized. Please restart the application."
                self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
                self.root.after(0, lambda: self.answer_text.insert('1.0', error_msg))
                self.update_status(error_msg)
                return
            
            self.update_status(f"🔍 Searching German immigration knowledge base...")
            
            # Classify the topic for better context
            topics = classify_topic(question)
            
            # Generate query embedding
            query_embedding = self.embedder.encode([question], convert_to_tensor=False)[0]
            
            # Search BürokratAI collection
            search_results = self.buerokratai_collection.query(
                query_embeddings=[query_embedding.tolist()],
                n_results=15,
                include=['documents', 'metadatas', 'distances']
            )
            
            # Extract results
            documents = search_results['documents'][0] if search_results['documents'] else []
            metadatas = search_results['metadatas'][0] if search_results['metadatas'] else []
            distances = search_results['distances'][0] if search_results['distances'] else []
            
            if not documents:
                no_results_msg = ("I don't have specific information about that in my German immigration knowledge base. "
                                "Please try rephrasing your question or ask about:\n\n"
                                "📋 Registration & Documents:\n"
                                "• Anmeldung (Address Registration)\n"
                                "• Tax ID (Steuer-ID) Application\n"
                                "• Health Insurance Requirements\n\n"
                                "🛂 Visas & Residence Permits:\n"
                                "• EU Blue Card\n"
                                "• Student Visa\n"
                                "• Job Seeker Visa\n"
                                "• Family Reunification\n\n"
                                "🏠 Living in Germany:\n"
                                "• Renting an Apartment\n"
                                "• Opening a Bank Account\n"
                                "• Driver's License Conversion")
                self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
                self.root.after(0, lambda: self.answer_text.insert('1.0', no_results_msg))
                self.update_status("⚠ No relevant immigration information found")
                return
            
            self.update_status(f"📊 Reranking {len(documents)} immigration knowledge chunks...")
            
            # Prepare contexts for reranking
            contexts_for_reranking = []
            for doc, meta, dist in zip(documents, metadatas, distances):
                contexts_for_reranking.append({
                    'text': doc,
                    'metadata': meta,
                    'distance': dist
                })
            
            # Rerank using CrossEncoder
            pairs = [[question, ctx['text']] for ctx in contexts_for_reranking]
            rerank_scores = self.reranker.predict(pairs)
            
            # Sort by rerank score (higher is better)
            for i, ctx in enumerate(contexts_for_reranking):
                ctx['rerank_score'] = float(rerank_scores[i])
            
            reranked_contexts = sorted(contexts_for_reranking, 
                                      key=lambda x: x['rerank_score'], 
                                      reverse=True)
            
            # Take top 4 after reranking (immigration needs more context)
            top_contexts = reranked_contexts[:4]
            
            # Build context string for LLM
            context_str = ""
            for i, ctx in enumerate(top_contexts, 1):
                source = ctx['metadata'].get('source', 'Unknown')
                truncated_text = ctx['text'][:500] + "..." if len(ctx['text']) > 500 else ctx['text']
                context_str += f"[Source {i} - {source}]\n{truncated_text}\n\n"
            
            self.update_status(f"💭 Generating immigration guidance...")
            
            # Generate answer using LLM with BürokratAI system prompt
            answer = self.generate_answer(question, context_str, custom_system_prompt=BUEROKRATAI_SYSTEM_PROMPT)
            
            # Calculate elapsed time
            elapsed_time = time.time() - start_time
            
            # Get relevant links
            relevant_links = get_relevant_links(topics)
            
            # Format result for display
            result = f"🏛️ BÜROKRATAI - IMMIGRATION ASSISTANT\n{'='*60}\n\n"
            result += f"Question: {question}\n\n"
            result += f"{'─'*60}\n\n"
            result += f"{answer}\n\n"
            result += f"{'─'*60}\n\n"
            result += f"📚 KNOWLEDGE BASE SOURCES USED\n\n"
            
            for i, ctx in enumerate(top_contexts, 1):
                source = ctx['metadata'].get('source', 'Unknown')
                score = ctx['rerank_score']
                snippet = ctx['text'][:200].replace('\n', ' ') + "..."
                result += f"  [{i}] {source} (Relevance: {score:.3f})\n"
                result += f"      {snippet}\n\n"
            
            if relevant_links:
                result += f"🔗 USEFUL OFFICIAL LINKS\n\n"
                for name, url in relevant_links[:3]:
                    result += f"  • {name}: {url}\n"
                result += "\n"
            
            result += f"✓ Processed in {elapsed_time:.2f}s using German immigration knowledge base\n\n"
            result += "⚠️ DISCLAIMER: This information is for guidance only.\n"
            result += "   Please verify with official German authorities for current regulations."
            
            # Update UI
            self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
            self.root.after(0, lambda: self.answer_text.insert('1.0', result))
            
            # Also show context in context pane
            context_display = "🏛️ BÜROKRATAI KNOWLEDGE BASE CONTEXT\n" + "="*60 + "\n\n"
            for i, ctx in enumerate(top_contexts, 1):
                source = ctx['metadata'].get('source', 'Unknown')
                score = ctx['rerank_score']
                context_display += f"SOURCE {i}: {source}\n"
                context_display += f"Relevance Score: {score:.3f}\n"
                context_display += f"{'-'*60}\n"
                context_display += f"{ctx['text']}\n\n"
                context_display += f"{'='*60}\n\n"
            
            self.root.after(0, lambda: self.context_text.delete('1.0', 'end'))
            self.root.after(0, lambda: self.context_text.insert('1.0', context_display))
            
            self.update_status(f"✓ BürokratAI answer generated in {elapsed_time:.2f}s")
            
        except Exception as e:
            error_msg = f"❌ Error processing immigration question: {str(e)}\n\nPlease try again or contact support."
            self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
            self.root.after(0, lambda: self.answer_text.insert('1.0', error_msg))
            self.update_status(f"❌ Error: {str(e)}")
        finally:
            self.processing = False  # Stop animation
            self.root.after(0, lambda: self.ask_btn.config(state='normal'))
            self.root.after(0, lambda: self.summary_btn.config(state='normal'))
    
    def process_ecommerce_question(self, question, mode, start_time):
        """Process e-commerce product search with professional agent"""
        import time
        from ecommerce_agent import ECommerceAgent
        
        try:
            # Show loading animation
            self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
            self.root.after(0, lambda: self.answer_text.insert('1.0', "🛒 Extracting search keywords...\n\n⏳ Processing •"))
            
            # Start animation
            self.processing = True
            self.root.after(400, self._animate_processing)
            
            # Initialize professional e-commerce agent
            agent = ECommerceAgent(cache_dir='./cache/ecommerce')
            
            try:
                # ALWAYS use LLM to extract proper search keywords first
                print(f"🔑 Extracting search keywords from: '{question}'")
                search_keywords = agent.extract_search_keywords(question, self.llm)
                print(f"✓ Search keywords: '{search_keywords}'")
                
                # Update UI to show keyword extraction
                self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
                self.root.after(0, lambda: self.answer_text.insert('1.0', f"🛒 Searching German retailers...\n\n🔑 Keywords: {search_keywords}\n\n⏳ Processing •"))
                
                optimized_query = search_keywords  # Keep track of what we searched for
                
                # First attempt: Search with extracted keywords
                products = agent.search_products(search_keywords, max_results=15)
                
                # If no products found, try additional optimization
                if not products or len(products) == 0:
                    print(f"⚠ No products found for: {search_keywords}")
                    print(f"🤖 Using LLM to further optimize search query...")
                    
                    try:
                        # Use LLM to optimize the query
                        optimization_prompt = """Du bist ein E-Commerce-Suchexperte. Optimiere die folgende Produktsuchanfrage für bessere Ergebnisse:

1. Verwende gängige Produktkategorien (z.B. "günstig" → "budget", "billig" → "preiswert")
2. Füge relevante Suchbegriffe hinzu (z.B. "phone" → "smartphone")
3. Entferne zu spezifische oder unklare Begriffe
4. Nutze deutsche Standardbegriffe, die Händler verwenden

Gib NUR die optimierte Suchanfrage zurück, keine Erklärung."""
                        
                        opt_response = self.llm.create_chat_completion(
                            messages=[
                                {"role": "system", "content": optimization_prompt},
                                {"role": "user", "content": f"Optimiere diese Suche: {search_keywords}"}
                            ],
                            max_tokens=50,
                            temperature=0.3
                        )
                        
                        optimized_query = opt_response['choices'][0]['message']['content'].strip()
                        print(f"✓ Further optimized query: {optimized_query}")
                        
                        # Update UI to show optimization
                        self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
                        self.root.after(0, lambda: self.answer_text.insert('1.0', f"🛒 Searching German retailers...\n\n🔑 Keywords: {optimized_query}\n\n⏳ Processing •"))
                        
                        # Second attempt with optimized query
                        if optimized_query and optimized_query != search_keywords:
                            products = agent.search_products(optimized_query, max_results=15)
                            print(f"✓ Found {len(products)} products with optimized query")
                        
                    except Exception as e:
                        print(f"✗ Query optimization failed: {e}")
                
                # Final fallback
                if not products or len(products) == 0:
                    print(f"📋 Using intelligent fallback results...")
                    products = agent.get_fallback_results(optimized_query or search_keywords)
                
                # Format results with LLM summarization
                formatted_result = self.format_ecommerce_results_pro(question, products, optimized_query)
                
                # Calculate elapsed time
                elapsed_time = time.time() - start_time
                
                # Update UI with results
                self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
                self.root.after(0, lambda: self.answer_text.insert('1.0', formatted_result + f"\n\n✓ Search completed in {elapsed_time:.2f}s"))
                
                # Update context pane with detailed product data
                context_display = "🛒 PROFESSIONAL PRODUCT SEARCH RESULTS\n" + "="*60 + "\n\n"
                context_display += f"Original Query: {question}\n"
                context_display += f"Search Keywords: {optimized_query}\n\n"
                for i, product in enumerate(products[:10], 1):
                    context_display += f"PRODUCT {i}:\n"
                    context_display += f"Title: {product.title}\n"
                    context_display += f"Merchant: {product.merchant}\n"
                    if product.price:
                        context_display += f"Price: {product.price:.2f} {product.currency}\n"
                    if product.rating:
                        context_display += f"Rating: {product.rating}/5.0\n"
                    context_display += f"Availability: {product.availability}\n"
                    context_display += f"URL: {product.url}\n"
                    context_display += f"{'-'*60}\n\n"
                
                self.root.after(0, lambda: self.context_text.delete('1.0', 'end'))
                self.root.after(0, lambda: self.context_text.insert('1.0', context_display))
                
                self.update_status(f"✓ E-commerce search completed in {elapsed_time:.2f}s")
                
            finally:
                agent.close()
            
        except Exception as e:
            error_msg = f"❌ Error searching products: {str(e)}\n\nPlease try again with a different search query."
            self.root.after(0, lambda: self.answer_text.delete('1.0', 'end'))
            self.root.after(0, lambda: self.answer_text.insert('1.0', error_msg))
            self.update_status(f"❌ Error: {str(e)}")
        finally:
            self.processing = False
            self.root.after(0, lambda: self.ask_btn.config(state='normal'))
            self.root.after(0, lambda: self.summary_btn.config(state='normal'))
    
    def format_ecommerce_results_pro(self, query, products, optimized_query=None):
        """Format professional product search results with LLM analysis"""
        from ecommerce_agent import Product
        
        if not products:
            return "❌ Keine Produkte gefunden. Versuchen Sie eine andere Suchanfrage."
        
        # Create rich context from product objects
        context = "PROFESSIONELLE PRODUKTSUCHE - ERGEBNISSE:\n\n"
        if optimized_query and optimized_query != query:
            context += f"[Original: {query} → Optimiert: {optimized_query}]\n\n"
        
        for i, product in enumerate(products[:10], 1):
            context += f"{i}. {product.title}\n"
            context += f"   Händler: {product.merchant}\n"
            
            if product.price:
                context += f"   Preis: {product.price:.2f} {product.currency}\n"
            
            if product.rating:
                context += f"   Bewertung: {product.rating}/5.0 Sterne\n"
            
            if product.reviews_count:
                context += f"   Anzahl Bewertungen: {product.reviews_count}\n"
            
            context += f"   Verfügbarkeit: {product.availability}\n"
            context += f"   URL: {product.url}\n"
            
            if product.description:
                context += f"   Beschreibung: {product.description[:200]}\n"
            
            context += "\n"
        
        # Use LLM to analyze and compare products professionally
        system_prompt = """Sie sind ein Experte für E-Commerce im deutschen Markt. Analysieren Sie die Produktsuchergebnisse und bieten Sie:

1. Eine kurze Zusammenfassung der gefundenen Produkte
2. Wichtige Features und Spezifikationen (falls verfügbar)
3. Preisvergleich und Preis-Leistungs-Verhältnis
4. Empfehlung: Welches Produkt bietet das beste Preis-Leistungs-Verhältnis
5. Kaufberatung mit klaren Gründen

Formatieren Sie Ihre Antwort benutzerfreundlich mit Aufzählungspunkten und klaren Abschnitten. Fokus auf Kaufentscheidungshilfe."""
        
        user_prompt = f"Benutzer sucht nach: {query}\n\n{context}\n\nBitte umfassenden Produktvergleich und Empfehlung bereitstellen."
        
        try:
            # Generate LLM response with professional analysis
            answer = self.generate_answer(context, user_prompt, custom_system_prompt=system_prompt)
            
            # Build professional result with product cards
            result = "🛒 PROFESSIONELLE PRODUKTSUCHE (DEUTSCHER MARKT)\n" + "="*70 + "\n\n"
            result += answer + "\n\n"
            result += "📦 DIREKTE PRODUKTLINKS:\n" + "-"*70 + "\n\n"
            
            # Add formatted product cards
            for i, product in enumerate(products[:8], 1):
                result += f"╔══ PRODUKT {i} ══╗\n"
                result += f"║ {product.title[:65]}\n"
                result += f"║ \n"
                result += f"║ Händler: {product.merchant}\n"
                
                if product.price:
                    result += f"║ 💰 Preis: {product.price:.2f} {product.currency}\n"
                
                if product.rating:
                    stars = "⭐" * int(product.rating)
                    result += f"║ {stars} ({product.rating}/5.0)\n"
                
                result += f"║ \n"
                result += f"║ 🔗 Link: {product.url[:60]}\n"
                result += f"╚═══════════════╝\n\n"
            
            return result
            
        except Exception as e:
            print(f"LLM analysis error: {e}")
            # Fallback to structured listing
            result = "🛒 PRODUKTSUCHE (DEUTSCHER MARKT)\n" + "="*70 + "\n\n"
            result += f"Gefunden: {len(products)} Produkte für '{query}'\n\n"
            
            for i, product in enumerate(products[:8], 1):
                result += f"{i}. {product.title}\n"
                result += f"   Händler: {product.merchant}\n"
                if product.price:
                    result += f"   Preis: {product.price:.2f} {product.currency}\n"
                if product.rating:
                    result += f"   Bewertung: {product.rating}/5.0 ⭐\n"
                result += f"   🔗 {product.url}\n\n"
            
            return result
    
    def advanced_hybrid_search_filtered(self, question, doc_ids, n_results=15, alpha=0.7):
        """Enhanced hybrid search filtered by specific document IDs"""
        try:
            # Filter collection to only include specified IDs
            filtered_results = self.collection.get(ids=doc_ids)
            
            if not filtered_results['documents']:
                return []
            
            question_embedding = self.embedder.encode(question).tolist()
            
            # Create temporary collection for filtered search
            combined_docs = {}
            k = 60
            
            # Vector search on filtered documents
            for rank, (doc, metadata, doc_id) in enumerate(zip(
                filtered_results['documents'],
                filtered_results['metadatas'], 
                filtered_results['ids']
            )):
                doc_embedding = self.embedder.encode(doc).tolist()
                
                # Calculate similarity
                similarity = np.dot(question_embedding, doc_embedding) / (
                    np.linalg.norm(question_embedding) * np.linalg.norm(doc_embedding)
                )
                
                # Skip very dissimilar documents
                if similarity < 0.3:
                    continue
                
                doc_key = doc[:150]
                combined_docs[doc_key] = {
                    'doc': doc,
                    'metadata': metadata,
                    'score': alpha * (1 / (k + rank)),
                    'vector_similarity': similarity
                }
            
            # BM25 search on filtered corpus if available
            if self.bm25 and self.bm25_corpus:
                # Find indices of filtered documents in BM25 corpus
                filtered_indices = []
                for doc_id in doc_ids:
                    try:
                        all_ids = self.collection.get()['ids']
                        if doc_id in all_ids:
                            idx = all_ids.index(doc_id)
                            if idx < len(self.bm25_corpus):
                                filtered_indices.append(idx)
                    except:
                        continue
                
                if filtered_indices:
                    tokenized_query = question.lower().split()
                    bm25_scores = self.bm25.get_scores(tokenized_query)
                    
                    # Score only filtered documents
                    for rank, idx in enumerate(filtered_indices):
                        if idx < len(self.bm25_corpus):
                            doc = self.bm25_corpus[idx]
                            doc_key = doc[:150]
                            
                            if doc_key not in combined_docs:
                                all_docs = self.collection.get()
                                if idx < len(all_docs['metadatas']):
                                    metadata = all_docs['metadatas'][idx]
                                    combined_docs[doc_key] = {
                                        'doc': doc,
                                        'metadata': metadata,
                                        'score': 0
                                    }
                            
                            if doc_key in combined_docs:
                                combined_docs[doc_key]['score'] += (1 - alpha) * (bm25_scores[idx] / (1 + bm25_scores[idx]))
            
            sorted_docs = sorted(combined_docs.values(), key=lambda x: x['score'], reverse=True)
            return sorted_docs[:n_results]
            
        except Exception as e:
            # Fallback to standard search
            return self.advanced_hybrid_search(question, n_results)
    
    def run(self):
        """Start"""
        self.root.mainloop()


def check_license_and_run():
    """Check license/trial status before running the application"""
    # Initialize license manager
    license_manager = LicenseManager()
    
    # Get overall status
    can_run, is_trial, days_remaining, message = license_manager.get_status()
    
    if can_run:
        # Either licensed or valid trial - run the app
        trial_info = (is_trial, days_remaining)
        app = DocumentQAApp(trial_info=trial_info)
        app.run()
    else:
        # Trial expired or no license - show license dialog
        # Create a hidden root window for the dialog
        root = tk.Tk()
        root.withdraw()  # Hide the main window
        
        # Show license dialog
        dialog = LicenseDialog(root, license_manager, days_remaining=0, is_expired=True)
        result = dialog.show()
        
        if result:
            # License activated or user chose to continue (shouldn't happen if expired)
            root.destroy()
            can_run, is_trial, days_remaining, message = license_manager.get_status()
            if can_run:
                trial_info = (is_trial, days_remaining)
                app = DocumentQAApp(trial_info=trial_info)
                app.run()
        else:
            # User chose to exit
            root.destroy()
            sys.exit(0)


if __name__ == "__main__":
    check_license_and_run()