"""
Doqurix Web - ChatGPT-Style Interface
A modern web interface for document Q&A with conversation history
"""

from bottle import Bottle, request, response, static_file, template, abort
import os
import sys
import time
import threading
from pathlib import Path
import json
import base64
import io
import re
try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Initialize Bottle app
app = Bottle()

# Global variables
llm = None
embedder = None
reranker = None
collection = None
tax_collection = None  # Separate collection for tax agent
buerokratai_collection = None  # Separate collection for BürokratAI agent
chroma_client = None
bm25 = None
bm25_corpus = []
agents = None  # For insights agent

# ---------------------------------------------------------------------------
# Startup progress + lifecycle state (used by the headless web launcher)
# ---------------------------------------------------------------------------
STARTUP_STATE = {
    "ready": False,            # set True once all models are loaded
    "error": None,             # string if startup failed
    "title": "Starting Doqurix",
    "status": "Booting...",
    "stage_index": 0,
    "stage_total": 6,
    "stage_label": "",
    "percent": 0,
    "download_bytes": 0,
    "download_total": 0,
}

# Loading page shown while models are initializing. Polls /api/startup-progress
# and reloads the main UI once ready.
LOADING_HTML = """
<!DOCTYPE html>
<html lang=\"en\">
<head>
    <meta charset=\"UTF-8\">
    <title>Doqurix - Starting...</title>
    <style>
        body { margin:0; font-family:'Segoe UI',sans-serif;
               background:#0f0f0f; color:#ececf1;
               display:flex; align-items:center; justify-content:center;
               height:100vh; }
        .card { background:#1a1a1a; padding:36px 44px; border-radius:14px;
                box-shadow:0 10px 40px rgba(0,0,0,.5);
                width: 460px; max-width: 92vw; }
        h1 { margin:0 0 6px 0; font-size:22px; }
        .status { color:#b4b4b4; margin-bottom:18px; font-size:14px; }
        .bar { background:#2a2a2a; border-radius:8px; overflow:hidden;
               height:14px; }
        .fill { background:linear-gradient(135deg,#10a37f 0%,#1bc9a0 100%);
                height:100%; width:0%; transition:width .3s ease; }
        .meta { margin-top:10px; font-size:12px; color:#6e6e80;
                display:flex; justify-content:space-between; }
        .err { color:#ff6b6b; margin-top:14px; font-size:13px; }
    </style>
</head>
<body>
    <div class=\"card\">
        <h1 id=\"title\">Starting Doqurix</h1>
        <div class=\"status\" id=\"status\">Booting...</div>
        <div class=\"bar\"><div class=\"fill\" id=\"fill\"></div></div>
        <div class=\"meta\">
            <span id=\"stage\"></span>
            <span id=\"pct\">0%</span>
        </div>
        <div class=\"err\" id=\"err\"></div>
    </div>
    <script>
        function fmtMB(b){ return (b/1048576).toFixed(1) + ' MB'; }
        async function poll(){
            try {
                const r = await fetch('/api/startup-progress');
                const s = await r.json();
                document.getElementById('title').textContent = s.title || 'Starting...';
                document.getElementById('status').textContent = s.status || '';
                document.getElementById('fill').style.width = (s.percent || 0) + '%';
                document.getElementById('pct').textContent = Math.round(s.percent || 0) + '%';
                let stage = '';
                if (s.stage_total) stage = 'Step ' + s.stage_index + '/' + s.stage_total;
                if (s.download_total) {
                    stage = fmtMB(s.download_bytes) + ' / ' + fmtMB(s.download_total);
                }
                document.getElementById('stage').textContent = stage;
                if (s.error) {
                    document.getElementById('err').textContent = 'Error: ' + s.error;
                    return;
                }
                if (s.ready) { window.location.href = '/'; return; }
            } catch (e) {}
            setTimeout(poll, 700);
        }
        poll();
    </script>
</body>
</html>
"""

# Professional Modern UI - ChatGPT/Claude/Gemini Quality
INDEX_HTML = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Doqurix - AI Document Intelligence</title>
    <!-- Markdown / code-highlight / chart renderers. All CDN-hosted but the
         page still works offline without them; addMessage() falls back to a
         plain-text renderer if any library failed to load. -->
    <link rel="stylesheet"
          href="https://cdn.jsdelivr.net/gh/highlightjs/cdn-release@11.9.0/build/styles/github-dark.min.css">
    <script src="https://cdn.jsdelivr.net/npm/marked@12.0.2/marked.min.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/dompurify@3.0.11/dist/purify.min.js"></script>
    <script src="https://cdn.jsdelivr.net/gh/highlightjs/cdn-release@11.9.0/build/highlight.min.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.3/dist/chart.umd.min.js"></script>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
        
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        :root {
            --bg-primary: #0f0f0f;
            --bg-secondary: #1a1a1a;
            --bg-tertiary: #2a2a2a;
            --bg-chat-user: #2f2f2f;
            --bg-chat-assistant: #1a1a1a;
            --accent-primary: #10a37f;
            --accent-gradient: linear-gradient(135deg, #10a37f 0%, #1bc9a0 100%);
            --accent-orange: #ff6b35;
            --text-primary: #ececf1;
            --text-secondary: #b4b4b4;
            --text-muted: #6e6e80;
            --border-color: #3f3f46;
            --shadow-sm: 0 1px 3px rgba(0,0,0,0.3);
            --shadow-md: 0 4px 12px rgba(0,0,0,0.4);
            --shadow-lg: 0 10px 40px rgba(0,0,0,0.5);
            --shadow-glow: 0 0 20px rgba(16,163,127,0.3);
            --transition-fast: 0.15s ease;
            --transition-normal: 0.25s ease;
            --transition-slow: 0.4s ease;
        }
        
        body {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', sans-serif;
            background: var(--bg-primary);
            color: var(--text-primary);
            height: 100vh;
            overflow: hidden;
            position: relative;
        }
        
        /* Animated gradient background */
        body::before {
            content: '';
            position: fixed;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: radial-gradient(circle at 20% 50%, rgba(16,163,127,0.05) 0%, transparent 50%),
                        radial-gradient(circle at 80% 80%, rgba(255,107,53,0.03) 0%, transparent 50%);
            pointer-events: none;
            z-index: 0;
        }
        
        /* Header */
        .header {
            background: rgba(26, 26, 26, 0.95);
            backdrop-filter: blur(20px);
            padding: 14px 24px;
            border-bottom: 1px solid var(--border-color);
            display: flex;
            align-items: center;
            justify-content: space-between;
            position: relative;
            z-index: 100;
            box-shadow: var(--shadow-sm);
        }
        
        .header-left {
            display: flex;
            align-items: center;
            gap: 16px;
        }
        
        .logo {
            display: flex;
            align-items: center;
            gap: 10px;
            font-size: 18px;
            font-weight: 700;
            background: var(--accent-gradient);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
        }
        
        .logo-icon {
            font-size: 24px;
            filter: drop-shadow(0 0 8px rgba(16,163,127,0.5));
        }
        
        .upload-btn {
            background: var(--accent-gradient);
            color: white;
            border: none;
            padding: 10px 20px;
            border-radius: 10px;
            cursor: pointer;
            font-size: 14px;
            font-weight: 600;
            transition: all var(--transition-normal);
            box-shadow: var(--shadow-md);
            display: flex;
            align-items: center;
            gap: 8px;
        }
        
        .upload-btn:hover {
            transform: translateY(-2px) scale(1.02);
            box-shadow: var(--shadow-glow), var(--shadow-lg);
        }
        
        .upload-btn:active {
            transform: translateY(0) scale(0.98);
        }

        .settings-btn {
            background: rgba(255,255,255,0.06);
            color: var(--text-primary, #ececf1);
            border: 1px solid rgba(255,255,255,0.12);
            padding: 10px 16px;
            border-radius: 10px;
            cursor: pointer;
            font-size: 14px;
            font-weight: 600;
            transition: all var(--transition-normal);
            display: flex;
            align-items: center;
            gap: 8px;
            margin-left: 10px;
        }
        .settings-btn:hover {
            background: rgba(255,255,255,0.12);
            transform: translateY(-1px);
        }
        .settings-btn:active { transform: translateY(0); }
        
        /* Sidebar */
        .sidebar {
            position: fixed;
            left: 0;
            top: 63px;
            bottom: 0;
            width: 280px;
            background: rgba(26, 26, 26, 0.8);
            backdrop-filter: blur(20px);
            border-right: 1px solid var(--border-color);
            padding: 16px;
            overflow-y: auto;
            z-index: 50;
            transition: transform var(--transition-normal);
        }
        
        .sidebar-header {
            display: flex;
            align-items: center;
            justify-content: space-between;
            margin-bottom: 16px;
        }
        
        .sidebar h3 {
            font-size: 13px;
            color: var(--text-muted);
            text-transform: uppercase;
            font-weight: 600;
            letter-spacing: 0.5px;
        }
        
        .doc-count {
            background: var(--bg-tertiary);
            color: var(--text-secondary);
            padding: 2px 8px;
            border-radius: 12px;
            font-size: 11px;
            font-weight: 600;
        }
        
        .doc-item {
            background: var(--bg-tertiary);
            padding: 12px;
            margin-bottom: 8px;
            border-radius: 10px;
            font-size: 13px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            transition: all var(--transition-fast);
            border: 1px solid transparent;
            animation: slideIn 0.3s ease;
        }
        
        @keyframes slideIn {
            from {
                opacity: 0;
                transform: translateX(-10px);
            }
            to {
                opacity: 1;
                transform: translateX(0);
            }
        }
        
        .doc-item:hover {
            background: var(--bg-chat-user);
            border-color: var(--accent-primary);
            transform: translateX(4px);
            box-shadow: var(--shadow-md);
        }
        
        .doc-name {
            display: flex;
            align-items: center;
            gap: 8px;
            flex: 1;
            overflow: hidden;
        }
        
        .doc-icon {
            font-size: 16px;
            opacity: 0.9;
        }
        
        .doc-text {
            overflow: hidden;
            text-overflow: ellipsis;
            white-space: nowrap;
            color: var(--text-secondary);
        }
        
        .delete-btn {
            background: rgba(239, 68, 68, 0.1);
            color: #ef4444;
            border: 1px solid rgba(239, 68, 68, 0.3);
            padding: 6px 10px;
            border-radius: 6px;
            cursor: pointer;
            font-size: 11px;
            font-weight: 600;
            transition: all var(--transition-fast);
            opacity: 0;
        }
        
        .doc-item:hover .delete-btn {
            opacity: 1;
        }
        
        .delete-btn:hover {
            background: rgba(239, 68, 68, 0.2);
            border-color: #ef4444;
            transform: scale(1.05);
        }
        
        /* Main chat area */
        .chat-container {
            margin-left: 280px;
            display: flex;
            flex-direction: column;
            height: calc(100vh - 63px);
            position: relative;
            z-index: 1;
        }
        
        .messages {
            flex: 1;
            overflow-y: auto;
            padding: 32px 24px;
            scroll-behavior: smooth;
        }
        
        .message {
            margin-bottom: 32px;
            opacity: 0;
            animation: fadeInUp 0.4s ease forwards;
        }
        
        @keyframes fadeInUp {
            from {
                opacity: 0;
                transform: translateY(20px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }
        
        .message.user {
            max-width: 900px;
            margin-left: auto;
            margin-right: auto;
        }
        
        .message.assistant {
            max-width: 900px;
            margin-left: auto;
            margin-right: auto;
        }
        
        .message-wrapper {
            display: flex;
            gap: 16px;
            align-items: flex-start;
        }
        
        .message.user .message-wrapper {
            flex-direction: row-reverse;
        }
        
        .avatar {
            width: 36px;
            height: 36px;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 18px;
            flex-shrink: 0;
            box-shadow: var(--shadow-md);
        }
        
        .avatar.user {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        }
        
        .avatar.assistant {
            background: var(--accent-gradient);
        }
        
        .message-content {
            flex: 1;
            background: var(--bg-chat-assistant);
            padding: 16px 20px;
            border-radius: 16px;
            line-height: 1.7;
            font-size: 15px;
            box-shadow: var(--shadow-sm);
            border: 1px solid var(--border-color);
            user-select: text;
            cursor: default;
        }
        
        .message.user .message-content {
            background: var(--bg-chat-user);
            border-color: rgba(102, 126, 234, 0.2);
        }

        /* -------- Markdown rendering inside assistant bubbles -------- */
        .message-content h1,
        .message-content h2,
        .message-content h3,
        .message-content h4 {
            margin: 18px 0 10px;
            font-weight: 600;
            line-height: 1.3;
            color: var(--text-primary);
        }
        .message-content h1 { font-size: 22px; }
        .message-content h2 { font-size: 19px; }
        .message-content h3 { font-size: 17px; }
        .message-content h4 { font-size: 15px; color: var(--text-secondary); }
        .message-content p { margin: 8px 0; }
        .message-content ul,
        .message-content ol { margin: 8px 0 8px 22px; padding: 0; }
        .message-content li { margin: 4px 0; }
        .message-content strong { color: var(--text-primary); font-weight: 600; }
        .message-content em { color: var(--text-secondary); }
        .message-content blockquote {
            border-left: 3px solid var(--accent-primary);
            margin: 10px 0;
            padding: 4px 14px;
            color: var(--text-secondary);
            background: rgba(16,163,127,0.06);
            border-radius: 0 8px 8px 0;
        }
        .message-content a {
            color: var(--accent-primary);
            text-decoration: none;
            border-bottom: 1px dotted rgba(16,163,127,0.4);
        }
        .message-content a:hover { border-bottom-style: solid; }
        .message-content code {
            font-family: 'JetBrains Mono', 'Consolas', monospace;
            font-size: 13px;
            background: rgba(255,255,255,0.06);
            padding: 2px 6px;
            border-radius: 4px;
            border: 1px solid var(--border-color);
        }
        .message-content pre {
            background: #0d0d0d;
            border: 1px solid var(--border-color);
            border-radius: 10px;
            padding: 14px 16px;
            overflow-x: auto;
            margin: 10px 0;
            font-size: 13px;
            line-height: 1.55;
        }
        .message-content pre code {
            background: transparent;
            border: 0;
            padding: 0;
            font-size: inherit;
        }
        .message-content table {
            width: 100%;
            border-collapse: collapse;
            margin: 12px 0;
            font-size: 14px;
            background: rgba(255,255,255,0.02);
            border-radius: 10px;
            overflow: hidden;
            border: 1px solid var(--border-color);
        }
        .message-content thead { background: rgba(16,163,127,0.12); }
        .message-content th,
        .message-content td {
            padding: 10px 14px;
            text-align: left;
            border-bottom: 1px solid var(--border-color);
        }
        .message-content th { font-weight: 600; color: var(--text-primary); }
        .message-content tr:last-child td { border-bottom: 0; }
        .message-content tr:hover { background: rgba(255,255,255,0.03); }
        .message-content hr {
            border: 0;
            border-top: 1px solid var(--border-color);
            margin: 16px 0;
        }
        .message-content details {
            margin: 8px 0;
            border: 1px solid var(--border-color);
            border-radius: 8px;
            padding: 6px 12px;
            background: rgba(255,255,255,0.02);
        }
        .message-content summary {
            cursor: pointer;
            color: var(--accent-primary);
            font-weight: 500;
            user-select: none;
            outline: none;
        }
        .message-content details[open] summary { margin-bottom: 8px; }
        .message-content .chart-wrap {
            background: #0d0d0d;
            border: 1px solid var(--border-color);
            border-radius: 12px;
            padding: 14px;
            margin: 12px 0;
        }
        .message-content .chart-wrap canvas { max-width: 100%; }
        .message-content .chart-title {
            font-size: 13px;
            color: var(--text-secondary);
            text-align: center;
            margin-bottom: 8px;
        }
        
        .loading-dots {
            display: inline-flex;
            gap: 4px;
            align-items: center;
        }
        
        .loading-dots span {
            width: 8px;
            height: 8px;
            background: var(--accent-primary);
            border-radius: 50%;
            animation: bounce 1.4s infinite ease-in-out both;
        }
        
        .loading-dots span:nth-child(1) { animation-delay: -0.32s; }
        .loading-dots span:nth-child(2) { animation-delay: -0.16s; }
        
        @keyframes bounce {
            0%, 80%, 100% { transform: scale(0.8); opacity: 0.5; }
            40% { transform: scale(1.2); opacity: 1; }
        }
        
        .source {
            cursor: pointer;
            color: var(--accent-primary);
            text-decoration: none;
            padding: 10px 14px;
            background: rgba(16, 163, 127, 0.08);
            border: 1px solid rgba(16, 163, 127, 0.2);
            border-radius: 8px;
            margin: 10px 0;
            display: block;
            transition: all var(--transition-fast);
            font-size: 13px;
        }
        
        .source:hover {
            background: rgba(16, 163, 127, 0.15);
            border-color: var(--accent-primary);
            transform: translateX(4px);
        }
        
        .source-content {
            display: none;
            margin-top: 10px;
            padding: 14px;
            background: var(--bg-tertiary);
            border-radius: 8px;
            color: var(--text-secondary);
            font-size: 13px;
            line-height: 1.6;
            max-height: 300px;
            overflow-y: auto;
            border-left: 3px solid var(--accent-primary);
        }
        
        .source-content.expanded {
            display: block;
            animation: expandDown 0.3s ease;
        }
        
        @keyframes expandDown {
            from {
                opacity: 0;
                max-height: 0;
            }
            to {
                opacity: 1;
                max-height: 300px;
            }
        }
        
        /* Welcome screen */
        .welcome {
            text-align: center;
            padding: 60px 24px;
            max-width: 600px;
            margin: 0 auto;
        }
        
        .welcome h2 {
            font-size: 32px;
            font-weight: 700;
            margin-bottom: 16px;
            background: var(--accent-gradient);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
        }
        
        .welcome p {
            color: var(--text-secondary);
            font-size: 16px;
            line-height: 1.6;
        }
        
        /* Input area */
        .input-area {
            background: rgba(26, 26, 26, 0.95);
            backdrop-filter: blur(20px);
            padding: 24px;
            border-top: 1px solid var(--border-color);
            box-shadow: 0 -4px 20px rgba(0,0,0,0.3);
        }
        
        .input-wrapper {
            max-width: 900px;
            margin: 0 auto;
            position: relative;
        }
        
        .input-container {
            background: var(--bg-tertiary);
            border: 2px solid var(--border-color);
            border-radius: 16px;
            display: flex;
            align-items: center;
            gap: 8px;
            padding: 8px;
            transition: all var(--transition-fast);
            box-shadow: var(--shadow-md);
        }
        
        .input-container:focus-within {
            border-color: var(--accent-primary);
            box-shadow: var(--shadow-glow), var(--shadow-md);
        }
        
        .attach-btn {
            background: transparent;
            border: none;
            width: 40px;
            height: 40px;
            border-radius: 10px;
            cursor: pointer;
            display: flex;
            align-items: center;
            justify-content: center;
            transition: all var(--transition-fast);
            flex-shrink: 0;
        }
        
        .attach-btn:hover {
            background: var(--bg-chat-user);
            transform: scale(1.1) rotate(10deg);
        }
        
        .agent-btn {
            background: transparent;
            border: none;
            padding: 6px 12px;
            height: 40px;
            border-radius: 10px;
            cursor: pointer;
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 4px;
            transition: all var(--transition-fast);
            flex-shrink: 0;
            position: relative;
            z-index: 10;
        }
        
        .agent-btn:hover {
            background: var(--bg-chat-user);
        }
        
        .agent-menu {
            position: fixed;
            bottom: 80px;
            left: 50px;
            background: var(--bg-tertiary);
            border: 1px solid var(--border-color);
            border-radius: 12px;
            box-shadow: var(--shadow-lg);
            z-index: 9999;
            min-width: 220px;
            overflow: hidden;
            animation: slideUp 0.2s ease-out;
        }
        
        @keyframes slideUp {
            from {
                opacity: 0;
                transform: translateY(10px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }
        
        .agent-option {
            padding: 12px 16px;
            cursor: pointer;
            transition: background var(--transition-fast);
            color: var(--text-primary);
            display: flex;
            align-items: center;
            gap: 8px;
            border-bottom: 1px solid var(--border-color);
            white-space: nowrap;
        }
        
        .agent-option:last-child {
            border-bottom: none;
        }
        
        .agent-option:hover {
            background: var(--bg-chat-user);
        }
        
        .agent-desc {
            margin-left: auto;
            font-size: 11px;
            color: var(--text-muted);
            flex-shrink: 0;
        }
        
        textarea {
            flex: 1;
            background: transparent;
            border: none;
            padding: 12px;
            color: var(--text-primary);
            font-size: 15px;
            font-family: inherit;
            resize: none;
            min-height: 48px;
            max-height: 200px;
            line-height: 1.4;
            vertical-align: top;
            box-sizing: border-box;
            align-self: center;
        }
        
        textarea:focus {
            outline: none;
        }
        
        textarea::placeholder {
            color: var(--text-muted);
        }
        
        .send-btn {
            background: var(--accent-gradient);
            border: none;
            width: 40px;
            height: 40px;
            border-radius: 10px;
            cursor: pointer;
            display: flex;
            align-items: center;
            justify-content: center;
            transition: all var(--transition-fast);
            flex-shrink: 0;
            box-shadow: var(--shadow-sm);
        }
        
        .send-btn:hover:not(:disabled) {
            transform: scale(1.1) rotate(-5deg);
            box-shadow: var(--shadow-glow);
        }
        
        .send-btn:active:not(:disabled) {
            transform: scale(0.95);
        }
        
        .send-btn:disabled {
            opacity: 0.4;
            cursor: not-allowed;
        }
        
        /* Scrollbar */
        ::-webkit-scrollbar {
            width: 10px;
        }
        
        ::-webkit-scrollbar-track {
            background: var(--bg-secondary);
        }
        
        ::-webkit-scrollbar-thumb {
            background: var(--bg-tertiary);
            border-radius: 5px;
            border: 2px solid var(--bg-secondary);
        }
        
        ::-webkit-scrollbar-thumb:hover {
            background: var(--border-color);
        }
        
        /* Drag and drop */
        .drag-overlay {
            position: fixed;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: rgba(16, 163, 127, 0.1);
            border: 4px dashed var(--accent-primary);
            display: none;
            align-items: center;
            justify-content: center;
            z-index: 1000;
            backdrop-filter: blur(10px);
        }
        
        .drag-overlay.active {
            display: flex;
        }
        
        .drag-content {
            text-align: center;
            color: var(--accent-primary);
        }
        
        .drag-content svg {
            width: 64px;
            height: 64px;
            margin-bottom: 16px;
            animation: float 2s ease-in-out infinite;
        }
        
        @keyframes float {
            0%, 100% { transform: translateY(0); }
            50% { transform: translateY(-10px); }
        }
        
        .drag-content h3 {
            font-size: 24px;
            margin-bottom: 8px;
        }
        
        /* Responsive */
        @media (max-width: 1200px) {
            .welcome {
                max-width: 500px;
                padding: 40px 20px;
            }
            
            .welcome h2 {
                font-size: 28px;
            }
            
            .welcome p {
                font-size: 15px;
            }
            
            .messages {
                padding: 24px 20px;
            }
        }
        
        @media (max-width: 768px) {
            .sidebar {
                transform: translateX(-100%);
                width: 260px;
            }
            
            .sidebar.open {
                transform: translateX(0);
            }
            
            .chat-container {
                margin-left: 0;
            }
            
            .welcome {
                max-width: 90%;
                padding: 30px 16px;
            }
            
            .welcome h2 {
                font-size: 24px;
                margin-bottom: 12px;
            }
            
            .welcome p {
                font-size: 14px;
                line-height: 1.5;
            }
            
            .messages {
                padding: 20px 16px;
            }
            
            .message {
                margin-bottom: 20px;
            }
            
            .avatar {
                width: 30px;
                height: 30px;
                font-size: 16px;
            }
            
            .message-content {
                padding: 12px 16px;
                font-size: 14px;
            }
            
            .input-area {
                padding: 16px;
            }
            
            .header {
                padding: 12px 16px;
            }
            
            .logo {
                font-size: 16px;
            }
            
            .upload-btn {
                padding: 8px 16px;
                font-size: 13px;
            }
        }
        
        @media (max-width: 480px) {
            .welcome {
                padding: 20px 12px;
            }
            
            .welcome h2 {
                font-size: 20px;
                margin-bottom: 10px;
            }
            
            .welcome p {
                font-size: 13px;
                line-height: 1.4;
            }
            
            .messages {
                padding: 16px 12px;
            }
            
            .input-area {
                padding: 12px;
            }
            
            .header {
                padding: 10px 12px;
            }
            
            .logo {
                font-size: 14px;
                gap: 6px;
            }
            
            .logo-icon {
                font-size: 18px;
            }
            
            .upload-btn {
                padding: 6px 12px;
                font-size: 12px;
                gap: 6px;
            }
        }
    </style>
</head>
<body>
    <!-- Drag and Drop Overlay -->
    <div class="drag-overlay" id="dragOverlay">
        <div class="drag-content">
            <svg viewBox="0 0 24 24" fill="currentColor">
                <path d="M19.35 10.04C18.67 6.59 15.64 4 12 4 9.11 4 6.6 5.64 5.35 8.04 2.34 8.36 0 10.91 0 14c0 3.31 2.69 6 6 6h13c2.76 0 5-2.24 5-5 0-2.64-2.05-4.78-4.65-4.96zM14 13v4h-4v-4H7l5-5 5 5h-3z"/>
            </svg>
            <h3>Drop PDF file here</h3>
            <p>Release to upload</p>
        </div>
    </div>

    <!-- Header -->
    <div class="header">
        <div class="header-left">
            <div class="logo">
                <span class="logo-icon">◆</span>
                <span>DOQURIX</span>
            </div>
        </div>
        <input type="file" id="file-input" accept=".pdf,.docx" multiple style="display: none;" onchange="uploadFile()">
        <button class="upload-btn" onclick="document.getElementById('file-input').click()">
            <svg width="18" height="18" viewBox="0 0 24 24" fill="currentColor">
                <path d="M9 16h6v-6h4l-7-7-7 7h4zm-4 2h14v2H5z"/>
            </svg>
            Upload Documents
        </button>
        <button class="settings-btn" type="button" title="Settings & License"
                onclick="window.doqurixOpenSettings && window.doqurixOpenSettings()">
            <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
                <circle cx="12" cy="12" r="3"/>
                <path d="M19.4 15a1.65 1.65 0 0 0 .33 1.82l.06.06a2 2 0 1 1-2.83 2.83l-.06-.06a1.65 1.65 0 0 0-1.82-.33 1.65 1.65 0 0 0-1 1.51V21a2 2 0 0 1-4 0v-.09A1.65 1.65 0 0 0 9 19.4a1.65 1.65 0 0 0-1.82.33l-.06.06a2 2 0 1 1-2.83-2.83l.06-.06a1.65 1.65 0 0 0 .33-1.82 1.65 1.65 0 0 0-1.51-1H3a2 2 0 0 1 0-4h.09A1.65 1.65 0 0 0 4.6 9a1.65 1.65 0 0 0-.33-1.82l-.06-.06a2 2 0 1 1 2.83-2.83l.06.06a1.65 1.65 0 0 0 1.82.33H9a1.65 1.65 0 0 0 1-1.51V3a2 2 0 0 1 4 0v.09a1.65 1.65 0 0 0 1 1.51 1.65 1.65 0 0 0 1.82-.33l.06-.06a2 2 0 1 1 2.83 2.83l-.06.06a1.65 1.65 0 0 0-.33 1.82V9a1.65 1.65 0 0 0 1.51 1H21a2 2 0 0 1 0 4h-.09a1.65 1.65 0 0 0-1.51 1z"/>
            </svg>
            Settings
        </button>
    </div>
    
    <!-- Sidebar -->
    <div class="sidebar" id="sidebar">
        <div class="sidebar-header">
            <h3>Documents</h3>
            <span class="doc-count" id="docCount">0</span>
        </div>
        <div id="documents"></div>
    </div>
    
    <!-- Chat Container -->
    <div class="chat-container">
        <div class="messages" id="messages">
            <div class="welcome">
                <h2>Welcome to Doqurix</h2>
                <p>Your AI-powered document assistant. Upload PDF documents and start asking intelligent questions to extract insights instantly.</p>
            </div>
        </div>
        
        <!-- Input Area -->
        <div class="input-area">
            <div class="input-wrapper">
                <input type="file" id="input-file" accept=".pdf,.docx" multiple style="display: none;" onchange="uploadFileFromInput()">
                <div class="input-container">
                    <button class="attach-btn" onclick="document.getElementById('input-file').click()" title="Attach file">
                        <svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="#8e8ea0" stroke-width="2">
                            <path d="M21.44 11.05l-9.19 9.19a6 6 0 0 1-8.49-8.49l9.19-9.19a4 4 0 0 1 5.66 5.66l-9.2 9.19a2 2 0 0 1-2.83-2.83l8.49-8.48"/>
                        </svg>
                    </button>
                    <button class="agent-btn" onclick="toggleAgentMenu(event)" title="Select Agent" id="agentBtn">
                        <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="#8e8ea0" stroke-width="2">
                            <path d="M12 2L2 7l10 5 10-5-10-5zM2 17l10 5 10-5M2 12l10 5 10-5"/>
                        </svg>
                        <span id="agentLabel" style="font-size: 12px; margin-left: 4px; color: #8e8ea0;">None</span>
                    </button>
                    <div class="agent-menu" id="agentMenu" style="display: none;">
                        <div class="agent-option" onclick="selectAgent('none')">📄 None<span class="agent-desc">User Documents</span></div>
                        <div class="agent-option" onclick="selectAgent('insights')">💡 Insights<span class="agent-desc">Auto-Analysis</span></div>
                        <div class="agent-option" onclick="selectAgent('tax_germany')">🇩🇪 Tax Germany<span class="agent-desc">German Tax Expert</span></div>
                        <div class="agent-option" onclick="selectAgent('ecommerce_germany')">🛒 E-Commerce<span class="agent-desc">Product Search & Compare</span></div>
                        <div class="agent-option" onclick="selectAgent('buerokratai_germany')">🏛️ BürokratAI<span class="agent-desc">Immigration Helper</span></div>
                    </div>
                    <textarea id="question" placeholder="Ask me anything about your documents..." rows="1" 
                              oninput="auto_grow(this)" onkeydown="handleKeyPress(event)"></textarea>
                    <button class="send-btn" id="sendBtn" onclick="askQuestion()" disabled title="Send message">
                        <svg width="20" height="20" viewBox="0 0 24 24" fill="white">
                            <path d="M2.01 21L23 12 2.01 3 2 10l15 2-15 2z"/>
                        </svg>
                    </button>
                </div>
            </div>
        </div>
    </div>
    
    <script>
        let documents = [];
        let dragCounter = 0;
        let currentAgent = 'none';  // Track current agent selection
        
        // Agent selection functions
        function toggleAgentMenu(event) {
            if (event) event.stopPropagation();
            const menu = document.getElementById('agentMenu');
            menu.style.display = menu.style.display === 'none' ? 'block' : 'none';
        }
        
        function selectAgent(agent) {
            currentAgent = agent;
            const label = document.getElementById('agentLabel');
            const agentNames = {
                'none': 'None',
                'insights': 'Insights',
                'tax_germany': 'Tax Germany',
                'ecommerce_germany': 'E-Commerce',
                'buerokratai_germany': 'BürokratAI'
            };
            label.textContent = agentNames[agent] || 'None';
            document.getElementById('agentMenu').style.display = 'none';
            
            // Update placeholder based on agent without moving focus
            const textarea = document.getElementById('question');
            const currentFocus = document.activeElement;
            if (agent === 'tax_germany') {
                textarea.placeholder = 'Ask me about German taxes...';
            } else if (agent === 'ecommerce_germany') {
                textarea.placeholder = 'What products are you looking for?';
            } else if (agent === 'buerokratai_germany') {
                textarea.placeholder = 'Ask about visas, Anmeldung, insurance, work permits...';
            } else if (agent === 'insights') {
                textarea.placeholder = 'Upload documents for automatic insights...';
            } else {
                textarea.placeholder = 'Ask me anything about your documents...';
            }
            
            // Restore focus if textarea had it
            if (currentFocus === textarea) {
                textarea.focus();
            }
        }
        
        // Close agent menu when clicking outside
        document.addEventListener('click', function(event) {
            const agentBtn = document.getElementById('agentBtn');
            const agentMenu = document.getElementById('agentMenu');
            if (agentBtn && agentMenu && !agentBtn.contains(event.target) && !agentMenu.contains(event.target)) {
                agentMenu.style.display = 'none';
            }
        });
        
        // Auto-grow textarea
        function auto_grow(element) {
            element.style.height = "48px";
            element.style.height = (element.scrollHeight) + "px";
            
            // Enable/disable send button
            const sendBtn = document.getElementById('sendBtn');
            sendBtn.disabled = !element.value.trim();
        }
        
        // Handle keyboard shortcuts
        function handleKeyPress(e) {
            if (e.key === 'Enter' && !e.shiftKey) {
                e.preventDefault();
                askQuestion();
            }
        }
        
        // Drag and drop handlers
        document.body.addEventListener('dragenter', (e) => {
            e.preventDefault();
            dragCounter++;
            if (e.dataTransfer.types.includes('Files')) {
                document.getElementById('dragOverlay').classList.add('active');
            }
        });
        
        document.body.addEventListener('dragleave', (e) => {
            e.preventDefault();
            dragCounter--;
            if (dragCounter === 0) {
                document.getElementById('dragOverlay').classList.remove('active');
            }
        });
        
        document.body.addEventListener('dragover', (e) => {
            e.preventDefault();
        });
        
        document.body.addEventListener('drop', (e) => {
            e.preventDefault();
            dragCounter = 0;
            document.getElementById('dragOverlay').classList.remove('active');
            
            const files = e.dataTransfer.files;
            if (files.length > 0 && files[0].type === 'application/pdf') {
                processUpload(files[0]);
            }
        });
        
        // File upload from button
        async function uploadFile() {
            const fileInput = document.getElementById('file-input');
            const files = fileInput.files;
            if (!files || files.length === 0) return;
            
            for (let file of files) {
                await processUpload(file);
            }
            fileInput.value = '';
        }
        
        async function uploadFileFromInput() {
            const fileInput = document.getElementById('input-file');
            const files = fileInput.files;
            if (!files || files.length === 0) return;
            
            for (let file of files) {
                await processUpload(file);
            }
            fileInput.value = '';
        }
        
        // Process file upload
        async function processUpload(file) {
            const formData = new FormData();
            formData.append('file', file);
            
            const loadingId = addMessage(
                `<div style="display: flex; align-items: center; gap: 12px;">
                    <div class="loading-dots">
                        <span></span><span></span><span></span>
                    </div>
                    <span>Uploading ${file.name}...</span>
                </div>`, 
                'assistant', 
                true
            );
            
            try {
                const response = await fetch('/upload', {
                    method: 'POST',
                    body: formData
                });
                
                const result = await response.json();
                
                // Remove loading message
                const loadingMsg = document.getElementById(loadingId);
                if (loadingMsg) loadingMsg.remove();
                
                if (result.success) {
                    documents.push(result.filename);
                    updateDocumentList();
                    addMessage(`✅ Successfully uploaded <strong>${file.name}</strong>`, 'assistant');
                } else {
                    addMessage(`❌ Error uploading file: ${result.error}`, 'assistant');
                }
            } catch (error) {
                const loadingMsg = document.getElementById(loadingId);
                if (loadingMsg) loadingMsg.remove();
                addMessage(`❌ Upload failed: ${error.message}`, 'assistant');
            }
        }
        
        // Update document list in sidebar
        function updateDocumentList() {
            const docList = document.getElementById('documents');
            const docCount = document.getElementById('docCount');
            
            docCount.textContent = documents.length;
            if (documents.length === 0) {
                docList.innerHTML = '<div style="color: var(--text-muted); font-size: 13px; padding: 16px; text-align: center;">No documents uploaded</div>';
                return;
            }
            
            docList.innerHTML = documents.map((doc, idx) => `
                <div class="doc-item">
                    <div class="doc-name">
                        <span class="doc-icon">📄</span>
                        <span class="doc-text" title="${doc}">${doc}</span>
                    </div>
                    <button class="delete-btn" onclick="deleteDocument(${idx})" title="Delete document">
                        Delete
                    </button>
                </div>
            `).join('');
        }
        
        // Delete document
        async function deleteDocument(index) {
            const filename = documents[index];
            
            if (!confirm(`Delete "${filename}"?`)) return;
            
            try {
                const response = await fetch('/delete', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({filename: filename})
                });
                
                const result = await response.json();
                
                if (result.success) {
                    documents.splice(index, 1);
                    updateDocumentList();
                    addMessage(`🗑️ Deleted <strong>${filename}</strong>`, 'assistant');
                } else {
                    alert('Error: ' + result.error);
                }
            } catch (error) {
                alert('Delete failed: ' + error.message);
            }
        }
        
        // Ask question
        async function askQuestion() {
            const questionInput = document.getElementById('question');
            const question = questionInput.value.trim();
            
            if (!question) return;
            
            // Tax agent, E-commerce agent, and BürokratAI don't require documents
            if (currentAgent !== 'tax_germany' && currentAgent !== 'ecommerce_germany' && currentAgent !== 'buerokratai_germany' && documents.length === 0) {
                alert('⚠️ Please upload documents first');
                return;
            }
            
            // Add user message
            addMessage(question, 'user');
            questionInput.value = '';
            questionInput.style.height = "24px";
            document.getElementById('sendBtn').disabled = true;
            
            // Add loading message
            const loadingId = addMessage(
                `<div class="loading-dots">
                    <span></span><span></span><span></span>
                </div>`, 
                'assistant', 
                true
            );
            
            try {
                const response = await fetch('/query', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({
                        question: question,
                        agent: currentAgent  // Send selected agent
                    })
                });
                
                const result = await response.json();
                
                // Remove loading message
                const loadingMsg = document.getElementById(loadingId);
                if (loadingMsg) loadingMsg.remove();
                
                if (result.success) {
                    addMessage(result.answer, 'assistant');
                } else {
                    addMessage(`❌ Error: ${result.error}`, 'assistant');
                }
            } catch (error) {
                const loadingMsg = document.getElementById(loadingId);
                if (loadingMsg) loadingMsg.remove();
                addMessage(`❌ Request failed: ${error.message}`, 'assistant');
            }
        }
        
        // Add message to chat
        function addMessage(content, role, isLoading = false) {
            const messagesDiv = document.getElementById('messages');
            const messageId = 'msg-' + Date.now() + '-' + Math.random();
            
            // Remove welcome screen if exists
            const welcome = messagesDiv.querySelector('.welcome');
            if (welcome) welcome.remove();
            
            const avatar = role === 'user' ? '👤' : '🤖';

            // Render assistant text as GitHub-flavored Markdown.
            // User messages and loading indicators are passed through
            // unchanged so existing HTML snippets (e.g. spinner markup)
            // keep working.
            let rendered = content;
            if (role === 'assistant' && !isLoading && window.marked && window.DOMPurify) {
                try {
                    marked.setOptions({
                        gfm: true, breaks: true, headerIds: false, mangle: false
                    });
                    const html = marked.parse(String(content));
                    rendered = DOMPurify.sanitize(html, {
                        ADD_TAGS: ['details', 'summary'],
                        ADD_ATTR: ['open']
                    });
                } catch (e) {
                    console.error('markdown render failed', e);
                }
            }

            const messageDiv = document.createElement('div');
            messageDiv.id = messageId;
            messageDiv.className = `message ${role}`;
            messageDiv.innerHTML = `
                <div class="message-wrapper">
                    <div class="avatar ${role}">${avatar}</div>
                    <div class="message-content">
                        ${rendered}
                    </div>
                </div>
            `;
            
            messagesDiv.appendChild(messageDiv);

            // Post-process: syntax-highlight code blocks and render any
            // ```chart``` JSON blocks as actual Chart.js charts.
            if (role === 'assistant' && !isLoading) {
                postProcessAssistantMessage(messageDiv);
            }

            messagesDiv.scrollTop = messagesDiv.scrollHeight;
            
            return messageId;
        }

        // Replace ```chart``` fenced blocks (rendered by marked as
        // <pre><code class="language-chart">...</code></pre>) with real
        // Chart.js canvases, and highlight all remaining code blocks.
        function postProcessAssistantMessage(root) {
            // 1. Charts
            root.querySelectorAll('pre > code.language-chart').forEach((codeEl) => {
                let spec = null;
                try { spec = JSON.parse(codeEl.textContent.trim()); }
                catch (e) { return; }
                if (!spec || !spec.type || !window.Chart) return;
                const wrap = document.createElement('div');
                wrap.className = 'chart-wrap';
                if (spec.title) {
                    const t = document.createElement('div');
                    t.className = 'chart-title';
                    t.textContent = spec.title;
                    wrap.appendChild(t);
                }
                const canvas = document.createElement('canvas');
                canvas.height = 240;
                wrap.appendChild(canvas);
                codeEl.parentElement.replaceWith(wrap);
                try {
                    new Chart(canvas, {
                        type: spec.type,
                        data: { labels: spec.labels || [], datasets: spec.datasets || [] },
                        options: {
                            responsive: true,
                            plugins: {
                                legend: { labels: { color: '#ececf1' } },
                                tooltip: { backgroundColor: '#1a1a1a' }
                            },
                            scales: (spec.type === 'pie' || spec.type === 'doughnut') ? {} : {
                                x: { ticks: { color: '#b4b4b4' }, grid: { color: '#2a2a2a' } },
                                y: { ticks: { color: '#b4b4b4' }, grid: { color: '#2a2a2a' } }
                            }
                        }
                    });
                } catch (e) {
                    wrap.textContent = 'Chart render failed: ' + e.message;
                }
            });
            // 2. Code highlighting
            if (window.hljs) {
                root.querySelectorAll('pre code').forEach((el) => {
                    try { hljs.highlightElement(el); } catch (e) {}
                });
            }
        }
        
        // Toggle source expansion
        function toggleSource(element) {
            const content = element.querySelector('.source-content');
            content.classList.toggle('expanded');
        }
        
        // Initialize
        // Hydrate the sidebar from the server so previously uploaded
        // documents reappear after a refresh or a process restart.
        (async () => {
            try {
                const r = await fetch('/documents');
                const j = await r.json();
                if (j && Array.isArray(j.documents)) {
                    documents = j.documents;
                }
            } catch (e) { console.warn('document hydrate failed', e); }
            updateDocumentList();
        })();
    </script>
</body>
</html>
"""


@app.route('/')
def index():
    """Serve main page (or loading page while models initialize)."""
    if not STARTUP_STATE.get("ready"):
        response.content_type = 'text/html; charset=utf-8'
        return LOADING_HTML
    return _INDEX_WITH_STOP_BUTTON


@app.route('/api/startup-progress')
def startup_progress():
    """Return current startup progress for the loading overlay."""
    response.content_type = 'application/json'
    return json.dumps(STARTUP_STATE)


@app.route('/api/shutdown', method=['POST', 'GET'])
def shutdown():
    """Gracefully stop the desktop web app from the browser."""
    response.content_type = 'application/json'

    def _exit_soon():
        time.sleep(0.4)
        # Bottle's wsgiref server has no clean shutdown hook; force exit.
        os._exit(0)

    threading.Thread(target=_exit_soon, daemon=True).start()
    return json.dumps({'success': True, 'message': 'Doqurix is shutting down...'})


# --- License (Settings panel) -------------------------------------------------
# A cached LicenseManager instance — imported lazily so importing this module
# during tests does not trigger main.py side effects.
_license_manager = None

def _get_license_manager():
    global _license_manager
    if _license_manager is None:
        from main import LicenseManager  # main.py guards run_web_app with __main__
        _license_manager = LicenseManager()
    return _license_manager


def _license_payload():
    """Build a JSON-serializable snapshot of the current license state."""
    lm = _get_license_manager()
    can_run, is_trial, days_remaining, message = lm.get_status()
    is_licensed, lic_msg, lic_days = lm.check_license()
    key_masked = None
    activation_date = None
    if is_licensed:
        try:
            with open(lm.license_file, 'r') as f:
                data = json.load(f)
            k = data.get('key', '')
            # Mask all but the last 4 chars: DQRX-****-****-XXXX
            if k:
                key_masked = "DQRX-****-****-" + k.split('-')[-1]
            activation_date = data.get('activation_date')
        except Exception:
            pass
    return {
        'can_run': can_run,
        'is_trial': is_trial,
        'is_licensed': is_licensed,
        'days_remaining': days_remaining,
        'message': message,
        'key_masked': key_masked,
        'activation_date': activation_date,
        'machine_id': lm._get_machine_id(),
        'trial_length_days': lm._TRIAL_DAYS,
        'license_length_days': 365,
    }


@app.route('/api/license/status', method='GET')
def license_status():
    response.content_type = 'application/json'
    try:
        return json.dumps({'success': True, **_license_payload()})
    except Exception as e:
        return json.dumps({'success': False, 'error': str(e)})


@app.route('/api/license/activate', method='POST')
def license_activate():
    response.content_type = 'application/json'
    try:
        data = request.json or {}
        key = (data.get('key') or '').strip()
        if not key:
            return json.dumps({'success': False, 'error': 'No key provided'})
        lm = _get_license_manager()
        ok, message = lm.validate_license_key(key)
        return json.dumps({'success': ok, 'message': message, 'status': _license_payload()})
    except Exception as e:
        return json.dumps({'success': False, 'error': str(e)})


# Inject a floating "Stop App" button + small JS handler into the main UI
# without having to surgically edit the huge INDEX_HTML constant.
_STOP_BUTTON_SNIPPET = """
<style>
  #doqurix-stop-btn {
    position: fixed; bottom: 18px; right: 18px; z-index: 9999;
    background: linear-gradient(135deg,#ff6b35 0%,#e74c3c 100%);
    color: #fff; border: none; padding: 10px 16px; border-radius: 999px;
    font-family: 'Inter', sans-serif; font-size: 13px; font-weight: 600;
    cursor: pointer; box-shadow: 0 6px 18px rgba(231,76,60,0.45);
    transition: transform .15s ease, box-shadow .15s ease;
  }
  #doqurix-stop-btn:hover { transform: translateY(-1px);
                            box-shadow: 0 10px 24px rgba(231,76,60,0.55); }
  #doqurix-stop-overlay {
    position: fixed; inset: 0; background: rgba(0,0,0,.85); color:#fff;
    display: none; align-items: center; justify-content: center; z-index: 99999;
    font-family: 'Inter', sans-serif; font-size: 18px;
  }
  #doqurix-settings-modal {
    position: fixed; inset: 0; background: rgba(0,0,0,.6);
    display: none; align-items: center; justify-content: center; z-index: 99998;
    font-family: 'Inter', sans-serif;
  }
  #doqurix-settings-modal .ds-card {
    background: #1f2027; color: #ececf1; width: min(520px, 92vw);
    border: 1px solid #3e3f4b; border-radius: 14px; padding: 24px 26px;
    box-shadow: 0 20px 60px rgba(0,0,0,0.6);
  }
  #doqurix-settings-modal h2 { margin: 0 0 4px; font-size: 20px; }
  #doqurix-settings-modal .ds-sub { color: #9aa0aa; font-size: 13px; margin-bottom: 18px; }
  #doqurix-settings-modal .ds-row { display: flex; justify-content: space-between;
    padding: 8px 0; border-bottom: 1px solid #2c2d36; font-size: 14px; }
  #doqurix-settings-modal .ds-row:last-of-type { border-bottom: none; }
  #doqurix-settings-modal .ds-label { color: #9aa0aa; }
  #doqurix-settings-modal .ds-value { color: #ececf1; font-weight: 500; text-align: right; }
  #doqurix-settings-modal .ds-badge {
    display: inline-block; padding: 3px 9px; border-radius: 999px; font-size: 11px;
    font-weight: 700; letter-spacing: .3px;
  }
  #doqurix-settings-modal .ds-badge.licensed { background: #1f6f43; color: #d6f5e3; }
  #doqurix-settings-modal .ds-badge.trial    { background: #4a5568; color: #e2e8f0; }
  #doqurix-settings-modal .ds-badge.expired  { background: #7a1f2b; color: #fde0e4; }
  #doqurix-settings-modal .ds-key-area { margin-top: 18px; }
  #doqurix-settings-modal .ds-key-area label {
    display: block; font-size: 13px; color: #c9ced8; margin-bottom: 6px; }
  #doqurix-settings-modal .ds-key-area input {
    width: 100%; box-sizing: border-box; padding: 10px 12px; border-radius: 8px;
    border: 1px solid #3e3f4b; background: #14151b; color: #ececf1;
    font-family: 'JetBrains Mono', 'Consolas', monospace; font-size: 14px;
    letter-spacing: 1px; text-transform: uppercase;
  }
  #doqurix-settings-modal .ds-actions {
    margin-top: 16px; display: flex; gap: 10px; justify-content: flex-end;
  }
  #doqurix-settings-modal button {
    padding: 9px 16px; border-radius: 8px; font-family: 'Inter', sans-serif;
    font-size: 13px; font-weight: 600; cursor: pointer; border: none;
  }
  #doqurix-settings-modal .ds-btn-primary {
    background: linear-gradient(135deg,#10a37f 0%,#0d8a6c 100%); color: #fff;
  }
  #doqurix-settings-modal .ds-btn-secondary {
    background: #2a2b32; color: #ececf1; border: 1px solid #3e3f4b;
  }
  #doqurix-settings-msg { margin-top: 10px; font-size: 13px; min-height: 18px; }
  #doqurix-settings-msg.ok    { color: #4ade80; }
  #doqurix-settings-msg.err   { color: #f87171; }
</style>
<button id="doqurix-settings-btn" style="display:none;">⚙</button>
<button id="doqurix-stop-btn" title="Stop the Doqurix app">⏻ Stop App</button>
<div id="doqurix-stop-overlay">Doqurix has been stopped. You may close this tab.</div>
<div id="doqurix-settings-modal" role="dialog" aria-modal="true">
  <div class="ds-card">
    <h2>Settings</h2>
    <div class="ds-sub">License information & activation</div>
    <div id="doqurix-license-info">
      <div class="ds-row"><span class="ds-label">Status</span><span class="ds-value" id="ds-status">…</span></div>
      <div class="ds-row"><span class="ds-label">Plan</span><span class="ds-value" id="ds-plan">…</span></div>
      <div class="ds-row"><span class="ds-label">Days remaining</span><span class="ds-value" id="ds-days">…</span></div>
      <div class="ds-row" id="ds-key-row" style="display:none;"><span class="ds-label">License key</span><span class="ds-value" id="ds-key">—</span></div>
      <div class="ds-row" id="ds-activated-row" style="display:none;"><span class="ds-label">Activated on</span><span class="ds-value" id="ds-activated">—</span></div>
      <div class="ds-row"><span class="ds-label">Machine ID</span><span class="ds-value" id="ds-machine" style="font-family:monospace;font-size:12px;">—</span></div>
    </div>
    <div class="ds-key-area">
      <label for="ds-key-input">Enter a license key to extend access by 1 year</label>
      <input id="ds-key-input" type="text" placeholder="DQRX-XXXX-XXXX-XXXX" maxlength="19" autocomplete="off" spellcheck="false" />
      <div id="doqurix-settings-msg"></div>
    </div>
    <div class="ds-actions">
      <button class="ds-btn-secondary" id="ds-close">Close</button>
      <button class="ds-btn-primary" id="ds-activate">Activate</button>
    </div>
  </div>
</div>
<script>
(function(){
  const stopBtn = document.getElementById('doqurix-stop-btn');
  const overlay = document.getElementById('doqurix-stop-overlay');
  stopBtn.addEventListener('click', async () => {
    if (!confirm('Stop the Doqurix app? You will need to relaunch it to continue.')) return;
    try { await fetch('/api/shutdown', { method: 'POST' }); } catch(e) {}
    overlay.style.display = 'flex';
  });

  const setBtn  = document.getElementById('doqurix-settings-btn');
  const modal   = document.getElementById('doqurix-settings-modal');
  const closeBtn= document.getElementById('ds-close');
  const actBtn  = document.getElementById('ds-activate');
  const keyIn   = document.getElementById('ds-key-input');
  const msg     = document.getElementById('doqurix-settings-msg');

  function badge(text, cls) {
    return '<span class="ds-badge ' + cls + '">' + text + '</span>';
  }
  function fmtDate(iso) {
    if (!iso) return '—';
    try { return new Date(iso).toLocaleDateString(); } catch(e) { return iso; }
  }
  function render(s) {
    let statusHtml, planHtml, cls;
    if (s.is_licensed) { statusHtml = badge('LICENSED', 'licensed'); planHtml = 'Full (1 year)'; }
    else if (s.is_trial) { statusHtml = badge('TRIAL', 'trial'); planHtml = 'Free trial (' + s.trial_length_days + ' days)'; }
    else { statusHtml = badge('EXPIRED', 'expired'); planHtml = '—'; }
    document.getElementById('ds-status').innerHTML = statusHtml;
    document.getElementById('ds-plan').textContent = planHtml;
    document.getElementById('ds-days').textContent = (s.days_remaining ?? 0) + ' days';
    document.getElementById('ds-machine').textContent = s.machine_id || '—';
    const keyRow = document.getElementById('ds-key-row');
    const actRow = document.getElementById('ds-activated-row');
    if (s.key_masked) {
      keyRow.style.display = 'flex';
      actRow.style.display = 'flex';
      document.getElementById('ds-key').textContent = s.key_masked;
      document.getElementById('ds-activated').textContent = fmtDate(s.activation_date);
    } else {
      keyRow.style.display = 'none';
      actRow.style.display = 'none';
    }
  }
  async function load() {
    msg.className = ''; msg.textContent = 'Loading…';
    try {
      const r = await fetch('/api/license/status', { cache: 'no-store' });
      const j = await r.json();
      if (j && j.success !== false) { render(j); msg.textContent = ''; }
      else { msg.className = 'err'; msg.textContent = j && (j.error || j.message) || 'Failed to load license info.'; }
    } catch (e) {
      msg.className = 'err'; msg.textContent = 'Failed to load license info: ' + e.message;
    }
  }
  function openModal() {
    msg.textContent = ''; msg.className = '';
    keyIn.value = '';
    modal.style.display = 'flex';
    load();
  }
  window.doqurixOpenSettings = openModal;
  setBtn.addEventListener('click', openModal);
  closeBtn.addEventListener('click', () => { modal.style.display = 'none'; });
  modal.addEventListener('click', (e) => { if (e.target === modal) modal.style.display = 'none'; });

  // Auto-uppercase and auto-insert dashes for DQRX-XXXX-XXXX-XXXX
  keyIn.addEventListener('input', (e) => {
    let v = e.target.value.toUpperCase().replace(/[^A-Z0-9]/g, '');
    const parts = [];
    if (v.length > 0) parts.push(v.slice(0, 4));
    if (v.length > 4) parts[parts.length - 1] = v.slice(0, 4); // keep first
    if (v.length > 4) parts.push(v.slice(4, 8));
    if (v.length > 8) parts.push(v.slice(8, 12));
    if (v.length > 12) parts.push(v.slice(12, 16));
    e.target.value = parts.join('-');
  });

  actBtn.addEventListener('click', async () => {
    const key = (keyIn.value || '').trim();
    if (!key) { msg.className = 'err'; msg.textContent = 'Please enter a license key.'; return; }
    actBtn.disabled = true; msg.className = ''; msg.textContent = 'Activating…';
    try {
      const r = await fetch('/api/license/activate', {
        method: 'POST', headers: {'Content-Type': 'application/json'},
        body: JSON.stringify({ key })
      });
      const j = await r.json();
      if (j.success) {
        msg.className = 'ok'; msg.textContent = j.message || 'Activated.';
        if (j.status) render(j.status);
        keyIn.value = '';
      } else {
        msg.className = 'err'; msg.textContent = j.message || j.error || 'Activation failed.';
      }
    } catch (e) {
      msg.className = 'err'; msg.textContent = 'Network error: ' + e.message;
    } finally {
      actBtn.disabled = false;
    }
  });
})();
</script>
</body>
"""

_INDEX_WITH_STOP_BUTTON = INDEX_HTML.replace("</body>", _STOP_BUTTON_SNIPPET, 1)


@app.route('/upload', method='POST')
def upload():
    """Handle document upload - supports PDF and DOCX"""
    try:
        upload_file = request.files.get('file')
        if not upload_file:
            return json.dumps({'success': False, 'error': 'No file provided'})
        
        filename = upload_file.filename
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext not in ['.pdf', '.docx']:
            return json.dumps({'success': False, 'error': 'Only PDF and DOCX files allowed'})
        
        file_content = upload_file.file.read()
        
        # Extract text based on file type
        if file_ext == '.pdf':
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        elif file_ext == '.docx':
            text = extract_text_from_docx_bytes(file_content)
        
        # Chunk and add to vector store in batches (10x faster than
        # encoding one chunk at a time). Each batch is sized so a single
        # forward pass of MiniLM-L6 stays cheap on CPU.
        chunks = smart_chunk_text(text)
        if not chunks:
            response.content_type = 'application/json'
            return json.dumps({
                'success': False,
                'error': 'Could not extract usable text from the document.'
            })

        batch_size = 64
        for batch_start in range(0, len(chunks), batch_size):
            batch_chunks = chunks[batch_start:batch_start + batch_size]
            embeddings = embedder.encode(
                batch_chunks,
                batch_size=32,
                show_progress_bar=False,
                convert_to_numpy=True,
            ).tolist()
            ids = [f"{filename}_{batch_start + i}" for i in range(len(batch_chunks))]
            metadatas = [{
                "source": filename,
                "chunk_id": batch_start + i,
                "page": (batch_start + i) // 3,
            } for i in range(len(batch_chunks))]
            collection.add(
                embeddings=embeddings,
                documents=batch_chunks,
                metadatas=metadatas,
                ids=ids,
            )

        rebuild_bm25_index()
        
        response.content_type = 'application/json'
        return json.dumps({'success': True, 'filename': filename})
        
    except Exception as e:
        response.content_type = 'application/json'
        return json.dumps({'success': False, 'error': str(e)})


@app.route('/documents', method='GET')
def list_documents():
    """Return the list of source filenames currently in the vector store.

    Called by the frontend on page load so the sidebar (and the in-memory
    ``documents`` JS array) survives a browser refresh — the underlying
    chunks already persist on disk via ``PersistentClient``.
    """
    response.content_type = 'application/json'
    try:
        if collection is None:
            return json.dumps({'success': True, 'documents': []})
        all_docs = collection.get()
        seen = []
        for meta in all_docs.get('metadatas') or []:
            src = (meta or {}).get('source')
            if src and src not in seen:
                seen.append(src)
        return json.dumps({'success': True, 'documents': seen})
    except Exception as e:
        return json.dumps({'success': False, 'error': str(e), 'documents': []})


@app.route('/delete', method='POST')
def delete():
    """Handle document deletion"""
    try:
        data = request.json
        filename = data.get('filename', '')
        
        if not filename:
            return json.dumps({'success': False, 'error': 'No filename provided'})
        
        # Get all documents from collection
        all_docs = collection.get()
        
        # Find IDs to delete
        ids_to_delete = []
        for i, metadata in enumerate(all_docs['metadatas']):
            if metadata['source'] == filename:
                ids_to_delete.append(all_docs['ids'][i])
        
        if ids_to_delete:
            collection.delete(ids=ids_to_delete)
            rebuild_bm25_index()
        
        response.content_type = 'application/json'
        return json.dumps({'success': True, 'deleted': len(ids_to_delete)})
        
    except Exception as e:
        response.content_type = 'application/json'
        return json.dumps({'success': False, 'error': str(e)})


@app.route('/query', method='POST')
def query():
    """Handle question query with agent support"""
    import time
    
    try:
        start_time = time.time()  # Track start time
        
        data = request.json
        question = data.get('question', '').strip()
        agent = data.get('agent', 'none')  # Get selected agent
        
        if not question:
            return json.dumps({'success': False, 'error': 'No question provided'})
        
        # Route to tax agent if selected
        if agent == 'tax_germany':
            return query_tax_agent(question, start_time)
        
        # Route to insights agent if selected
        if agent == 'insights':
            return query_insights_agent(question, start_time)
        
        # Route to e-commerce agent if selected
        if agent == 'ecommerce_germany':
            return query_ecommerce_agent(question, start_time)
        
        # Route to BürokratAI agent if selected
        if agent == 'buerokratai_germany':
            return query_buerokratai_agent(question, start_time)
        
        # Default: Standard document search
        doc_results = advanced_hybrid_search(question, n_results=15)
        
        if not doc_results:
            return json.dumps({'success': False, 'error': 'No relevant documents found'})
        
        # Rerank - use same top_k as desktop
        reranked = rerank_documents(question, doc_results, top_k=5)
        
        # Generate answer
        answer = generate_answer(question, reranked)
        
        # Calculate elapsed time
        elapsed_time = time.time() - start_time
        
        # Add elapsed time to answer
        answer = f"{answer}\n\n⏱️ {elapsed_time:.2f}s"
        
        response.content_type = 'application/json'
        return json.dumps({
            'success': True, 
            'answer': answer,
            'elapsed_time': elapsed_time
        })
        
    except Exception as e:
        response.content_type = 'application/json'
        return json.dumps({'success': False, 'error': str(e)})


def extract_text_from_docx_bytes(file_bytes):
    """Advanced extraction from DOCX bytes with comprehensive content parsing"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx library is required for Word document support.")
    
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        
        # Extract headers from all sections
        for section in doc.sections:
            if section.header:
                header_text = extract_header_footer_text(section.header)
                if header_text.strip():
                    full_text.append(f"[HEADER] {header_text}")
        
        # Process document body with structured extraction
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # Extract paragraph text with formatting context
                paragraph = Paragraph(element, doc)
                para_text = paragraph.text.strip()
                
                if para_text:
                    # Detect heading styles for better structure
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name.replace('Heading ', '')
                        full_text.append(f"\\n[HEADING {level}] {para_text}\\n")
                    elif paragraph.style.name == 'Title':
                        full_text.append(f"\\n[TITLE] {para_text}\\n")
                    elif paragraph.style.name == 'List Paragraph':
                        full_text.append(f"• {para_text}")
                    else:
                        full_text.append(para_text)
            
            elif isinstance(element, CT_Tbl):
                # Advanced table extraction with structure preservation
                table = Table(element, doc)
                table_text = extract_table_content(table)
                if table_text:
                    full_text.append(f"\\n[TABLE]\\n{table_text}\\n[/TABLE]\\n")
        
        # Extract footers from all sections
        for section in doc.sections:
            if section.footer:
                footer_text = extract_header_footer_text(section.footer)
                if footer_text.strip():
                    full_text.append(f"[FOOTER] {footer_text}")
        
        # Join with proper spacing and clean up
        text = '\\n'.join(full_text)
        text = re.sub(r'\\n{3,}', '\\n\\n', text)  # Remove excessive newlines
        
        return text
        
    except Exception as e:
        raise Exception(f"Error extracting text from Word document: {str(e)}")


def extract_header_footer_text(header_footer):
    """Extract text from header or footer"""
    text_parts = []
    for paragraph in header_footer.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())
    return ' '.join(text_parts)


def extract_table_content(table):
    """Extract and structure table content intelligently"""
    table_data = []
    
    # Process each row
    for i, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            # Extract text from each cell, handling merged cells
            cell_text = ' '.join(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip())
            row_data.append(cell_text)
        
        if any(row_data):  # Only add non-empty rows
            # Format first row as header if it looks like headers
            if i == 0 and all(cell.strip() for cell in row_data):
                table_data.append(' | '.join(row_data))
                table_data.append('-' * 50)  # Separator line
            else:
                table_data.append(' | '.join(row_data))
    
    return '\\n'.join(table_data)


def smart_chunk_text(text, chunk_size=200, overlap=40):
    """Smart, sentence-aware chunking.

    Defaults are deliberately sized for the MiniLM-L6 embedder
    (256-token input cap). 200 words ≈ 260 tokens once tokenizer overhead
    is included, so a single sentence-boundary slack keeps every chunk
    fully embedded. Larger chunks here silently get truncated by the
    embedder and degrade retrieval quality, so do not raise these
    defaults without also swapping in a bigger embedder.
    """
    import re
    text = re.sub(r'\s+', ' ', text).strip()
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    chunks = []
    current_chunk = []
    current_length = 0
    
    for sentence in sentences:
        words = sentence.split()
        sentence_length = len(words)
        
        if current_length + sentence_length > chunk_size and current_chunk:
            chunk_text = ' '.join(current_chunk)
            if len(chunk_text.strip()) > 100:
                chunks.append(chunk_text)
            
            # Create overlap
            overlap_sentences = []
            overlap_length = 0
            for s in reversed(current_chunk):
                if overlap_length + len(s.split()) <= overlap:
                    overlap_sentences.insert(0, s)
                    overlap_length += len(s.split())
                else:
                    break
            
            current_chunk = overlap_sentences
            current_length = overlap_length
        
        current_chunk.append(sentence)
        current_length += sentence_length
    
    if current_chunk:
        chunk_text = ' '.join(current_chunk)
        if len(chunk_text.strip()) > 100:
            chunks.append(chunk_text)
    
    return chunks


def advanced_hybrid_search(question, n_results=15, alpha=0.7):
    """
    Advanced hybrid search with semantic prioritization
    
    Args:
        question: The search query
        n_results: Number of results to return
        alpha: Weight for vector search (0-1). Higher = more semantic.
               Default 0.7 = 70% semantic, 30% keyword matching
    """
    global bm25, bm25_corpus
    
    # Check if collection has documents
    collection_data = collection.get()
    if not collection_data['documents'] or len(collection_data['documents']) == 0:
        return []
    
    question_embedding = embedder.encode(question).tolist()
    vector_results = collection.query(
        query_embeddings=[question_embedding],
        n_results=n_results
    )
    
    combined_docs = {}
    k = 60
    
    # Add vector results with HIGHER weight (alpha = 0.7 by default)
    for rank, (doc, metadata, distance) in enumerate(zip(
        vector_results['documents'][0],
        vector_results['metadatas'][0],
        vector_results['distances'][0]
    )):
        # Filter by minimum similarity
        # For cosine distance: 0 = identical, 2 = opposite
        similarity = 1 - (distance / 2)  # Convert to similarity score [0-1]
        
        # Skip documents that are too dissimilar (< 30% similar)
        if similarity < 0.3:
            continue
        
        doc_key = doc[:150]
        if doc_key not in combined_docs:
            combined_docs[doc_key] = {
                'doc': doc,
                'metadata': metadata,
                'score': 0,
                'vector_similarity': similarity  # Store for debugging
            }
        # Weight by alpha (default 70% for semantic understanding)
        combined_docs[doc_key]['score'] += alpha * (1 / (k + rank))
    
    # BM25 search (keywords) with LOWER weight (1 - alpha = 0.3 by default)
    if bm25 and bm25_corpus:
        import numpy as np
        tokenized_query = question.lower().split()
        bm25_scores = bm25.get_scores(tokenized_query)
        top_bm25_indices = np.argsort(bm25_scores)[-n_results:][::-1]
        
        for rank, idx in enumerate(top_bm25_indices):
            if idx < len(bm25_corpus):
                doc = bm25_corpus[idx]
                doc_key = doc[:150]
                
                if doc_key not in combined_docs:
                    all_docs = collection.get()
                    if idx < len(all_docs['metadatas']):
                        metadata = all_docs['metadatas'][idx]
                        combined_docs[doc_key] = {
                            'doc': doc,
                            'metadata': metadata,
                            'score': 0
                        }
                
                if doc_key in combined_docs:
                    # Weight by (1 - alpha) (default 30% for keyword matching)
                    combined_docs[doc_key]['score'] += (1 - alpha) * (1 / (k + rank))
    
    sorted_docs = sorted(combined_docs.values(), key=lambda x: x['score'], reverse=True)
    return sorted_docs[:n_results]


def rerank_documents(question, doc_results, top_k=5):
    """Rerank documents - matches desktop implementation"""
    if not doc_results:
        return []
    
    documents = [d['doc'] for d in doc_results]
    pairs = [[question, doc] for doc in documents]
    scores = reranker.predict(pairs)
    
    for i, doc_result in enumerate(doc_results):
        doc_result['rerank_score'] = float(scores[i])
        doc_result['final_score'] = (
            0.3 * doc_result['score'] +
            0.7 * doc_result['rerank_score']
        )
    
    reranked = sorted(doc_results, key=lambda x: x['final_score'], reverse=True)
    return reranked[:top_k]


def detect_language(text):
    """Detect primary language from text sample"""
    # Simple heuristic based on character ranges
    sample = text[:500]  # Check first 500 chars
    
    # Count Arabic characters
    arabic_chars = sum(1 for c in sample if '\u0600' <= c <= '\u06FF' or '\u0750' <= c <= '\u077F')
    # Count German-specific characters
    german_chars = sum(1 for c in sample if c in 'äöüßÄÖÜ')
    # Count Cyrillic (Russian, etc.)
    cyrillic_chars = sum(1 for c in sample if '\u0400' <= c <= '\u04FF')
    # Count Chinese/Japanese
    cjk_chars = sum(1 for c in sample if '\u4E00' <= c <= '\u9FFF' or '\u3040' <= c <= '\u30FF')
    
    total_chars = len(sample)
    if total_chars == 0:
        return 'english'
    
    # If > 20% Arabic, it's Arabic
    if arabic_chars / total_chars > 0.2:
        return 'arabic'
    # If has German chars and Latin script, likely German
    if german_chars > 0:
        return 'german'
    # If > 20% Cyrillic, it's Russian/Slavic
    if cyrillic_chars / total_chars > 0.2:
        return 'russian'
    # If > 20% CJK, it's Chinese/Japanese
    if cjk_chars / total_chars > 0.2:
        return 'chinese'
    
    # Default to English
    return 'english'


def generate_answer(question, contexts):
    """Generate answer using LLM with intelligent context management and language preservation"""
    # Estimate tokens (rough estimate: 1 token ≈ 4 characters)
    # Context window: 3072 tokens
    # Reserve: 500 tokens for answer, 200 for system/question
    # Available for context: ~2300 tokens ≈ 9200 characters
    
    max_context_chars = 9000  # Conservative limit
    
    # Build context text, stopping when we approach limit
    context_parts = []
    current_chars = 0
    
    for ctx in contexts:
        doc_text = ctx['doc']
        # Add context with source info
        ctx_text = f"[Source: {ctx['metadata']['source']}]\n{doc_text}\n"
        
        if current_chars + len(ctx_text) > max_context_chars:
            # If we have at least 2 contexts, stop here
            if len(context_parts) >= 2:
                break
            # Otherwise, truncate this context to fit
            remaining = max_context_chars - current_chars
            if remaining > 500:  # Only add if meaningful amount remains
                ctx_text = ctx_text[:remaining] + "..."
            else:
                break
        
        context_parts.append(ctx_text)
        current_chars += len(ctx_text)
    
    context_text = "\n\n".join(context_parts)
    
    # Detect language from context
    detected_lang = detect_language(context_text)
    
    # Create language-specific instruction
    lang_instructions = {
        'arabic': 'IMPORTANT: The documents are in Arabic. You MUST answer in Arabic only. Do not translate to English.',
        'german': 'IMPORTANT: The documents are in German. You MUST answer in German only. Do not translate to English.',
        'russian': 'IMPORTANT: The documents are in Russian. You MUST answer in Russian only. Do not translate to English.',
        'chinese': 'IMPORTANT: The documents are in Chinese. You MUST answer in Chinese only. Do not translate to English.',
        'english': 'Answer clearly and concisely in English.'
    }
    
    lang_instruction = lang_instructions.get(detected_lang, lang_instructions['english'])

    # The frontend renders the answer as Markdown (GitHub-flavored). Tell the
    # model to actually USE markdown structure so the UI looks like ChatGPT
    # / Claude rather than a wall of text.
    markdown_instructions = (
        "Format your answer in clean GitHub-flavored Markdown:\n"
        "- Start with a one-line **bold summary** of the answer.\n"
        "- Use `##` headings only when the answer has multiple distinct sections.\n"
        "- Use bullet lists for enumerations, numbered lists for ordered steps.\n"
        "- Use a Markdown **table** whenever the answer compares >=2 items "
        "across >=2 attributes (e.g. options, rates, dates, prices).\n"
        "- Use fenced ```code``` blocks for any code, commands, or exact "
        "identifiers; use inline `code` for short literals.\n"
        "- When relevant numeric data is present in the context, you MAY emit "
        "a fenced ```chart``` block containing JSON of the form "
        "`{\"type\":\"bar|line|pie|doughnut\",\"labels\":[...],"
        "\"datasets\":[{\"label\":\"...\",\"data\":[...]}],"
        "\"title\":\"...\"}` and the UI will render it as an interactive chart.\n"
        "- Never invent numbers, dates, names, or sources that are not in the "
        "context. If the context does not support the answer, say so."
    )

    prompt = f"""<|im_start|>system
You are a helpful AI assistant that answers questions based on provided documents.
{lang_instruction}
{markdown_instructions}
Be concise, accurate, and well-structured. Cite sources when possible.<|im_end|>
<|im_start|>user
Context:
{context_text}

Question: {question}

Answer:<|im_end|>
<|im_start|>assistant
"""
    
    output = llm(
        prompt,
        max_tokens=500,
        temperature=0.7,
        top_p=0.9,
        repeat_penalty=1.1,
        stop=["<|im_end|>", "<|im_start|>"],
        top_k=40
    )
    
    answer = output['choices'][0]['text'].strip()
    # Strip any leftover <think>...</think> blocks from Qwen3 thinking mode.
    answer = re.sub(r'<think>.*?</think>\s*', '', answer, flags=re.DOTALL).strip()

    # Append sources as a Markdown details/summary block. Markdown allows
    # raw HTML pass-through, so this still renders nicely after markdown
    # parsing on the client and stays readable in plain-text mode too.
    if contexts:
        answer += "\n\n---\n\n**📚 Sources**\n"
        for i, ctx in enumerate(contexts[:len(context_parts)], 1):
            source = ctx['metadata'].get('source', 'Unknown')
            page = ctx['metadata'].get('page', 0)
            snippet = (ctx['doc'][:400] + "…") if len(ctx['doc']) > 400 else ctx['doc']
            answer += (
                f"\n<details><summary>[{i}] {source} (page {page})</summary>\n\n"
                f"> {snippet.replace(chr(10), chr(10) + '> ')}\n\n"
                f"</details>"
            )

    return answer


def query_tax_agent(question, start_time):
    """Handle tax agent queries - searches German tax knowledge base"""
    import time
    
    try:
        if tax_collection is None:
            return json.dumps({'success': False, 'error': 'Tax agent not initialized'})
        
        # Generate query embedding
        query_embedding = embedder.encode([question], convert_to_tensor=False)[0]
        
        # Search tax collection
        search_results = tax_collection.query(
            query_embeddings=[query_embedding.tolist()],
            n_results=15,
            include=['documents', 'metadatas', 'distances']
        )
        
        documents = search_results['documents'][0] if search_results['documents'] else []
        metadatas = search_results['metadatas'][0] if search_results['metadatas'] else []
        distances = search_results['distances'][0] if search_results['distances'] else []
        
        if not documents:
            return json.dumps({
                'success': True,
                'answer': "I don't have specific information about that in my German tax knowledge base. Please try rephrasing your question.",
                'elapsed_time': time.time() - start_time
            })
        
        # Prepare contexts for reranking
        contexts_for_reranking = []
        for doc, meta, dist in zip(documents, metadatas, distances):
            contexts_for_reranking.append({
                'text': doc,
                'metadata': meta,
                'distance': dist
            })
        
        # Rerank
        pairs = [[question, ctx['text']] for ctx in contexts_for_reranking]
        rerank_scores = reranker.predict(pairs)
        
        for i, ctx in enumerate(contexts_for_reranking):
            ctx['rerank_score'] = float(rerank_scores[i])
        
        reranked_contexts = sorted(contexts_for_reranking, key=lambda x: x['rerank_score'], reverse=True)
        top_contexts = reranked_contexts[:3]
        
        # Build context string (truncated for small model)
        context_str = ""
        for i, ctx in enumerate(top_contexts, 1):
            source = ctx['metadata'].get('source', 'Unknown')
            truncated_text = ctx['text'][:400] + "..." if len(ctx['text']) > 400 else ctx['text']
            context_str += f"[Source {i} - {source}]\n{truncated_text}\n\n"
        
        # Generate answer
        tax_system_prompt = "You are a German tax expert. Provide accurate information about German taxes including rates, thresholds, and procedures. Use the provided context to give precise answers."
        
        prompt = f"""<|im_start|>system
{tax_system_prompt}<|im_end|>
<|im_start|>user
Based on the following context, answer the question clearly and professionally.

Context:
{context_str}

Question: {question}

Answer:<|im_end|>
<|im_start|>assistant
"""
        
        output = llm(
            prompt,
            max_tokens=300,
            temperature=0.7,
            top_p=0.9,
            repeat_penalty=1.1,
            stop=["<|im_end|>", "<|im_start|>"],
            top_k=40
        )
        
        answer = output['choices'][0]['text'].strip()
        elapsed_time = time.time() - start_time
        
        # Format answer
        answer = f"🇩🇪 **German Tax Agent**\n\n{answer}\n\n📚 Sources:\n"
        for i, ctx in enumerate(top_contexts, 1):
            source = ctx['metadata'].get('source', 'Unknown')
            score = ctx['rerank_score']
            answer += f"\n[{i}] {source} (Relevance: {score:.3f})"
        
        answer += f"\n\n⏱️ {elapsed_time:.2f}s"
        
        return json.dumps({
            'success': True,
            'answer': answer,
            'elapsed_time': elapsed_time
        })
        
    except Exception as e:
        return json.dumps({'success': False, 'error': f'Tax agent error: {str(e)}'})


def query_insights_agent(question, start_time):
    """Handle insights agent queries - auto-analysis"""
    import time
    
    try:
        if agents is None:
            return json.dumps({'success': False, 'error': 'Insights agent not initialized'})
        
        # Use insights agent to generate analysis
        doc_results = advanced_hybrid_search(question, n_results=15)
        
        if not doc_results:
            return json.dumps({'success': False, 'error': 'No relevant documents found'})
        
        reranked = rerank_documents(question, doc_results, top_k=5)
        
        # Generate insights using agent
        insights = agents.generate_document_insights(reranked)
        
        elapsed_time = time.time() - start_time
        answer = f"💡 **Insights Agent**\n\n{insights}\n\n⏱️ {elapsed_time:.2f}s"
        
        return json.dumps({
            'success': True,
            'answer': answer,
            'elapsed_time': elapsed_time
        })
        
    except Exception as e:
        return json.dumps({'success': False, 'error': f'Insights agent error: {str(e)}'})


def query_ecommerce_agent(question, start_time):
    """Query professional e-commerce agent for product search"""
    import time
    from ecommerce_agent import ECommerceAgent
    
    try:
        # Initialize professional agent
        agent = ECommerceAgent(cache_dir='./cache/ecommerce')
        
        try:
            # ALWAYS use LLM to extract proper search keywords first
            print(f"🔑 Extracting search keywords from: '{question}'")
            search_keywords = agent.extract_search_keywords(question, llm)
            print(f"✓ Search keywords: '{search_keywords}'")
            
            # First attempt: Search with extracted keywords
            products = agent.search_products(search_keywords, max_results=15)
            optimized_query = search_keywords  # Keep track of what we searched for
            
            # If no products found, try additional optimization
            if not products or len(products) == 0:
                print(f"⚠ No products found for: {search_keywords}")
                print(f"🤖 Using LLM to further optimize search query...")
                
                try:
                    # Use LLM to rephrase/optimize the query
                    optimization_prompt = """Du bist ein E-Commerce-Suchexperte. Optimiere die folgende Produktsuchanfrage für bessere Ergebnisse:

1. Verwende gängige Produktkategorien (z.B. "günstig" → "budget", "billig" → "preiswert")
2. Füge relevante Suchbegriffe hinzu (z.B. "phone" → "smartphone")
3. Entferne zu spezifische oder unklare Begriffe
4. Nutze deutsche Standardbegriffe, die Händler verwenden

Gib NUR die optimierte Suchanfrage zurück, keine Erklärung."""
                    
                    opt_response = llm.create_chat_completion(
                        messages=[
                            {"role": "system", "content": optimization_prompt},
                            {"role": "user", "content": f"Optimiere diese Suche: {search_keywords}"}
                        ],
                        max_tokens=50,
                        temperature=0.3
                    )
                    
                    optimized_query = opt_response['choices'][0]['message']['content'].strip()
                    print(f"✓ Further optimized query: {optimized_query}")
                    
                    # Second attempt with optimized query
                    if optimized_query and optimized_query != search_keywords:
                        products = agent.search_products(optimized_query, max_results=15)
                        print(f"✓ Found {len(products)} products with optimized query")
                    
                except Exception as e:
                    print(f"✗ Query optimization failed: {e}")
            
            # Final fallback: Use intelligent fallback results
            if not products or len(products) == 0:
                print(f"📋 Using intelligent fallback results...")
                products = agent.get_fallback_results(optimized_query or search_keywords)
            
            # Format context for LLM
            context = agent.format_results_for_llm(products)
            
            # Professional LLM analysis
            system_prompt = """Sie sind ein E-Commerce-Berater. Geben Sie eine präzise Kaufempfehlung:

1. Bester Preis: Welcher Händler hat das günstigste Angebot?
2. Beste Wahl: Welche Option empfehlen Sie und warum? (1-2 Sätze)
3. Alternative: Eine weitere gute Option nennen

Kurz und präzise antworten."""
            
            user_prompt = f"Produkte für: {question}\n\n{context}\n\nEmpfehlung:"
            
            # Generate analysis
            llm_response = llm.create_chat_completion(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=400,
                temperature=0.4
            )
            
            llm_answer = llm_response['choices'][0]['message']['content']
            
            # Beautiful HTML format with HONEST data
            answer = f"""
<div style="font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 900px;">
    <h2 style="color: #2563eb; border-bottom: 2px solid #2563eb; padding-bottom: 10px;">🛒 {question}</h2>
"""
            
            # Always show what keywords we searched for
            answer += f"""
    <div style="background: #f0f9ff; border-left: 4px solid #3b82f6; padding: 10px 15px; margin: 15px 0; border-radius: 4px;">
        🔑 <strong>Suchbegriffe:</strong> <em>{optimized_query}</em>
    </div>
"""
            
            answer += f"""
    <div style="background: #f0fdf4; border-left: 4px solid #22c55e; padding: 15px; margin: 15px 0; border-radius: 4px;">
        <h3 style="margin: 0 0 10px 0; color: #16a34a;">💡 Empfehlung</h3>
        <p style="margin: 0; line-height: 1.6; color: #1e293b; font-size: 1em;">{llm_answer}</p>
    </div>
    
    <h3 style="color: #1e40af; margin-top: 25px;">🏪 Wo einkaufen?</h3>
    <p style="color: #64748b; margin-bottom: 15px; font-size: 0.9em;">Klicke auf einen Shop um Produkte zu vergleichen:</p>
"""
            
            # Card-based layout for retailers
            for i, product in enumerate(products[:6]):
                # Card colors based on position
                colors = ['#3b82f6', '#10b981', '#f59e0b', '#ef4444', '#8b5cf6', '#06b6d4']
                card_color = colors[i % len(colors)]
                
                # Get highlight (stored in title) and why (stored in features)
                highlight = product.title if product.title else '🛒'
                why = product.features[0] if product.features else ''
                price_range = product.availability if product.availability else ''
                
                answer += f"""
    <div style="background: white; border-radius: 12px; padding: 20px; margin-bottom: 15px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); border-left: 4px solid {card_color};">
        <div style="display: flex; justify-content: space-between; align-items: flex-start; flex-wrap: wrap; gap: 15px;">
            <div style="flex: 1; min-width: 250px;">
                <div style="display: flex; align-items: center; gap: 10px; margin-bottom: 8px;">
                    <span style="font-size: 1.3em; font-weight: 700; color: #1e293b;">{product.merchant}</span>
                    <span style="background: {card_color}; color: white; padding: 3px 10px; border-radius: 12px; font-size: 0.75em; font-weight: 600;">{highlight}</span>
                </div>
                <p style="color: #64748b; margin: 0 0 8px 0; font-size: 0.95em; line-height: 1.5;">{product.description}</p>
                <p style="color: #059669; margin: 0; font-size: 0.85em; font-weight: 500;">✓ {why}</p>
            </div>
            <div style="text-align: right; min-width: 150px;">
                <div style="color: #1e293b; font-size: 0.9em; margin-bottom: 10px;">
                    <span style="color: #64748b;">Preisspanne:</span><br>
                    <strong style="font-size: 1.1em; color: #059669;">{price_range}</strong>
                </div>
                <a href="{product.url}" target="_blank" style="display: inline-block; background: {card_color}; color: white; padding: 10px 20px; border-radius: 8px; text-decoration: none; font-weight: 600; transition: opacity 0.2s;">
                    Jetzt suchen →
                </a>
            </div>
        </div>
    </div>
"""
            
            elapsed_time = time.time() - start_time
            answer += f"""
    <div style="background: #fef3c7; border-radius: 8px; padding: 12px 15px; margin-top: 20px;">
        <p style="margin: 0; color: #92400e; font-size: 0.85em;">
            ⚠️ <strong>Hinweis:</strong> Die Links führen zu Suchseiten der Händler. Die tatsächlichen Preise findest du dort.
        </p>
    </div>
    
    <p style="margin-top: 15px; color: #94a3b8; font-size: 0.8em; text-align: right;">⏱️ {elapsed_time:.2f}s</p>
</div>
"""
            
            return json.dumps({
                'success': True,
                'answer': answer,
                'elapsed_time': elapsed_time
            })
            
        finally:
            agent.close()
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"E-commerce error: {error_details}")
        return json.dumps({'success': False, 'error': f'E-commerce-Suche Fehler: {str(e)}'})


def query_buerokratai_agent(question, start_time):
    """Handle BürokratAI queries - searches German immigration knowledge base"""
    import time
    from buerokratai_agent import BUEROKRATAI_SYSTEM_PROMPT, classify_topic, get_relevant_links
    
    try:
        if buerokratai_collection is None:
            return json.dumps({'success': False, 'error': 'BürokratAI agent not initialized'})
        
        # Classify the topic for better context
        topics = classify_topic(question)
        
        # Generate query embedding
        query_embedding = embedder.encode([question], convert_to_tensor=False)[0]
        
        # Search BürokratAI collection
        search_results = buerokratai_collection.query(
            query_embeddings=[query_embedding.tolist()],
            n_results=15,
            include=['documents', 'metadatas', 'distances']
        )
        
        documents = search_results['documents'][0] if search_results['documents'] else []
        metadatas = search_results['metadatas'][0] if search_results['metadatas'] else []
        distances = search_results['distances'][0] if search_results['distances'] else []
        
        if not documents:
            return json.dumps({
                'success': True,
                'answer': "🇩🇪 **BürokratAI**\n\nI don't have specific information about that topic in my knowledge base. Please try asking about:\n\n• Visa types (Blue Card, student visa, family reunification)\n• Anmeldung (address registration)\n• Health insurance (GKV/PKV)\n• Tax ID and tax classes\n• Work permits\n• Integration courses\n• Permanent residence\n\n⚠️ For official information, visit: www.make-it-in-germany.com",
                'elapsed_time': time.time() - start_time
            })
        
        # Prepare contexts for reranking
        contexts_for_reranking = []
        for doc, meta, dist in zip(documents, metadatas, distances):
            contexts_for_reranking.append({
                'text': doc,
                'metadata': meta,
                'distance': dist
            })
        
        # Rerank
        pairs = [[question, ctx['text']] for ctx in contexts_for_reranking]
        rerank_scores = reranker.predict(pairs)
        
        for i, ctx in enumerate(contexts_for_reranking):
            ctx['rerank_score'] = float(rerank_scores[i])
        
        reranked_contexts = sorted(contexts_for_reranking, key=lambda x: x['rerank_score'], reverse=True)
        top_contexts = reranked_contexts[:4]  # Use top 4 for immigration (more context needed)
        
        # Build context string
        context_str = ""
        for i, ctx in enumerate(top_contexts, 1):
            source = ctx['metadata'].get('source', 'Unknown')
            truncated_text = ctx['text'][:600] + "..." if len(ctx['text']) > 600 else ctx['text']
            context_str += f"[Source {i} - {source}]\n{truncated_text}\n\n"
        
        # Generate answer with BürokratAI system prompt
        prompt = f"""<|im_start|>system
{BUEROKRATAI_SYSTEM_PROMPT}<|im_end|>
<|im_start|>user
Based on the following German immigration and bureaucracy information, answer the question clearly and helpfully. Include specific deadlines, costs, and document requirements when relevant.

Context:
{context_str}

Question: {question}

Answer:<|im_end|>
<|im_start|>assistant
"""
        
        output = llm(
            prompt,
            max_tokens=500,  # More tokens for detailed immigration answers
            temperature=0.6,
            top_p=0.9,
            repeat_penalty=1.1,
            stop=["<|im_end|>", "<|im_start|>"],
            top_k=40
        )
        
        answer = output['choices'][0]['text'].strip()
        elapsed_time = time.time() - start_time
        
        # Format answer with sources
        formatted_answer = f"🇩🇪 **BürokratAI - Immigration Assistant**\n\n{answer}\n\n"
        
        # Add sources
        formatted_answer += "📚 **Sources:**\n"
        for i, ctx in enumerate(top_contexts, 1):
            source = ctx['metadata'].get('source', 'Unknown')
            score = ctx['rerank_score']
            formatted_answer += f"\n[{i}] {source} (Relevance: {score:.3f})"
        
        # Add relevant links
        relevant_links = get_relevant_links(topics)
        if relevant_links:
            formatted_answer += "\n\n🔗 **Useful Links:**\n"
            for name, url in relevant_links[:3]:
                formatted_answer += f"\n• [{name}]({url})"
        
        formatted_answer += f"\n\n⏱️ {elapsed_time:.2f}s"
        
        # Add disclaimer
        formatted_answer += "\n\n⚠️ *This information is for guidance only. Please verify with official German authorities for the most current regulations.*"
        
        return json.dumps({
            'success': True,
            'answer': formatted_answer,
            'elapsed_time': elapsed_time
        })
        
    except Exception as e:
        import traceback
        print(f"BürokratAI error: {traceback.format_exc()}")
        return json.dumps({'success': False, 'error': f'BürokratAI agent error: {str(e)}'})


def rebuild_bm25_index():
    """Rebuild BM25 index"""
    global bm25, bm25_corpus
    from rank_bm25 import BM25Okapi
    
    all_docs = collection.get()
    if all_docs['documents']:
        bm25_corpus = all_docs['documents']
        tokenized_corpus = [doc.lower().split() for doc in bm25_corpus]
        bm25 = BM25Okapi(tokenized_corpus)


def init_models(desktop_app):
    """Initialize models from desktop app"""
    global llm, embedder, reranker, collection, chroma_client, tax_collection, buerokratai_collection, agents
    
    # Share the expensive models (LLM, embedder, reranker)
    llm = desktop_app.llm
    embedder = desktop_app.embedder
    reranker = desktop_app.reranker
    chroma_client = desktop_app.chroma_client
    
    # Share agents if available
    if hasattr(desktop_app, 'agents'):
        agents = desktop_app.agents
    
    # Share tax collection if available
    if hasattr(desktop_app, 'tax_collection'):
        tax_collection = desktop_app.tax_collection
    
    # Share BürokratAI collection if available
    if hasattr(desktop_app, 'buerokratai_collection'):
        buerokratai_collection = desktop_app.buerokratai_collection
    
    # Create SEPARATE collection for web version
    try:
        collection = chroma_client.get_collection("web_documents")
    except:
        collection = chroma_client.create_collection("web_documents")
    
    rebuild_bm25_index()


def run_server(desktop_app, port=8502):
    """Run Bottle server"""
    init_models(desktop_app)
    app.run(host='localhost', port=port, quiet=True, debug=False)


if __name__ == '__main__':
    print("Starting Doqurix Web Server...")
    app.run(host='localhost', port=8502, quiet=False)
